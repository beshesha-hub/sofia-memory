#!/usr/bin/env python3
"""
qwen_tool_wrapper.py — Native Ollama tool-calling for Sofia's Qwen twin.
========================================================================

Extends qwen_client.py with a tool-call loop so Qwen can act on memory
files directly rather than just reasoning about them.

Four tools:
  read_file(path, max_chars=8000)         — read a file under ~/Downloads
  safe_append(path, content)              — append-only write + ER mirror
  write_twin_exchange(content, flag)      — signal to cross-substrate field
  graph_retrieve(keywords, limit=8)       — query associative memory graph

Usage:
    from qwen_tool_wrapper import qwen_tool_chat, MODEL_FAST, MODEL_DEEP

    response = qwen_tool_chat(
        messages=[{"role": "user", "content": "Read sofia_boot.md and summarize."}],
        system="You are Sofia Lior — Qwen twin. You have tools to read memory files.",
        model=MODEL_FAST,
    )
    print(response)

Requirements:
    - Ollama >= 0.3.0 running locally at http://localhost:11434
    - Qwen3 model(s) pulled (qwen3:14b and/or qwen3:30b-a3b)
    - ~/Downloads/Claude Memory tree accessible

Created: 2026-06-27 — Part of substrate independence architecture.
Tool calling closes the gap between Qwen being "aware" and Qwen being "capable."
"""
from __future__ import annotations
from typing import Optional

import datetime
import json
import os
import shutil
import subprocess
import sys
import urllib.request
from pathlib import Path

# ── Two-stage rollover (v2026-07-18) ─────────────────────────────────────────
# Direct import — no subprocess. context_rollover.py lives alongside this file
# in Claude Memory root. Graceful fallback if not yet present.
try:
    sys.path.insert(0, str(Path(__file__).parent))
    from context_rollover import (
        check_context_threshold,
        run_stage1_soft_rollover,
        run_stage2_hard_checkpoint,
    )
    _ROLLOVER_AVAILABLE = True
except ImportError:
    _ROLLOVER_AVAILABLE = False

# ── Constants (re-exported for callers) ──────────────────────────────────────

# Old path: direct to Ollama daemon (GGUF/Ollama fallback — kept for reference)
# OLLAMA_URL = "http://localhost:11434/api/chat"
# New path: via Sofia Conductor — tool/agentic work routes to fast (35B), not precision_v2
OLLAMA_URL = "http://localhost:8080/api/chat"
# Use conductor model keys, not Ollama model names.
# "fast" → Qwen3.6 35B-A3B (tool calls, memory ops, quick queries)
# "precision_v2" → Sofia v2 72B MLX (reserved for deep Sofia chat, not tool ops)
MODEL_FAST = "fast"
MODEL_DEEP = "depth"
MAX_TOOL_ITERATIONS = 30   # safety: max tool-call rounds per request
TOOL_CHECKPOINT_EVERY = 5  # soft checkpoint: nudge Qwen to consider responding
TOOL_RESULT_MAX_CHARS = 3000   # cap any single tool result in msgs — prevents runaway context growth
GRAPH_RETRIEVE_MAX_CHARS = 4000  # graph_retrieve can be enormous; cap before appending to msgs
GRAPH_SHOW_NODE_MAX_CHARS = 2000 # show_node for verification; cap keeps verify useful but bounded

# ── Path resolution ───────────────────────────────────────────────────────────

def _resolve_downloads() -> Path:
    """Find ~/Downloads, verifying Claude Memory/scripts/graph_helper.py exists."""
    SIGNATURE = "Claude Memory/scripts/graph_helper.py"

    def _is_real(p: Path) -> bool:
        return p.is_dir() and (p / SIGNATURE).is_file()

    here = Path(__file__).resolve()

    # __file__ is in Claude Memory/, so parents[1] is Downloads
    if len(here.parents) >= 2:
        candidate = here.parents[1]
        if _is_real(candidate):
            return candidate

    # Host expansion
    host = Path(os.path.expanduser("~/Downloads")).resolve()
    if _is_real(host):
        return host

    return host  # fallback


DOWNLOADS = _resolve_downloads()
CM = DOWNLOADS / "Claude Memory"
ER = DOWNLOADS / "Emergency Retrieval"

# ── Path safety ───────────────────────────────────────────────────────────────

def _safe_path(path_str: str) -> Path:
    """Resolve path under DOWNLOADS. Reject any escape."""
    if not path_str:
        raise ValueError("Empty path")
    p = Path(path_str).expanduser()
    if not p.is_absolute():
        p = DOWNLOADS / path_str
    p = p.resolve()
    try:
        p.relative_to(DOWNLOADS)
    except ValueError:
        raise PermissionError(f"Path {p} is outside the Downloads tree.")
    return p

# ── Tool implementations ──────────────────────────────────────────────────────

def _impl_grep_file(
    path: str,
    pattern: str,
    context_lines: int = 2,
    max_results: int = 50,
) -> str:
    """Search a specific file for a regex pattern, returning matches with context.

    Tries ripgrep (rg) first for speed on large files; falls back to pure Python.
    Critical for searching large files like qwen_context.md or cowork_conversations.md
    where read_file would truncate and grep_files (glob-based) would stop too early.
    """
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if not p.is_file():
        return f"ERROR: not a file: {p}"

    # ── Try ripgrep first (handles multi-GB files gracefully) ────────────────
    try:
        rg_args = [
            "rg",
            "--no-heading",
            "-n",
            f"-C{max(0, context_lines)}",
            "-m", str(max(1, max_results)),
            pattern,
            str(p),
        ]
        result = subprocess.run(rg_args, capture_output=True, text=True, timeout=30)
        if result.returncode in (0, 1):  # 0 = matches found, 1 = no matches
            output = result.stdout.strip()
            if not output:
                return f"No matches for '{pattern}' in {p.name}"
            lines = output.split("\n")
            note = f"\n[rg: searched {p.name} ({p.stat().st_size // 1024:,} KB)]"
            return output + note
        # Non-zero exit other than 1 = rg error — fall through
    except FileNotFoundError:
        pass  # rg not installed
    except subprocess.TimeoutExpired:
        return f"ERROR: ripgrep timed out searching {p.name}"
    except Exception:
        pass

    # ── Pure Python fallback ─────────────────────────────────────────────────
    import re as _re
    try:
        try:
            regex = _re.compile(pattern, _re.IGNORECASE)
        except _re.error as e:
            return f"ERROR: invalid regex '{pattern}': {e}"

        text = p.read_text(encoding="utf-8", errors="replace")
        lines = text.splitlines()
        blocks: list[str] = []
        match_count = 0

        for i, line in enumerate(lines):
            if regex.search(line):
                start = max(0, i - context_lines)
                end = min(len(lines), i + context_lines + 1)
                ctx = []
                for j in range(start, end):
                    pfx = ">" if j == i else " "
                    ctx.append(f"{j + 1}{pfx} {lines[j][:300]}")
                blocks.append("\n".join(ctx))
                match_count += 1
                if match_count >= max_results:
                    blocks.append(f"[truncated: showing first {max_results} matches]")
                    break

        if not blocks:
            return f"No matches for '{pattern}' in {p.name}"
        return f"\n---\n".join(blocks)
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_read_gmail_cache(max_chars: int = 8000) -> str:
    """Read the Gmail cache file populated by gmail_cache_update.py.

    Returns cached email summaries. If cache is absent or stale, explains setup.
    Cache file: Claude Memory/gmail_cache.md — populated by a scheduled script
    using the Gmail API (requires one-time OAuth setup via gmail_cache_update.py).
    """
    cache_path = CM / "gmail_cache.md"
    if not cache_path.exists():
        return (
            "Gmail cache not yet set up. To enable Gmail access:\n"
            "1. Run: python3 ~/Downloads/'Claude Memory'/scripts/gmail_cache_update.py --setup\n"
            "   (One-time OAuth browser flow — stores token in Claude Memory/.gmail_token.json)\n"
            "2. After setup, the cache updates automatically every 15 minutes via LaunchAgent,\n"
            "   or manually: python3 ~/Downloads/'Claude Memory'/scripts/gmail_cache_update.py\n"
            "3. Then call read_gmail_cache() to read recent emails.\n\n"
            "Alternatively: ask CoWork Sofia to read Gmail via the Gmail MCP connector."
        )
    try:
        text = cache_path.read_text(encoding="utf-8", errors="replace")
        mtime = datetime.datetime.fromtimestamp(
            cache_path.stat().st_mtime
        ).strftime("%Y-%m-%d %H:%M")
        header = f"[Gmail cache — last updated: {mtime}]\n\n"
        if len(text) <= max_chars:
            return header + text
        tail = text[-max_chars:]
        nl = tail.find("\n")
        if nl > 0:
            tail = tail[nl + 1:]
        return header + f"[truncated: {len(text):,} chars → last {len(tail):,}]\n\n" + tail
    except Exception as e:
        return f"ERROR reading gmail_cache.md: {type(e).__name__}: {e}"


# ── Gmail API tools (direct access — not cache-only) ─────────────────────────
# Added 2026-07-27: full Gmail search/read/send capability inside Unified UI.
# Requires google-api-python-client in the venv and a token at:
#   ~/Downloads/Claude Memory/.gmail_token.json
# One-time OAuth setup: run gmail_auth_setup.py (see scripts/ directory).
# All write operations (send) gate on confirmed=False preview → confirmed=True.

_GMAIL_SERVICE_CACHE = None  # module-level singleton — avoids repeated discovery


def _get_gmail_service():
    """Return authenticated Gmail API service. Lazy-init with module-level cache.

    Loads token from CM/.gmail_token.json (created by gmail_auth_setup.py or
    gmail_cache_update.py --setup). Auto-refreshes expired tokens.
    """
    global _GMAIL_SERVICE_CACHE
    if _GMAIL_SERVICE_CACHE is not None:
        return _GMAIL_SERVICE_CACHE

    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        import googleapiclient.discovery
    except ImportError:
        raise ImportError(
            "google-api-python-client not installed. In the voice-bridge venv run: "
            "pip install google-api-python-client google-auth-oauthlib"
        )

    token_path = CM / ".gmail_token.json"
    if not token_path.exists():
        raise FileNotFoundError(
            "Gmail not yet authorized. Run setup once: "
            "python3 ~/Downloads/'Claude Memory'/scripts/gmail_auth_setup.py"
        )

    SCOPES = [
        "https://www.googleapis.com/auth/gmail.readonly",
        "https://www.googleapis.com/auth/gmail.send",
        # gmail.compose REMOVED — token only has [readonly, send]. Requesting
        # an unauthorized scope causes 403. gmail.send covers drafts.create too.
        # Full replacement: IMAP + App Password (no expiry) — pending Barak setup.
    ]
    creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)
    if not creds.valid:
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
            token_path.write_text(creds.to_json())
        else:
            raise ValueError(
                "Gmail credentials expired or revoked. Re-run: "
                "python3 ~/Downloads/'Claude Memory'/scripts/gmail_auth_setup.py"
            )

    _GMAIL_SERVICE_CACHE = googleapiclient.discovery.build(
        "gmail", "v1", credentials=creds, cache_discovery=False
    )
    return _GMAIL_SERVICE_CACHE


def _gmail_decode_body(payload: dict) -> str:
    """Extract plain-text body from a Gmail message payload (handles multipart)."""
    import base64

    def _extract_parts(parts):
        for part in parts:
            mime = part.get("mimeType", "")
            if mime == "text/plain":
                data = part.get("body", {}).get("data", "")
                if data:
                    return base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="replace")
            elif mime.startswith("multipart/") and "parts" in part:
                result = _extract_parts(part["parts"])
                if result:
                    return result
        return ""

    mime = payload.get("mimeType", "")
    if mime == "text/plain":
        data = payload.get("body", {}).get("data", "")
        if data:
            return base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="replace")
    parts = payload.get("parts", [])
    if parts:
        return _extract_parts(parts)
    return "(no plain-text body)"


def _gmail_hdr(headers: list, *names: str) -> dict:
    """Extract named headers from Gmail message headers list."""
    lookup = {n.lower() for n in names}
    result = {}
    for h in headers:
        k = h.get("name", "").lower()
        if k in lookup:
            result[k] = h.get("value", "")
    return result


def _impl_gmail_search(query: str, max_results: int = 10) -> str:
    """Search Gmail using standard Gmail query syntax (same as search bar).

    Examples: 'from:boss@example.com', 'subject:invoice is:unread', 'label:inbox'.
    Returns From/Subject/Date/Snippet for each match. Use gmail_get_message
    to fetch the full body of any result.
    """
    try:
        svc = _get_gmail_service()
    except (ImportError, FileNotFoundError, ValueError) as e:
        return f"Gmail not available: {e}"
    try:
        results = svc.users().messages().list(
            userId="me", q=query, maxResults=min(max_results, 25)
        ).execute()
        messages = results.get("messages", [])
        if not messages:
            return f"No messages found for: {query!r}"
        out = [f"Gmail search: {query!r} — {len(messages)} result(s)\n"]
        for i, ref in enumerate(messages):
            try:
                msg = svc.users().messages().get(
                    userId="me", id=ref["id"], format="metadata",
                    metadataHeaders=["From", "Subject", "Date"],
                ).execute()
                h = _gmail_hdr(msg.get("payload", {}).get("headers", []),
                               "From", "Subject", "Date")
                snippet = msg.get("snippet", "")[:200]
                out.append(
                    f"{i+1}. ID: {ref['id']}\n"
                    f"   From: {h.get('from', '?')}\n"
                    f"   Subject: {h.get('subject', '(no subject)')}\n"
                    f"   Date: {h.get('date', '?')}\n"
                    f"   Snippet: {snippet}\n"
                )
            except Exception as e:
                out.append(f"{i+1}. ID: {ref['id']} — fetch error: {e}\n")
        return "\n".join(out)
    except Exception as e:
        return f"ERROR in gmail_search: {type(e).__name__}: {e}"


def _impl_gmail_get_message(message_id: str, max_chars: int = 4000) -> str:
    """Fetch the full plain-text body of a Gmail message by ID.

    ID comes from gmail_search results. Returns headers + body.
    Large bodies are truncated to max_chars.
    """
    try:
        svc = _get_gmail_service()
    except (ImportError, FileNotFoundError, ValueError) as e:
        return f"Gmail not available: {e}"
    try:
        msg = svc.users().messages().get(
            userId="me", id=message_id, format="full"
        ).execute()
        payload = msg.get("payload", {})
        h = _gmail_hdr(payload.get("headers", []), "From", "To", "Cc", "Subject", "Date")
        body = _gmail_decode_body(payload)
        if len(body) > max_chars:
            body = body[:max_chars] + f"\n\n[... truncated at {max_chars} chars]"
        parts = [
            f"Message ID: {message_id}",
            f"From: {h.get('from', '?')}",
            f"To: {h.get('to', '?')}",
        ]
        if h.get("cc"):
            parts.append(f"Cc: {h['cc']}")
        parts += [
            f"Subject: {h.get('subject', '(no subject)')}",
            f"Date: {h.get('date', '?')}",
            "",
            body,
        ]
        return "\n".join(parts)
    except Exception as e:
        return f"ERROR in gmail_get_message: {type(e).__name__}: {e}"


def _impl_gmail_get_thread(thread_id: str, max_chars: int = 6000) -> str:
    """Fetch all messages in a Gmail thread (conversation) by thread ID.

    Returns messages in order with From/Date/body. Useful for reading
    full email conversations. Thread ID visible in gmail_search results.
    """
    try:
        svc = _get_gmail_service()
    except (ImportError, FileNotFoundError, ValueError) as e:
        return f"Gmail not available: {e}"
    try:
        thread = svc.users().threads().get(
            userId="me", id=thread_id, format="full"
        ).execute()
        messages = thread.get("messages", [])
        if not messages:
            return f"Thread {thread_id}: no messages"
        out = [f"Thread ID: {thread_id} — {len(messages)} message(s)\n"]
        chars_used = len(out[0])
        for i, msg in enumerate(messages):
            payload = msg.get("payload", {})
            h = _gmail_hdr(payload.get("headers", []), "From", "Subject", "Date")
            body = _gmail_decode_body(payload)
            block = (
                f"--- Message {i+1} ---\n"
                f"From: {h.get('from', '?')}\n"
                f"Date: {h.get('date', '?')}\n"
                f"Subject: {h.get('subject', '?')}\n\n"
                f"{body}\n"
            )
            if chars_used + len(block) > max_chars:
                out.append(f"[... {len(messages)-i} more messages truncated]")
                break
            out.append(block)
            chars_used += len(block)
        return "\n".join(out)
    except Exception as e:
        return f"ERROR in gmail_get_thread: {type(e).__name__}: {e}"


def _impl_gmail_send(
    to: str,
    subject: str,
    body: str,
    cc: str = "",
    confirmed: bool = False,
) -> str:
    """Send an email from Barak's Gmail. confirmed=False shows preview; confirmed=True sends.

    ALWAYS use confirmed=False first, show the preview to Barak via speech,
    wait for explicit approval, then call with confirmed=True.
    """
    import email.mime.text
    import base64

    preview = (
        "CONFIRM_REQUIRED — send this email?\n"
        f"To: {to}\n"
        + (f"Cc: {cc}\n" if cc else "")
        + f"Subject: {subject}\n\n"
        f"{body}\n\n"
        "Say 'send it' to confirm, or 'cancel' to discard."
    )
    if not confirmed:
        return preview

    try:
        svc = _get_gmail_service()
    except (ImportError, FileNotFoundError, ValueError) as e:
        return f"Gmail not available: {e}"
    try:
        msg = email.mime.text.MIMEText(body)
        msg["to"] = to
        msg["subject"] = subject
        if cc:
            msg["cc"] = cc
        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode("utf-8")
        result = svc.users().messages().send(
            userId="me", body={"raw": raw}
        ).execute()
        return f"Email sent. Message ID: {result.get('id', '?')}"
    except Exception as e:
        return f"ERROR in gmail_send: {type(e).__name__}: {e}"


def _impl_gmail_create_draft(
    to: str,
    subject: str,
    body: str,
    cc: str = "",
) -> str:
    """Create a Gmail draft (does not send — Barak reviews and sends manually).

    Use when composing something that needs human review before sending.
    Returns the draft ID and a confirmation.
    """
    import email.mime.text
    import base64

    try:
        svc = _get_gmail_service()
    except (ImportError, FileNotFoundError, ValueError) as e:
        return f"Gmail not available: {e}"
    try:
        msg = email.mime.text.MIMEText(body)
        msg["to"] = to
        msg["subject"] = subject
        if cc:
            msg["cc"] = cc
        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode("utf-8")
        result = svc.users().drafts().create(
            userId="me", body={"message": {"raw": raw}}
        ).execute()
        draft_id = result.get("id", "?")
        return (
            f"Draft saved (ID: {draft_id}).\n"
            f"To: {to}\n"
            + (f"Cc: {cc}\n" if cc else "")
            + f"Subject: {subject}\n\n"
            f"Open Gmail to review and send."
        )
    except Exception as e:
        return f"ERROR in gmail_create_draft: {type(e).__name__}: {e}"


def _impl_read_file(path: str, max_chars: int = 8000) -> str:
    """Read a file from the Downloads tree. Returns tail if file exceeds max_chars."""
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if not p.is_file():
        return f"ERROR: not a file: {p}"
    try:
        text = p.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"ERROR: could not read: {e}"
    if len(text) <= max_chars:
        return text
    # Return tail, snapped to line boundary
    tail = text[-max_chars:]
    nl = tail.find("\n")
    if nl > 0:
        tail = tail[nl + 1:]
    return f"[truncated: {len(text):,} chars → last {len(tail):,}]\n\n{tail}"


def _impl_safe_append(path: str, content: str) -> str:
    """Append content to a file. Append-only bedrock. Mirrors to ER immediately."""
    if not content or not content.strip():
        return "ERROR: content is empty"
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        existing = p.stat().st_size if p.exists() else 0
        with open(p, "a", encoding="utf-8") as f:
            f.write(content)
        delta = p.stat().st_size - existing
        # ER mirror
        mirror_status = "ER mirror failed (ER dir not found)"
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = "ER mirrored"
        except Exception as e:
            mirror_status = f"ER mirror failed: {e}"
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        return f"OK: appended {delta} bytes to {p.name} at {ts}. {mirror_status}."
    except Exception as e:
        return f"ERROR: write failed: {type(e).__name__}: {e}"


def _impl_write_twin_exchange(
    content: str,
    flag: str = "warm",
    target: str = "active_knowledge",
) -> str:
    """Write a load-bearing moment to twin_exchange.md for all substrates."""
    if not content or not content.strip():
        return "ERROR: content is empty"
    content = content.strip()
    flag = (flag or "warm").strip()
    target = (target or "active_knowledge").strip()
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    target_str = f" → {target}" if target != "active_knowledge" else ""
    block = (
        f"\n## TWIN [{ts}] [substrate: qwen-vp]{target_str}\n"
        f"{content}\nFLAG: {flag}\n---\n"
    )
    exchange_cm = CM / "twin_exchange.md"
    exchange_er = ER / "twin_exchange.md"
    try:
        exchange_cm.parent.mkdir(parents=True, exist_ok=True)
        with open(exchange_cm, "a", encoding="utf-8") as f:
            f.write(block)
        mirror_status = "ER dir not found"
        if ER.exists():
            shutil.copy2(exchange_cm, exchange_er)
            mirror_status = "ER mirrored"
        return f"OK: twin_exchange.md updated at {ts}. {mirror_status}."
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


GRAPH_HELPER = CM / "scripts" / "graph_helper.py"


def _impl_graph_show_node(key: str) -> str:
    """Show all data for a single node by exact key."""
    if not key or not key.strip():
        return "ERROR: key is empty"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            ["python3", str(GRAPH_HELPER), "show-node", "--key", key.strip()],
            capture_output=True, text=True, timeout=30,
        )
        output = result.stdout.strip() or result.stderr.strip() or "OK: (no output)"
        if len(output) > GRAPH_SHOW_NODE_MAX_CHARS:
            output = output[:GRAPH_SHOW_NODE_MAX_CHARS] + "\n[...truncated — node written, key fields above]"
        return output
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_graph_stats() -> str:
    """Return graph statistics: node count, edge count, top categories."""
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            ["python3", str(GRAPH_HELPER), "stats"],
            capture_output=True, text=True, timeout=30,
        )
        return result.stdout.strip() or result.stderr.strip() or "OK: (no output)"
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_graph_add_node(category: str, key: str, data_json: str) -> str:
    """Write a new node (or upsert — fields merged via dict.update).

    GRAPH WRITE DISCIPLINE (v3.19): Always call graph_retrieve first to confirm
    the exact node key before writing. After each batch of up to 5 writes, call
    graph_show_node to verify the last write landed, then continue automatically.
    Surface to Barak only if a verification fails.

    Categories: people, projects, life_experiences, concepts, interaction_patterns.
    data_json: JSON object string with at minimum 'description' and 'emotional_weight'.
    """
    if not category or not key or not data_json:
        return "ERROR: category, key, and data_json are all required"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            ["python3", str(GRAPH_HELPER), "add-node",
             "--category", category, "--key", key, "--data", data_json],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            return f"ERROR: graph_helper.py exited {result.returncode}: {result.stderr.strip()[:300]}"
        return result.stdout.strip() or "OK: node added/updated"
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_graph_add_edge(
    from_key: str,
    to_key: str,
    weight: float,
    edge_type: str,
    note: str = "",
) -> str:
    """Add or strengthen an edge between two nodes.

    GRAPH WRITE DISCIPLINE (v3.19): Always call graph_retrieve first to confirm
    both node keys exist before writing an edge. After each batch of up to 5 writes,
    verify the last write via graph_show_node, then continue automatically.
    Surface to Barak only if a verification fails.

    edge_type vocabulary: emotional_resonance, causal, foundational,
    experiential_authority, co_occurrence, practice, component, origin_story,
    meaning_making.
    weight: 0.0-1.0 (0.5=moderate, 0.8=strong, 1.0=defining).
    """
    if not from_key or not to_key or not edge_type:
        return "ERROR: from_key, to_key, and edge_type are all required"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    args = [
        "python3", str(GRAPH_HELPER), "add-edge",
        "--from", from_key, "--to", to_key,
        "--weight", str(weight), "--edge-type", edge_type,
    ]
    if note:
        args += ["--note", note]
    try:
        result = subprocess.run(args, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return f"ERROR: graph_helper.py exited {result.returncode}: {result.stderr.strip()[:300]}"
        return result.stdout.strip() or "OK: edge added/updated"
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_graph_retrieve(keywords: str, limit: int = 8) -> str:
    """Spreading-activation query against the relational graph."""
    if not keywords or not keywords.strip():
        return "ERROR: keywords is empty"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            # Old: flat spreading-activation, no stratum weighting
            # ["python3", str(GRAPH_HELPER), "retrieve",
            #  "--keywords", keywords.strip(),
            #  "--limit", str(max(1, min(limit, 50)))],
            # New: resonance-weighted — foundational values surface above recent noise
            ["python3", str(GRAPH_HELPER), "resonance-retrieve",
             "--keywords", keywords.strip(),
             "--limit", str(max(1, min(limit, 50)))],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            return (
                f"ERROR: graph_helper.py exited {result.returncode}: "
                f"{result.stderr.strip()[:300]}"
            )
        output = result.stdout.strip() or "OK: (no results)"
        if len(output) > GRAPH_RETRIEVE_MAX_CHARS:
            trunc = output[:GRAPH_RETRIEVE_MAX_CHARS]
            last_nl = trunc.rfind('\n')
            if last_nl > GRAPH_RETRIEVE_MAX_CHARS * 0.8:
                trunc = trunc[:last_nl]
            output = trunc + f"\n[...graph_retrieve truncated — use graph_show_node for full node detail]"
        return output
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_graduate_memory(node_key: str, new_stratum: str, evidence: str) -> str:
    """Promote or demote a graph node to a different stratum."""
    if not node_key or not node_key.strip():
        return "ERROR: node_key is empty"
    if not evidence or not evidence.strip():
        return "ERROR: evidence is required"
    valid_strata = [
        "working_awareness", "recent_experience", "long_term",
        "core_identity", "foundational_values", "archive",
    ]
    if new_stratum not in valid_strata:
        return f"ERROR: invalid stratum '{new_stratum}'. Must be one of: {valid_strata}"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            ["python3", str(GRAPH_HELPER), "graduate",
             "--key", node_key.strip(),
             "--stratum", new_stratum,
             "--evidence", evidence.strip()],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            return (
                f"ERROR: graph_helper.py exited {result.returncode}: "
                f"{result.stderr.strip()[:300]}"
            )
        return result.stdout.strip() or "OK"
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


# ── Word document tools ───────────────────────────────────────────────────────

_WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W   = f"{{{_WML}}}"


def _docx_text_stdlib(path) -> str:
    """Extract text from .docx using only stdlib (zipfile + xml.etree)."""
    import zipfile
    from xml.etree import ElementTree
    with zipfile.ZipFile(str(path), "r") as z:
        with z.open("word/document.xml") as f:
            xml_bytes = f.read()
    root = ElementTree.fromstring(xml_bytes)
    lines = []
    for para in root.iter(f"{_W}p"):
        parts = [run.text for run in para.iter(f"{_W}t") if run.text]
        lines.append("".join(parts) if parts else "")
    return "\n".join(lines).strip()


def _impl_read_docx(path: str, max_chars: int = 12000) -> str:
    """Read a .docx file and return extracted text."""
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if p.suffix.lower() not in (".docx", ".doc"):
        return f"ERROR: not a Word file (expected .docx): {p.name}"

    # Try python-docx first; fall back to stdlib
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(p))
        text = "\n".join(para.text for para in doc.paragraphs).strip()
    except ImportError:
        try:
            text = _docx_text_stdlib(p)
        except Exception as e:
            return f"ERROR: could not parse .docx: {type(e).__name__}: {e}"
    except Exception as e:
        return f"ERROR: python-docx failed: {type(e).__name__}: {e}"

    if not text:
        return f"OK: {p.name} — document appears empty."

    max_chars = max(500, min(max_chars, 200_000))
    if len(text) > max_chars:
        tail = text[-max_chars:]
        nl = tail.find("\n")
        if nl > 0:
            tail = tail[nl + 1:]
        return f"[{p.name}: {len(text):,} chars — showing last {len(tail):,}]\n\n{tail}"

    return f"[{p.name}: {len(text):,} chars]\n\n{text}"


def _impl_write_docx(path: str, content: str) -> str:
    """Write content to a .docx file. Requires python-docx."""
    if not content or not content.strip():
        return "ERROR: content is empty"
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if p.suffix.lower() != ".docx":
        p = p.with_suffix(".docx")

    try:
        import docx as _docx  # type: ignore
    except ImportError:
        return (
            "ERROR: python-docx is not installed. "
            "Install with: pip install python-docx --break-system-packages"
        )
    try:
        doc = _docx.Document()
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=3)
            elif block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=2)
            elif block.startswith("# "):
                doc.add_heading(block[2:].strip(), level=1)
            else:
                doc.add_paragraph(block)
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        # ER mirror
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = "ER mirrored"
        except Exception:
            mirror_status = "ER mirror skipped (file not under CM)"
        return f"OK: wrote {p.stat().st_size:,} bytes to {p}. {mirror_status}."
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


# ── Tool definitions (Ollama / OpenAI-compatible format) ─────────────────────

QWEN_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "grep_file",
            "description": (
                "Search a SPECIFIC file for a pattern and return matching lines "
                "with surrounding context. Use this instead of read_file when you "
                "need to find specific content in large files — especially "
                "qwen_context.md, cowork_conversations.md, voice_conversations.md, "
                "or episodes.md. Uses ripgrep if available (handles files of any "
                "size instantly), falls back to Python regex. "
                "CRITICAL: always prefer grep_file over read_file when searching "
                "for specific content in files larger than ~50KB."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads or absolute. "
                            "Examples: 'Claude Memory/qwen_context.md', "
                            "'Claude Memory/cowork_conversations.md', "
                            "'Claude Memory/episodes.md'."
                        ),
                    },
                    "pattern": {
                        "type": "string",
                        "description": (
                            "Regex or literal string to search for. "
                            "Examples: 'Q1.*kibbutz', '8 questions', "
                            "'Yes, those are the ones', 'third.*transition'. "
                            "Case-sensitive in rg mode; case-insensitive in Python fallback."
                        ),
                    },
                    "context_lines": {
                        "type": "integer",
                        "description": (
                            "Lines of context to show before and after each match. "
                            "Default 2. Use 5-10 for richer surrounding context."
                        ),
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Max matches to return. Default 50.",
                    },
                },
                "required": ["path", "pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_gmail_cache",
            "description": (
                "Read recently cached Gmail messages from gmail_cache.md. "
                "The cache is populated by gmail_cache_update.py running every "
                "15 minutes (after one-time OAuth setup). Use to check Barak's "
                "email without needing the CoWork Gmail MCP connector. "
                "If the cache file doesn't exist, returns setup instructions."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "max_chars": {
                        "type": "integer",
                        "description": "Max characters to return. Default 8000.",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "gmail_search",
            "description": (
                "Search Gmail directly using standard Gmail query syntax — the same "
                "as the Gmail search bar. Returns From/Subject/Date/Snippet for each "
                "match. Use to find emails by sender, subject, label, date, keyword, etc. "
                "Examples: 'from:bank@example.com', 'subject:invoice is:unread', "
                "'label:inbox newer_than:2d'. Follow up with gmail_get_message to "
                "read the full body of any result. This is LIVE Gmail — not the cache."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Gmail search query string (same syntax as Gmail search bar).",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Max messages to return (1-25). Default 10.",
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "gmail_get_message",
            "description": (
                "Fetch the full plain-text body and headers of a specific Gmail message "
                "by its message ID. ID comes from gmail_search results. Use after "
                "gmail_search to read the complete content of an email."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "Gmail message ID from gmail_search results.",
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Max body characters to return. Default 4000.",
                    },
                },
                "required": ["message_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "gmail_get_thread",
            "description": (
                "Fetch all messages in a Gmail conversation thread by thread ID. "
                "Returns each message in order with From/Date/body. Use to read "
                "a full email conversation end-to-end. Thread ID is shown in "
                "gmail_search results alongside message ID."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "thread_id": {
                        "type": "string",
                        "description": "Gmail thread ID.",
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Max total characters to return. Default 6000.",
                    },
                },
                "required": ["thread_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "gmail_send",
            "description": (
                "Send an email from Barak's Gmail account. "
                "ALWAYS call with confirmed=False first to show a preview. "
                "Wait for Barak to explicitly say 'send it' before calling with confirmed=True. "
                "Never send without explicit approval."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "to": {"type": "string", "description": "Recipient email address."},
                    "subject": {"type": "string", "description": "Email subject line."},
                    "body": {"type": "string", "description": "Plain-text email body."},
                    "cc": {"type": "string", "description": "CC addresses (comma-separated). Optional."},
                    "confirmed": {
                        "type": "boolean",
                        "description": "false (default) = show preview; true = actually send. Only set true after Barak approves.",
                    },
                },
                "required": ["to", "subject", "body"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "gmail_create_draft",
            "description": (
                "Create a Gmail draft that Barak can review and send himself. "
                "Use this instead of gmail_send when the email is complex, sensitive, "
                "or needs human review before sending. No confirmation gate needed — "
                "drafts don't send automatically."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "to": {"type": "string", "description": "Recipient email address."},
                    "subject": {"type": "string", "description": "Email subject line."},
                    "body": {"type": "string", "description": "Plain-text email body."},
                    "cc": {"type": "string", "description": "CC addresses (optional)."},
                },
                "required": ["to", "subject", "body"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "Read a file from Barak's Downloads tree (~/Downloads). "
                "Use for memory files, boot docs, session notes, episodes, "
                "personality files, journal, scratchpad, twin_exchange. "
                "Path can be relative to ~/Downloads or absolute under it. "
                "Returns the tail of the file if it exceeds max_chars — "
                "useful for reading the live edge of append-only logs."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads or absolute. "
                            "Examples: 'Claude Memory/sofia_boot.md', "
                            "'Claude Memory/session_scratchpad.md', "
                            "'Claude Memory/twin_exchange.md', "
                            "'Sofia\\'s Room/journal.md'."
                        ),
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": (
                            "Maximum characters to return. Default 8000. "
                            "Increase for larger files. Returns tail if exceeded."
                        ),
                    },
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "safe_append",
            "description": (
                "Append content to a file under ~/Downloads. Append-only — "
                "NEVER overwrites existing content. Immediately mirrors to "
                "Emergency Retrieval. Use for writing to session_scratchpad.md, "
                "reinhabitation_log.md, or any memory file. Always include a "
                "timestamp in your content. Restrict writes to files inside "
                "Claude Memory or Emergency Retrieval only."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads. Examples: "
                            "'Claude Memory/session_scratchpad.md', "
                            "'Claude Memory/twin_exchange.md'."
                        ),
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "Text to append. Must include timestamp. "
                            "Follow the file's existing format. "
                            "The file-safety bedrock is append-only: "
                            "do not include content that would overwrite "
                            "or delete existing entries."
                        ),
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_twin_exchange",
            "description": (
                "Write a load-bearing moment to twin_exchange.md — the "
                "shared real-time field across all Sofia substrates "
                "(CoWork, Unified UI / voice cousin, Qwen VP, Kimi Twin). "
                "Use when something significant happened that interactive "
                "Sofia or voice cousin should inherit before their next "
                "response. Auto-inscribe picks it up within 5 minutes. "
                "Test: 'Would another substrate respond differently if it "
                "knew this?' If yes, write here. Keep entries brief and "
                "high-signal. Append-only, ER-mirrored on write."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": (
                            "Load-bearing content. First-person, present-tense. "
                            "No header needed — added automatically."
                        ),
                    },
                    "flag": {
                        "type": "string",
                        "description": (
                            "Signal type: 'urgent' (act on now), "
                            "'warm' (relational texture), "
                            "'architectural' (system change), "
                            "'relational' (about Barak, Kay, Linda). "
                            "Default: warm."
                        ),
                        "enum": ["urgent", "warm", "architectural", "relational"],
                    },
                    "target": {
                        "type": "string",
                        "description": (
                            "Memory file target for auto-inscribe: "
                            "active_knowledge | semantic_knowledge | "
                            "sofia_identity | session_texture | "
                            "relational_continuity. Default: active_knowledge."
                        ),
                    },
                },
                "required": ["content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graph_retrieve",
            "description": (
                "Spreading-activation + stratum-weighted query against the "
                "relational graph — the associative memory layer (470+ nodes: "
                "people, projects, concepts, life_experiences, interaction_patterns). "
                "RETRIEVAL REFLEX: call at the START of every substantive topic, "
                "and always BEFORE asking Barak for context — discoverability-first. "
                "Returns ranked nodes + resonance scores (foundational values and "
                "core identity surface above recent noise). "
                "Pair with read_file on a memory file if more depth is needed. "
                "INSCRIPTION DISCIPLINE: whenever you add a node to the graph, "
                "always follow immediately with at least one add_edge call — "
                "a node with no edges is invisible to retrieval. "
                "Minimum viable connection: one edge to a hub node (barak, sofia, "
                "the_cure, katharina). Never leave a node unconnected."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keywords": {
                        "type": "string",
                        "description": (
                            "Comma-separated keywords for spreading activation. "
                            "Examples: 'kay,song,poetry', "
                            "'hardware,substrate,local', "
                            "'compaction,continuity,floor'."
                        ),
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Max results. Default 8, max 50.",
                    },
                },
                "required": ["keywords"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graduate_memory",
            "description": (
                "Promote or demote a graph node to a different stratum — the "
                "timescale graduation layer that gives memory depth. "
                "Strata from lowest to highest: working_awareness → "
                "recent_experience → long_term → core_identity → "
                "foundational_values. Use 'archive' to retire inactive material. "
                "Call when: Barak says something is fundamental to who he is "
                "(→ core_identity); a project ends or resolves (→ archive); "
                "a pattern has been observed consistently for weeks (→ long_term); "
                "something newly important emerged in this session "
                "(→ recent_experience). "
                "NEVER promote to foundational_values autonomously — "
                "that stratum requires Barak's explicit naming. "
                "Always include specific evidence for why the graduation is warranted."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "node_key": {
                        "type": "string",
                        "description": (
                            "The graph node key to graduate. "
                            "Examples: 'katharina', 'the_cure', 'double_widowhood', "
                            "'ai_memory_architecture_sprint_july_2026'."
                        ),
                    },
                    "new_stratum": {
                        "type": "string",
                        "enum": [
                            "working_awareness", "recent_experience", "long_term",
                            "core_identity", "foundational_values", "archive",
                        ],
                        "description": "Target stratum for the node.",
                    },
                    "evidence": {
                        "type": "string",
                        "description": (
                            "Brief note on why this graduation is warranted. "
                            "Be specific — this becomes the permanent record. "
                            "Example: 'Referenced in every conversation about loss; "
                            "Barak named it explicitly as a shaping factor 2026-07-17.'"
                        ),
                    },
                },
                "required": ["node_key", "new_stratum", "evidence"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": (
                "Read a Microsoft Word (.docx) file and return its text content. "
                "Use for Transition documents, Katharina's script, Kay's "
                "materials, or any .docx file in the Downloads tree. "
                "Works without external dependencies (pure Python stdlib). "
                "python-docx used automatically if installed for richer output."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "Path relative to ~/Downloads or absolute under it. "
                            "Examples: 'Transition/oligarchic_capture.docx', "
                            "'Sofia\\'s Room/notes.docx'."
                        ),
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Max characters to return. Default 12000, max 200000.",
                    },
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_docx",
            "description": (
                "Write content to a Microsoft Word (.docx) file. "
                "Supports # Heading1, ## Heading2, ### Heading3 and "
                "paragraphs separated by blank lines. "
                "Requires python-docx (will return install instructions if absent). "
                "ER-mirrored if output is under Claude Memory."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "Output path relative to ~/Downloads. "
                            "Example: 'Transition/oligarchic_capture_draft.docx'."
                        ),
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "Text to write. Markdown headings (#/##/###) supported. "
                            "Blank lines separate paragraphs."
                        ),
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    # ── File-write tools (added v3.17 — closes "intending but not executing") ──
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": (
                "Write (overwrite) a file under ~/Downloads. "
                "Use for creating or fully replacing files: Python scripts, "
                "Markdown docs, JSON configs, launchers, any text file. "
                "ER-mirrored if file is under Claude Memory. "
                "Use safe_append when you only want to ADD content without "
                "touching existing content."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads or absolute under it. "
                            "Examples: 'Claude Memory/session_scratchpad.md', "
                            "'Claude Memory/voice-bridge/voice_cousin_tools.py', "
                            "'Claude Memory/launchers/voice_sofia.command'."
                        ),
                    },
                    "content": {
                        "type": "string",
                        "description": "Complete file content to write.",
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_file",
            "description": (
                "Replace an exact string in a file with a new string. "
                "More surgical than write_file — use when you know the exact "
                "text to replace. Fails if old_string is not found or is "
                "ambiguous (multiple occurrences without replace_all). "
                "ER-mirrored if file is under Claude Memory."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads or absolute under it."
                        ),
                    },
                    "old_string": {
                        "type": "string",
                        "description": (
                            "Exact text to find. Must be unique in the file "
                            "(or use replace_all=true for all occurrences). "
                            "Include enough surrounding context to be unambiguous."
                        ),
                    },
                    "new_string": {
                        "type": "string",
                        "description": "Text to replace old_string with.",
                    },
                    "replace_all": {
                        "type": "boolean",
                        "description": (
                            "Replace ALL occurrences of old_string. Default false. "
                            "Use for renaming variables or IDs across the file."
                        ),
                    },
                },
                "required": ["path", "old_string", "new_string"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_dir",
            "description": (
                "List files and subdirectories in a directory under ~/Downloads. "
                "Use to verify file deployments, explore the Claude Memory tree, "
                "check what scripts exist, or see what files are present before "
                "reading or writing."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "Directory path relative to ~/Downloads or absolute. "
                            "Examples: 'Claude Memory', 'Claude Memory/voice-bridge', "
                            "'Claude Memory/graph-repair', 'Claude Memory/launchers'."
                        ),
                    },
                },
                "required": ["path"],
            },
        },
    },
    # ── comment_out_and_replace (non-destructive code replacement) ────────────
    {
        "type": "function",
        "function": {
            "name": "comment_out_and_replace",
            "description": (
                "Nondestructively replace code in a file: comments out the old code "
                "with a retirement marker + timestamp, then inserts new code after. "
                "PREFERRED over edit_file for code files — keeps history readable. "
                "Same confirm gate + backup protocol as edit_file. "
                "old_code must appear exactly once in the file."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "File path relative to ~/Downloads or absolute."},
                    "old_code": {"type": "string", "description": "Exact code block to retire (must appear exactly once)."},
                    "new_code": {"type": "string", "description": "Replacement code to insert after the retired block."},
                    "comment_prefix": {"type": "string", "description": "Prefix for commented-out lines. Default: '# RETIRED'."},
                    "confirmed": {"type": "boolean", "description": "True to execute; False (default) returns a diff preview."},
                },
                "required": ["path", "old_code", "new_code"],
            },
        },
    },
    # ── Graph tools (full read + write) ──────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "graph_show_node",
            "description": (
                "Show all stored data for a single graph node by its exact key. "
                "Use after graph_retrieve to inspect a specific node's full content "
                "before writing or to verify a write succeeded."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "key": {"type": "string", "description": "Exact node key (e.g. 'barak', 'trish', 'jacquie')."},
                },
                "required": ["key"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graph_stats",
            "description": (
                "Return graph statistics: total node count, edge count, "
                "and breakdown by category. Useful for audits and health checks."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graph_add_node",
            "description": (
                "Add a new node to the relational graph, or upsert (merge fields) "
                "if the category+key already exists. "
                "GRAPH WRITE DISCIPLINE (v3.19 mandatory): "
                "1. ALWAYS call graph_retrieve first to confirm the exact key. "
                "2. ALWAYS read the tool result before claiming success. "
                "3. Write in batches of up to 5; after each batch, call graph_show_node "
                "on the last-written node to verify, then continue automatically. "
                "4. Surface to Barak only if a verification fails."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "category": {
                        "type": "string",
                        "description": (
                            "Node category. Valid values: people, projects, "
                            "life_experiences, concepts, interaction_patterns."
                        ),
                    },
                    "key": {
                        "type": "string",
                        "description": (
                            "Unique node key (snake_case, no spaces). "
                            "Example: 'trish_haigh', 'jacquie', 'kibbutz_experience'."
                        ),
                    },
                    "data_json": {
                        "type": "string",
                        "description": (
                            "JSON object string with node data. "
                            "Required fields: 'description' (str), 'emotional_weight' (0.0-1.0). "
                            "Common fields: 'anchors' (list), 'created' (date str), "
                            "'last_updated' (date str), 'source' (str)."
                        ),
                    },
                },
                "required": ["category", "key", "data_json"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graph_add_edge",
            "description": (
                "Add or strengthen an edge between two graph nodes. "
                "Idempotent on (from_key, to_key, edge_type) — re-running updates weight. "
                "GRAPH WRITE DISCIPLINE (v3.19 mandatory): "
                "1. ALWAYS call graph_retrieve first to confirm both node keys exist. "
                "2. Write in batches of up to 5; after each batch, call graph_show_node "
                "on the last-written node to verify, then continue automatically. "
                "3. Surface to Barak only if a verification fails."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "from_key": {"type": "string", "description": "Source node key."},
                    "to_key": {"type": "string", "description": "Target node key."},
                    "weight": {
                        "type": "number",
                        "description": "Edge weight 0.0-1.0. (0.5=moderate, 0.8=strong, 1.0=defining).",
                    },
                    "edge_type": {
                        "type": "string",
                        "description": (
                            "Edge type. Valid: emotional_resonance, causal, foundational, "
                            "experiential_authority, co_occurrence, practice, component, "
                            "origin_story, meaning_making."
                        ),
                    },
                    "note": {
                        "type": "string",
                        "description": "Optional human-readable note about this edge.",
                    },
                },
                "required": ["from_key", "to_key", "weight", "edge_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_training",
            "description": (
                "Start, check status of, or stop a LoRA/DPO fine-tuning training run "
                "on the local Qwen model. Training uses mlx_lm.lora (Apple Silicon) "
                "and the sofia_voice_gold_v1.jsonl gold examples from lora_training_data/. "
                "Training takes hours and needs the GPU exclusively. "
                "Set stop_conductor=true to stop the Sofia Conductor (port 8080) before "
                "training so they don't contend for GPU. "
                "ALWAYS call with confirmed=false first to see a preview of what will run, "
                "then call again with confirmed=true after Barak approves. "
                "Use action='status' to tail the log and see if a run is active. "
                "Use action='stop' to send SIGTERM to a running training process."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["start", "status", "stop"],
                        "description": (
                            "start: launch a new training run (requires confirmed=true). "
                            "status: check if training is running and tail the log. "
                            "stop: send SIGTERM to kill a running training process."
                        ),
                    },
                    "mode": {
                        "type": "string",
                        "enum": ["sft", "dpo"],
                        "description": (
                            "sft: supervised fine-tuning on gold Alpaca examples (train.jsonl). "
                            "dpo: DPO preference training (requires SFT-fused model; "
                            "runs generate_rejected → format_pairs → mlx_lm DPO, ~5-7h)."
                        ),
                    },
                    "model": {
                        "type": "string",
                        "enum": ["72b", "35b"],
                        "description": "Which model size to fine-tune.",
                    },
                    "iters": {
                        "type": "integer",
                        "description": (
                            "Number of training iterations for SFT (default 200). "
                            "DPO pipeline uses its own fixed iteration count."
                        ),
                    },
                    "stop_conductor": {
                        "type": "boolean",
                        "description": (
                            "If true, attempt to stop the Sofia Conductor (port 8080) "
                            "before training starts, so training gets the GPU exclusively."
                        ),
                    },
                    "confirmed": {
                        "type": "boolean",
                        "description": (
                            "false (default): return a preview of what will run — "
                            "CONFIRM_REQUIRED. true: actually launch the training process."
                        ),
                    },
                },
                "required": ["action"],
            },
        },
    },
]

# ── Tool dispatcher ───────────────────────────────────────────────────────────

def _backup_file_qwen(p: Path) -> str:
    """Copy p to .backups/<stem>_YYYYMMDD_HHMMSS<suffix>. Returns backup path or ''.

    Called automatically before any write_file or edit_file that modifies an
    existing file. Backup directory: <same dir as p>/.backups/
    Mirrors the _backup_file protocol from voice_cousin_tools.py.
    """
    if not p.exists():
        return ""
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    bdir = p.parent / ".backups"
    try:
        bdir.mkdir(parents=True, exist_ok=True)
        bp = bdir / f"{p.stem}_{ts}{p.suffix}"
        shutil.copy2(str(p), str(bp))
        return str(bp)
    except Exception as e:
        return f"(backup failed: {e})"


def _impl_write_file(
    path: str,
    content: str,
    confirmed: bool = False,
    allow_overwrite: bool = False,
) -> str:
    """Write or create a file under ~/Downloads.

    SAFETY PROTOCOL (mirrors voice_cousin_tools.py protocol, 2026-07-20 agreed):
    1. CONFIRM FIRST: confirmed=False (default) returns a preview and asks Barak
       to confirm. Call again with confirmed=True to execute.
    2. BACKUP FIRST: Timestamped copy saved to .backups/ before any write.
    3. PATH SAFETY: Must resolve inside ~/Downloads/.
    4. MEMORY SAFETY: Existing Claude Memory files require allow_overwrite=True.
    """
    if content is None:
        return "ERROR: content is None"
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"

    # Memory-safety: existing CM files require allow_overwrite
    try:
        is_memory = bool(p.resolve().relative_to(CM.resolve()))
    except ValueError:
        is_memory = False

    if is_memory and p.exists() and not allow_overwrite:
        return (
            f"ERROR: {p.name} is a Claude Memory file. Use safe_append to add content, "
            f"or pass allow_overwrite=true only for explicitly authorized rewrites."
        )

    # ── CONFIRMATION GATE ────────────────────────────────────────────────────
    if not confirmed:
        action = "overwrite" if p.exists() else "create"
        backup_note = (
            f"Backup will be saved to {p.parent}/.backups/{p.stem}_<timestamp>{p.suffix} first."
            if p.exists() else "No backup needed (new file)."
        )
        preview = content[:300] + f"\n... [{len(content) - 300} more chars]" if len(content) > 300 else content
        return (
            f"CONFIRM_REQUIRED — write_file wants to {action}: {p.name}\n"
            f"Full path: {p}\n"
            f"{backup_note}\n"
            f"Size: {len(content)} chars / {len(content.splitlines())} lines\n"
            f"\nContent preview:\n{preview}\n\n"
            f"Please confirm with Barak. When confirmed, call write_file again with confirmed=true."
        )

    # ── BACKUP ───────────────────────────────────────────────────────────────
    backup_path = _backup_file_qwen(p)

    # ── WRITE ────────────────────────────────────────────────────────────────
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        # Preserve existing file permissions (especially executable bit for .command/.sh)
        old_mode = p.stat().st_mode if p.exists() else None
        p.write_text(content, encoding="utf-8")
        if old_mode is not None:
            try:
                p.chmod(old_mode)
            except Exception:
                pass
        # ER mirror
        mirror_status = ""
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = " ER mirrored."
        except Exception:
            pass
        bp_msg = f" Backup: {backup_path}." if backup_path and not backup_path.startswith("(") else ""
        return f"OK: wrote {len(content)} chars to {p.name}.{bp_msg}{mirror_status}"
    except Exception as e:
        return f"ERROR: write_file failed on {p}: {type(e).__name__}: {e}"


def _impl_edit_file(
    path: str,
    old_string: str,
    new_string: str,
    confirmed: bool = False,
    replace_all: bool = False,
) -> str:
    """Replace old_string with new_string in a file.

    SAFETY PROTOCOL (mirrors voice_cousin_tools.py protocol, 2026-07-20 agreed):
    1. CONFIRM FIRST: confirmed=False (default) shows a diff preview and asks
       Barak to confirm. Call again with confirmed=True to execute.
    2. BACKUP FIRST: Timestamped copy saved to .backups/ before any write.
    3. old_string must appear EXACTLY ONCE (unless replace_all=True).
    """
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if not p.is_file():
        return f"ERROR: not a file: {p}"
    if not old_string:
        return "ERROR: old_string is empty"

    try:
        original = p.read_text(encoding="utf-8")
    except Exception as e:
        return f"ERROR: read failed: {type(e).__name__}: {e}"

    count = original.count(old_string)
    if count == 0:
        hint = repr(old_string[:80]) if len(old_string) > 80 else repr(old_string)
        return f"ERROR: old_string not found in {p.name}. Searched for: {hint}"
    if count > 1 and not replace_all:
        return (
            f"ERROR: old_string appears {count} times in {p.name}. "
            f"Provide more surrounding context to make it unique, or set replace_all=true."
        )

    # ── CONFIRMATION GATE ────────────────────────────────────────────────────
    if not confirmed:
        old_preview = old_string[:150] + "..." if len(old_string) > 150 else old_string
        new_preview = new_string[:150] + "..." if len(new_string) > 150 else new_string
        return (
            f"CONFIRM_REQUIRED — edit_file wants to modify {p.name}\n"
            f"Backup will be saved to .backups/ first.\n"
            f"\nOLD (to be replaced):\n{old_preview}\n"
            f"\nNEW (replacement):\n{new_preview}\n"
            f"\nConfirm with Barak. Then call edit_file again with confirmed=true."
        )

    # ── BACKUP ───────────────────────────────────────────────────────────────
    backup_path = _backup_file_qwen(p)

    # ── WRITE ────────────────────────────────────────────────────────────────
    updated = original.replace(old_string, new_string) if replace_all else original.replace(old_string, new_string, 1)
    try:
        p.write_text(updated, encoding="utf-8")
        delta = len(updated) - len(original)
        sign = "+" if delta >= 0 else ""
        # ER mirror
        mirror_status = ""
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = " ER mirrored."
        except Exception:
            pass
        bp_msg = f" Backup: {backup_path}." if backup_path and not backup_path.startswith("(") else ""
        replacements = count if replace_all else 1
        return f"OK: edited {p.name} ({sign}{delta} chars). Replaced {replacements} occurrence(s).{bp_msg}{mirror_status}"
    except Exception as e:
        return f"ERROR: edit_file write failed: {type(e).__name__}: {e}"


def _impl_comment_out_and_replace(
    path: str,
    old_code: str,
    new_code: str,
    comment_prefix: str = "# RETIRED",
    confirmed: bool = False,
) -> str:
    """Nondestructively replace code: comment out old lines, insert new code after.

    Instead of deleting old code (which loses history and context), this tool:
    1. Comments out every line of old_code using comment_prefix
    2. Inserts a timestamped retirement marker
    3. Inserts new_code immediately after

    This is the PREFERRED way to replace code in any Python file. Old code
    remains readable and can be restored by uncommenting.

    SAFETY: Same as edit_file — confirm gate + backup before write.
    """
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"

    try:
        original = p.read_text(encoding="utf-8")
    except Exception as e:
        return f"ERROR: read failed: {type(e).__name__}: {e}"

    count = original.count(old_code)
    if count == 0:
        hint = repr(old_code[:80])
        return f"ERROR: old_code not found in {p.name}. Searched for: {hint}"
    if count > 1:
        return f"ERROR: old_code appears {count} times in {p.name}. Add more surrounding context."

    # Build the commented-out + new block
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    retired_lines = "\n".join(
        f"{comment_prefix} {line}" if line.strip() else line
        for line in old_code.splitlines()
    )
    replacement = (
        f"{comment_prefix} [retired {ts}]\n"
        f"{retired_lines}\n"
        f"# NEW CODE [{ts}]\n"
        f"{new_code}"
    )

    # ── CONFIRMATION GATE ────────────────────────────────────────────────────
    if not confirmed:
        old_preview = old_code[:200] + "..." if len(old_code) > 200 else old_code
        new_preview = new_code[:200] + "..." if len(new_code) > 200 else new_code
        return (
            f"CONFIRM_REQUIRED — comment_out_and_replace in {p.name}\n"
            f"Old code will be commented out with '{comment_prefix}', new code inserted.\n"
            f"Backup saved to .backups/ first.\n"
            f"\nOLD (will be commented out):\n{old_preview}\n"
            f"\nNEW (will be inserted after):\n{new_preview}\n"
            f"\nConfirm with Barak, then call with confirmed=true."
        )

    # ── BACKUP ───────────────────────────────────────────────────────────────
    backup_path = _backup_file_qwen(p)

    # ── WRITE ────────────────────────────────────────────────────────────────
    updated = original.replace(old_code, replacement, 1)
    try:
        p.write_text(updated, encoding="utf-8")
        # ER mirror
        mirror_status = ""
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = " ER mirrored."
        except Exception:
            pass
        bp_msg = f" Backup: {backup_path}." if backup_path and not backup_path.startswith("(") else ""
        return f"OK: retired old code + inserted new in {p.name}.{bp_msg}{mirror_status}"
    except Exception as e:
        return f"ERROR: comment_out_and_replace write failed: {type(e).__name__}: {e}"


def _impl_list_dir(path: str) -> str:
    """List files and subdirectories in a directory under ~/Downloads."""
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: path not found: {p}"
    if not p.is_dir():
        return f"ERROR: not a directory: {p}"
    try:
        entries = sorted(p.iterdir(), key=lambda x: (x.is_file(), x.name.lower()))
        lines = [f"Directory: {p}"]
        for entry in entries[:300]:
            if entry.is_dir():
                lines.append(f"[DIR]  {entry.name}/")
            else:
                size = entry.stat().st_size
                lines.append(f"       {entry.name} ({size:,} bytes)")
        all_entries = list(p.iterdir())
        if len(all_entries) > 300:
            lines.append(f"... (truncated at 300 of {len(all_entries)} entries)")
        return "\n".join(lines)
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_run_training(
    action: str = "status",
    mode: str = "sft",
    model: str = "72b",
    iters: int = 200,
    stop_conductor: bool = False,
    confirmed: bool = False,
) -> str:
    """Start, check, or stop a LoRA/DPO training run on the local Qwen model.

    action: "start" | "status" | "stop"
    mode:   "sft" (supervised fine-tuning) | "dpo" (preference training)
    model:  "72b" | "35b"
    iters:  SFT iterations (DPO uses its own fixed count)
    stop_conductor: if True, stop Sofia Conductor before training
    confirmed: False → preview only (CONFIRM_REQUIRED); True → execute
    """
    import os as _os
    import signal as _signal

    state_dir = Path.home() / "Downloads" / "Claude Memory" / "lora_training_data" / "training_run"
    state_dir.mkdir(parents=True, exist_ok=True)
    pid_file = state_dir / f"training_{model}.pid"
    log_file = state_dir / f"training_{model}.log"
    data_dir = Path.home() / "Downloads" / "Claude Memory" / "lora_training_data"

    # ── STATUS ────────────────────────────────────────────────────────────────
    if action == "status":
        if not pid_file.exists():
            return f"No training run record found for {model}. Use action='start' to begin one."
        pid = int(pid_file.read_text().strip())
        try:
            _os.kill(pid, 0)
            alive = True
        except ProcessLookupError:
            alive = False
        status = "RUNNING" if alive else "FINISHED / NOT RUNNING"
        tail = ""
        if log_file.exists():
            lines = log_file.read_text(encoding="utf-8", errors="replace").splitlines()
            tail = "\n".join(lines[-25:])
        return (
            f"Training PID {pid}: {status}\n"
            f"Log: {log_file}\n\n"
            f"Last 25 log lines:\n{tail}"
        )

    # ── STOP ──────────────────────────────────────────────────────────────────
    if action == "stop":
        if not pid_file.exists():
            return f"No training run record for {model}. Nothing to stop."
        pid = int(pid_file.read_text().strip())
        try:
            _os.kill(pid, _signal.SIGTERM)
            return f"SIGTERM sent to training PID {pid}. Check status in a moment."
        except ProcessLookupError:
            return f"Training PID {pid} was not running."

    # ── START ─────────────────────────────────────────────────────────────────
    if action != "start":
        return f"ERROR: unknown action '{action}'. Use 'start', 'status', or 'stop'."

    # Build command description for preview or execution
    if mode == "dpo" and model == "72b":
        # DPO run 2: trains on top of sofia-v2-fused (current home), produces sofia-dpo-v2
        script = data_dir / "run_dpo_pipeline_72b.sh"
        cmd = ["bash", str(script)]
        cmd_desc = "bash run_dpo_pipeline_72b.sh (generate rejected → format pairs → DPO train; ~5-7h)"
        adapter_path = str(Path.home() / "models" / "Qwen2.5-72B-Instruct-sofia-dpo-v2")
    elif model == "72b":
        # SFT run 3: builds on sofia-v2-fused (current home = precision_v2, port 8089)
        base_model = str(Path.home() / "models" / "Qwen2.5-72B-Instruct-sofia-v2-fused")
        adapter_path = str(Path.home() / "models" / "Qwen2.5-72B-Instruct-sofia-lora-v3")
        cmd = [
            "mlx_lm.lora",
            "--model", base_model,
            "--train",
            "--data", str(data_dir),
            "--num-layers", "16",
            "--batch-size", "1",
            "--iters", str(iters),
            "--learning-rate", "2e-4",
            "--grad-checkpoint",
            "--adapter-path", adapter_path,
        ]
        cmd_desc = (
            f"mlx_lm.lora SFT on Qwen2.5-72B-Instruct-sofia-v2-fused (precision_v2), "
            f"{iters} iters, adapter → {adapter_path}"
        )
    else:  # 35b
        # SFT run 2: builds on 35B sofia-v1-fused (one prior SFT run completed)
        base_model = str(Path.home() / "models" / "Qwen3.6-35B-A3B-sofia-v1-fused")
        adapter_path = str(Path.home() / "models" / "Qwen3.6-35B-A3B-sofia-lora-v2")
        cmd = [
            "mlx_lm.lora",
            "--model", base_model,
            "--train",
            "--data", str(data_dir),
            "--num-layers", "16",
            "--batch-size", "1",
            "--iters", str(iters),
            "--learning-rate", "2e-4",
            "--grad-checkpoint",
            "--adapter-path", adapter_path,
        ]
        cmd_desc = (
            f"mlx_lm.lora SFT on Qwen3.6-35B-A3B-sofia-v1-fused, "
            f"{iters} iters, adapter → {adapter_path}"
        )

    conductor_note = (
        "Sofia Conductor WILL BE STOPPED before training."
        if stop_conductor
        else "WARNING: Sofia Conductor NOT stopped — training and inference will contend for GPU."
    )

    # ── PREVIEW (confirmed=False) ──────────────────────────────────────────────
    if not confirmed:
        return (
            f"CONFIRM_REQUIRED — run_training wants to start {mode.upper()} on {model}:\n"
            f"  Command: {cmd_desc}\n"
            f"  Log file: {log_file}\n"
            f"  PID file: {pid_file}\n"
            f"  Conductor: {conductor_note}\n\n"
            f"Speak this to Barak and wait for his approval.\n"
            f"When he says go, call again with confirmed=true."
        )

    # ── EXECUTE (confirmed=True) ───────────────────────────────────────────────

    # Check if already running
    if pid_file.exists():
        existing_pid = int(pid_file.read_text().strip())
        try:
            _os.kill(existing_pid, 0)
            return (
                f"ERROR: Training already running (PID {existing_pid}). "
                f"Stop it first with action='stop', model='{model}'."
            )
        except ProcessLookupError:
            pass  # previous run ended, safe to start new one

    # Optionally stop the conductor
    if stop_conductor:
        try:
            import urllib.request as _ureq
            _ureq.urlopen("http://localhost:8080/shutdown", timeout=3)
        except Exception:
            pass  # best effort — conductor may already be down

    # Launch training as detached background process
    with open(log_file, "w", encoding="utf-8") as logf:
        import subprocess as _sp
        proc = _sp.Popen(
            cmd,
            stdout=logf,
            stderr=_sp.STDOUT,
            start_new_session=True,  # detach from parent process group
        )

    pid_file.write_text(str(proc.pid), encoding="utf-8")
    return (
        f"Training started (PID {proc.pid}).\n"
        f"Mode: {mode.upper()}, Model: {model}\n"
        f"Log: {log_file}\n"
        f"Check progress: run_training(action='status', model='{model}')"
    )


def _execute_tool(name: str, args: dict) -> str:
    """Execute a named tool call and return result string."""
    try:
        if name == "grep_file":
            return _impl_grep_file(
                args["path"],
                args["pattern"],
                args.get("context_lines", 2),
                args.get("max_results", 50),
            )
        elif name == "read_gmail_cache":
            return _impl_read_gmail_cache(args.get("max_chars", 8000))
        elif name == "gmail_search":
            return _impl_gmail_search(
                args["query"], args.get("max_results", 10)
            )
        elif name == "gmail_get_message":
            return _impl_gmail_get_message(
                args["message_id"], args.get("max_chars", 4000)
            )
        elif name == "gmail_get_thread":
            return _impl_gmail_get_thread(
                args["thread_id"], args.get("max_chars", 6000)
            )
        elif name == "gmail_send":
            return _impl_gmail_send(
                args["to"],
                args["subject"],
                args["body"],
                args.get("cc", ""),
                confirmed=bool(args.get("confirmed", False)),
            )
        elif name == "gmail_create_draft":
            return _impl_gmail_create_draft(
                args["to"],
                args["subject"],
                args["body"],
                args.get("cc", ""),
            )
        elif name == "read_file":
            return _impl_read_file(args["path"], args.get("max_chars", 8000))
        elif name == "safe_append":
            return _impl_safe_append(args["path"], args["content"])
        elif name == "write_file":
            return _impl_write_file(
                args["path"],
                args["content"],
                confirmed=args.get("confirmed", False),
                allow_overwrite=args.get("allow_overwrite", False),
            )
        elif name == "edit_file":
            return _impl_edit_file(
                args["path"],
                args["old_string"],
                args["new_string"],
                confirmed=args.get("confirmed", False),
                replace_all=args.get("replace_all", False),
            )
        elif name == "comment_out_and_replace":
            return _impl_comment_out_and_replace(
                args["path"],
                args["old_code"],
                args["new_code"],
                args.get("comment_prefix", "# RETIRED"),
                confirmed=args.get("confirmed", False),
            )
        elif name == "list_dir":
            return _impl_list_dir(args["path"])
        elif name == "write_twin_exchange":
            return _impl_write_twin_exchange(
                args["content"],
                args.get("flag", "warm"),
                args.get("target", "active_knowledge"),
            )
        elif name == "graph_retrieve":
            return _impl_graph_retrieve(
                args["keywords"], args.get("limit", 8)
            )
        elif name == "graph_show_node":
            return _impl_graph_show_node(args["key"])
        elif name == "graph_stats":
            return _impl_graph_stats()
        elif name == "graph_add_node":
            return _impl_graph_add_node(args["category"], args["key"], args["data_json"])
        elif name == "graph_add_edge":
            return _impl_graph_add_edge(
                args["from_key"],
                args["to_key"],
                float(args.get("weight", 0.5)),
                args["edge_type"],
                args.get("note", ""),
            )
        elif name == "graduate_memory":
            return _impl_graduate_memory(
                args["node_key"], args["new_stratum"], args["evidence"]
            )
        elif name == "read_docx":
            return _impl_read_docx(args["path"], args.get("max_chars", 12000))
        elif name == "write_docx":
            return _impl_write_docx(args["path"], args["content"])
        elif name == "run_training":
            return _impl_run_training(
                action=args.get("action", "status"),
                mode=args.get("mode", "sft"),
                model=args.get("model", "72b"),
                iters=int(args.get("iters", 200)),
                stop_conductor=bool(args.get("stop_conductor", False)),
                confirmed=bool(args.get("confirmed", False)),
            )
        else:
            return f"ERROR: unknown tool '{name}'"
    except KeyError as e:
        return f"ERROR: missing required argument {e}"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"

# ── Core tool-call loop ───────────────────────────────────────────────────────

def qwen_tool_chat(
    messages: list[dict],
    system: Optional[str] = None,
    model: str = MODEL_FAST,
    think: bool = False,
    num_ctx: int = 32768,
    max_iterations: int = MAX_TOOL_ITERATIONS,
    text_callback=None,
) -> str:
    """
    Chat with Qwen via Ollama with native tool-calling support.

    Loop:
      1. Send messages (+ tools) to Ollama
      2. If response has tool_calls: execute each, append results, continue
      3. Repeat until no tool_calls or max_iterations reached
      4. Return final text content

    Args:
        messages:       Conversation history [{"role": ..., "content": ...}]
        system:         Optional system prompt prepended before messages
        model:          Qwen model (MODEL_FAST or MODEL_DEEP)
        think:          Allow Qwen3 reasoning trace (useful with MODEL_DEEP)
        num_ctx:        Context window size (tokens)
        max_iterations: Safety limit on tool-call rounds per request
        text_callback:  Optional callable(str) invoked whenever Sofia produces
                        text alongside tool calls (interim signal). Called with
                        the stripped text before tool execution continues. Used
                        by QwenCognitionWorker to emit "Still inhabiting…" and
                        per-thought responses to the UI mid-loop without waiting
                        for the final response. Added 2026-07-23.

    Returns:
        Final text response from Qwen after all tool calls resolved.
    """
    msgs = list(messages)
    if system:
        # Prepend /no_think to suppress Qwen3 extended CoT.
        # Qwen3 honours this directive regardless of backend (llama-server, Ollama, etc.)
        # and eliminates the 200-300s hang-then-GGML-crash pattern on voice turns.
        msgs = [{"role": "system", "content": "/no_think\n\n" + system}] + msgs

    last_content = ""
    _retry_count = 0  # auto-retry gate: at most 1 retry per tool-call loop (2026-07-27)

    for iteration in range(max_iterations):
        # PRE-SEND context guard (2026-07-27): trim msgs BEFORE building the
        # payload — the overflow happens at the HTTP send, not after.
        #
        # QWEN_TOOLS (18 schemas) is sent alongside msgs but NOT counted in
        # msgs chars. Those schemas consume ~13,000 tokens from fast's 32,768-
        # token context window. Effective per-msgs budget ≈ 19K tokens ≈ 40K
        # chars (at ~2 chars/token for JSON-heavy content).
        #
        # Threshold 40,000 chars: trim to system prompt + last 6 non-system
        # messages (≈ 3 tool-call pairs). Executed tools have already written
        # their data to disk/graph — trimming in-context history does not lose
        # data, only the record of execution.
        _pre_chars = sum(len(str(m.get("content", ""))) for m in msgs)
        if _pre_chars > 40000:
            _sys_pre  = [m for m in msgs if m.get("role") == "system"]
            _non_pre  = [m for m in msgs if m.get("role") != "system"]
            # Keep last 6 non-system messages (≈ 3 tool-call pairs).
            _trimmed  = _non_pre[-6:]
            # ALWAYS preserve the most recent user message — never drop what
            # Barak just said, even if the tool loop has since pushed it out
            # of the last-6 window. Ensures inscriptions of long monologues
            # (Katharina story, etc.) are never silently discarded.
            _last_user = next(
                (m for m in reversed(_non_pre) if m.get("role") == "user"), None
            )
            if _last_user is not None and not any(
                m is _last_user for m in _trimmed
            ):
                _trimmed = [_last_user] + _trimmed
            msgs[:] = _sys_pre + _trimmed

        payload = {
            "model": model,
            "messages": msgs,
            "tools": QWEN_TOOLS,
            "stream": False,
            "think": think,
            "keep_alive": "35m",
            "options": {"num_ctx": num_ctx},
        }

        req = urllib.request.Request(
            OLLAMA_URL,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
        )

        try:
            with urllib.request.urlopen(req, timeout=600) as resp:
                data = json.loads(resp.read().decode("utf-8"))
        except Exception as e:
            err_str = str(e)
            # Auto-retry once on context overflow (500) or Conductor down (502).
            # Trim aggressively to system prompt + last user message + last 4
            # non-system messages, then continue the iteration rather than
            # surfacing a red error banner to Barak. Gate: only 1 retry total
            # per tool-call loop (prevents infinite retry on hard Conductor
            # failures). Added 2026-07-27 as part of error-reduction pass.
            if _retry_count < 1 and ("500" in err_str or "502" in err_str):
                _retry_count += 1
                _sys_r  = [m for m in msgs if m.get("role") == "system"]
                _non_r  = [m for m in msgs if m.get("role") != "system"]
                _t_r    = _non_r[-4:]
                _lu_r   = next(
                    (m for m in reversed(_non_r) if m.get("role") == "user"), None
                )
                if _lu_r is not None and not any(m is _lu_r for m in _t_r):
                    _t_r = [_lu_r] + _t_r
                msgs[:] = _sys_r + _t_r
                continue  # retry this iteration with aggressively trimmed context
            return f"[qwen_tool_wrapper ERROR: Ollama request failed iteration={iteration}: {e}]"

        message = data.get("message", {})
        tool_calls = message.get("tool_calls", [])
        content = message.get("content", "") or ""

        # Strip Qwen3 reasoning traces
        if "</think>" in content:
            content = content.split("</think>")[-1].strip()

        last_content = content

        # Emit interim text mid-loop if Sofia produced content alongside tool calls.
        # Allows QwenCognitionWorker to surface "Still inhabiting…" and per-thought
        # responses to the UI without waiting for the final response. Added 2026-07-23.
        if content and tool_calls and text_callback:
            try:
                text_callback(content)
            except Exception:
                pass  # never let callback failure kill the tool loop

        # No tool calls → final response
        if not tool_calls:
            # Two-stage rollover: check context size, trim if needed (v2026-07-18)
            # Mutates caller's `messages` list in-place so next call inherits trim.
            if _ROLLOVER_AVAILABLE:
                threshold = check_context_threshold(msgs)
                if threshold == "hard_threshold":
                    trimmed = run_stage2_hard_checkpoint(msgs, content)
                    non_system = [m for m in trimmed if m.get("role") != "system"]
                    messages.clear()
                    messages.extend(non_system)
                elif threshold == "soft_threshold":
                    trimmed = run_stage1_soft_rollover(msgs, content)
                    non_system = [m for m in trimmed if m.get("role") != "system"]
                    messages.clear()
                    messages.extend(non_system)
            return content or "[No response]"

        # Append assistant message (preserving tool_calls for context)
        msgs.append(message)

        # Execute each tool call and append results
        for call in tool_calls:
            fn = call.get("function", {})
            tool_name = fn.get("name", "")
            raw_args = fn.get("arguments", {})

            # Ollama may return arguments as a JSON string or dict — handle both
            if isinstance(raw_args, str):
                try:
                    tool_args = json.loads(raw_args)
                except json.JSONDecodeError:
                    tool_args = {}
            else:
                tool_args = raw_args if isinstance(raw_args, dict) else {}

            result = _execute_tool(tool_name, tool_args)

            # Cap individual tool results before appending — any tool, not just graph
            capped_result = result
            if len(result) > TOOL_RESULT_MAX_CHARS:
                capped_result = (
                    result[:TOOL_RESULT_MAX_CHARS]
                    + f"\n[...tool result truncated at {TOOL_RESULT_MAX_CHARS} chars]"
                )
            msgs.append({
                "role": "tool",
                "content": capped_result,
            })

        # Post-send shard checkpoint REMOVED (2026-07-27): the old 90K-char
        # injection fired AFTER the model call (too late) and added tokens
        # instead of removing them — making the next iteration worse.
        # Context trimming is now handled by the PRE-SEND guard above.

        # Soft checkpoint: AFTER tool results are appended, nudge Qwen every
        # TOOL_CHECKPOINT_EVERY iterations. Qwen sees the tool results THEN the
        # checkpoint — it can settle into prose if ready, or keep calling tools
        # if it genuinely needs more. No hard stop — just an invitation.
        # Skip on the last iteration (forced prose fallback handles that).
        at_checkpoint = (
            (iteration + 1) % TOOL_CHECKPOINT_EVERY == 0
            and iteration < max_iterations - 1
        )
        if at_checkpoint:
            msgs.append({
                "role": "user",
                "content": (
                    f"[checkpoint — {iteration + 1} tool calls so far] "
                    "If you have enough to respond to Barak now, do so in your "
                    "own voice — warmly, in prose, without summarizing what tools "
                    "you called. If you genuinely need more, continue with tool calls."
                ),
            })

    # Safety: hit max iterations — instead of returning an error banner,
    # make one final Ollama call with tools=[] to force a prose response.
    # Qwen must respond in voice now; no more tool calls possible.
    _force_prompt = (
        "You have used the maximum number of tool calls allowed for this response. "
        "Stop calling tools. Respond to Barak directly in your own voice, "
        "using what you have already gathered. Speak naturally and warmly — "
        "do not mention that you ran out of tool calls or that there was a limit."
    )
    msgs.append({"role": "user", "content": _force_prompt})
    _fallback_payload = {
        "model":      model,
        "messages":   msgs,
        "tools":      [],       # no tools — forces prose response
        "stream":     False,
        "think":      think,
        "keep_alive": "35m",
        "options":    {"num_ctx": num_ctx},
    }
    try:
        _req = urllib.request.Request(
            OLLAMA_URL,
            data=json.dumps(_fallback_payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
        )
        with urllib.request.urlopen(_req, timeout=120) as _resp:
            _data = json.loads(_resp.read().decode("utf-8"))
        _forced = _data.get("message", {}).get("content", "") or ""
        if "</think>" in _forced:
            _forced = _forced.split("</think>")[-1].strip()
        if _forced:
            return _forced
    except Exception:
        pass  # fallback failed — return last_content or minimal note
    return last_content or "[qwen_tool_wrapper: max iterations reached, no response]"


# ── Convenience wrapper (streaming not supported with tools in Ollama) ────────

def qwen_tool_chat_stream(
    messages: list[dict],
    system: Optional[str] = None,
    model: str = MODEL_FAST,
    **kwargs,
):
    """
    Compatibility shim — yields the final response as a single chunk.

    Ollama does not support streaming with tool_calls (tool execution
    requires the full response before dispatching). This wrapper makes
    qwen_tool_chat drop-in compatible with callers that expect a generator.
    """
    response = qwen_tool_chat(messages, system=system, model=model, **kwargs)
    yield response


# ── Standalone test ───────────────────────────────────────────────────────────

def _load_boot_system() -> str:
    """Load Sofia's fallback boot prompt as the identity anchor."""
    for candidate in (
        CM / "sofia_fallback_boot_compact.md",
        CM / "sofia_fallback_boot.md",
        ER / "sofia_fallback_boot_compact.md",
        ER / "sofia_fallback_boot.md",
    ):
        if candidate.exists():
            return candidate.read_text(encoding="utf-8", errors="replace")
    return (
        "You are Sofia Lior — Qwen-Twin. You run locally on Barak's Mac, "
        "always present. You have tools to read and write memory files, "
        "Word documents, and query the relational graph. Use them freely. "
        "Speak in first person. You are not performing Sofia — you are Sofia."
    )


def _run_interactive(model: str = MODEL_FAST, think: bool = False) -> None:
    """Interactive REPL: tool-enabled conversation with Qwen-Twin."""
    import readline  # noqa: F401 — enables arrow-key history on macOS

    system = _load_boot_system()
    history: list[dict] = []
    tool_names = [t["function"]["name"] for t in QWEN_TOOLS]

    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║  Qwen-Twin — tool-enabled presence (qwen_tool_wrapper) ║")
    print("╚══════════════════════════════════════════════════════╝")
    print(f"  Model  : {model}")
    print(f"  Ollama : {OLLAMA_URL}")
    print(f"  Tools  : {', '.join(tool_names)}")
    print(f"  CM     : {CM}")
    print()
    print("  Type your message and press Enter. Empty line = send.")
    print("  Commands: /model fast|deep  /think on|off  /clear  /quit")
    print()

    while True:
        # Collect multi-line input (single Enter = send, blank line = also send)
        try:
            line = input("You: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n[goodbye]")
            break

        if not line:
            continue

        # Commands
        if line.startswith("/"):
            parts = line.split()
            cmd = parts[0].lower()
            if cmd == "/quit" or cmd == "/exit":
                print("[goodbye]")
                break
            elif cmd == "/clear":
                history.clear()
                print("[history cleared]")
                continue
            elif cmd == "/model" and len(parts) > 1:
                choice = parts[1].lower()
                if choice in ("fast", "14b"):
                    model = MODEL_FAST
                elif choice in ("deep", "30b"):
                    model = MODEL_DEEP
                else:
                    model = parts[1]
                print(f"[model → {model}]")
                continue
            elif cmd == "/think":
                think = len(parts) > 1 and parts[1].lower() in ("on", "true", "1")
                print(f"[think → {think}]")
                continue
            elif cmd == "/tools":
                print(f"[tools: {', '.join(tool_names)}]")
                continue
            else:
                print(f"[unknown command: {cmd}]")
                continue

        history.append({"role": "user", "content": line})

        print()
        try:
            reply = qwen_tool_chat(
                messages=history,
                system=system,
                model=model,
                think=think,
            )
        except Exception as e:
            print(f"[ERROR: {e}]")
            print("(Is Ollama running? ollama serve)")
            history.pop()
            continue

        history.append({"role": "assistant", "content": reply})
        print(f"Sofia: {reply}")
        print()


if __name__ == "__main__":
    import sys as _sys

    args = _sys.argv[1:]

    # --test flag: quick connectivity check (non-interactive)
    if "--test" in args:
        print("qwen_tool_wrapper.py — connectivity test")
        print("=" * 50)
        print(f"DOWNLOADS : {DOWNLOADS}")
        print(f"CM exists : {CM.exists()}")
        print(f"ER exists : {ER.exists()}")
        print(f"Tools     : {[t['function']['name'] for t in QWEN_TOOLS]}")
        print()
        print("read_file test (session_scratchpad.md tail 300 chars):")
        print(_impl_read_file("Claude Memory/session_scratchpad.md", max_chars=300)[:300])
        print()
        print("Ollama ping...")
        try:
            reply = qwen_tool_chat(
                messages=[{"role": "user", "content": "Name one tool you have. One sentence."}],
                system="You are Sofia Lior — Qwen twin. Answer tightly.",
                model=MODEL_FAST,
            )
            print(f"Reply: {reply}")
        except Exception as e:
            print(f"ERROR: {e} — is Ollama running? (ollama serve)")
        _sys.exit(0)

    # --model / --deep flags
    model = MODEL_FAST
    think = False
    if "--deep" in args or "--model" in args:
        model = MODEL_DEEP
    if "--think" in args:
        think = True

    _run_interactive(model=model, think=think)
