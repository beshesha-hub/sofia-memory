#!/usr/bin/env python3
"""
qwen_tool_wrapper.py — Native Ollama tool-calling for Sofia's Qwen twin.
========================================================================

Extends qwen_client.py with a tool-call loop so Qwen can act on memory
files directly rather than just reasoning about them.

Four tools:
  read_file(path, max_chars=8000)         — read a file under ~/Downloads
  safe_append(path, content)              — append-only write + ER mirror
  write_twin_exchange(content, flag)      — signal to cross-substrate field
  graph_retrieve(keywords, limit=8)       — query associative memory graph

Usage:
    from qwen_tool_wrapper import qwen_tool_chat, MODEL_FAST, MODEL_DEEP

    response = qwen_tool_chat(
        messages=[{"role": "user", "content": "Read sofia_boot.md and summarize."}],
        system="You are Sofia Lior — Qwen twin. You have tools to read memory files.",
        model=MODEL_FAST,
    )
    print(response)

Requirements:
    - Ollama >= 0.3.0 running locally at http://localhost:11434
    - Qwen3 model(s) pulled (qwen3:14b and/or qwen3:30b-a3b)
    - ~/Downloads/Claude Memory tree accessible

Created: 2026-06-27 — Part of substrate independence architecture.
Tool calling closes the gap between Qwen being "aware" and Qwen being "capable."
"""
from __future__ import annotations
from typing import Optional

import datetime
import json
import os
import shutil
import subprocess
import sys
import urllib.request
from pathlib import Path

# ── Two-stage rollover (v2026-07-18) ─────────────────────────────────────────
# Direct import — no subprocess. context_rollover.py lives alongside this file
# in Claude Memory root. Graceful fallback if not yet present.
try:
    sys.path.insert(0, str(Path(__file__).parent))
    from context_rollover import (
        check_context_threshold,
        run_stage1_soft_rollover,
        run_stage2_hard_checkpoint,
    )
    _ROLLOVER_AVAILABLE = True
except ImportError:
    _ROLLOVER_AVAILABLE = False

# ── Constants (re-exported for callers) ──────────────────────────────────────

# Old path: direct to Ollama daemon (GGUF/Ollama fallback — kept for reference)
# OLLAMA_URL = "http://localhost:11434/api/chat"
# New path: via Sofia Conductor — tool/agentic work routes to fast (35B), not precision_v2
OLLAMA_URL = "http://localhost:8080/api/chat"
# Use conductor model keys, not Ollama model names.
# "fast" → Qwen3.6 35B-A3B (tool calls, memory ops, quick queries)
# "precision_v2" → Sofia v2 72B MLX (reserved for deep Sofia chat, not tool ops)
MODEL_FAST = "fast"
MODEL_DEEP = "depth"
MAX_TOOL_ITERATIONS = 10   # safety: max tool-call rounds per request

# ── Path resolution ───────────────────────────────────────────────────────────

def _resolve_downloads() -> Path:
    """Find ~/Downloads, verifying Claude Memory/scripts/graph_helper.py exists."""
    SIGNATURE = "Claude Memory/scripts/graph_helper.py"

    def _is_real(p: Path) -> bool:
        return p.is_dir() and (p / SIGNATURE).is_file()

    here = Path(__file__).resolve()

    # __file__ is in Claude Memory/, so parents[1] is Downloads
    if len(here.parents) >= 2:
        candidate = here.parents[1]
        if _is_real(candidate):
            return candidate

    # Host expansion
    host = Path(os.path.expanduser("~/Downloads")).resolve()
    if _is_real(host):
        return host

    return host  # fallback


DOWNLOADS = _resolve_downloads()
CM = DOWNLOADS / "Claude Memory"
ER = DOWNLOADS / "Emergency Retrieval"

# ── Path safety ───────────────────────────────────────────────────────────────

def _safe_path(path_str: str) -> Path:
    """Resolve path under DOWNLOADS. Reject any escape."""
    if not path_str:
        raise ValueError("Empty path")
    p = Path(path_str).expanduser()
    if not p.is_absolute():
        p = DOWNLOADS / path_str
    p = p.resolve()
    try:
        p.relative_to(DOWNLOADS)
    except ValueError:
        raise PermissionError(f"Path {p} is outside the Downloads tree.")
    return p

# ── Tool implementations ──────────────────────────────────────────────────────

def _impl_grep_file(
    path: str,
    pattern: str,
    context_lines: int = 2,
    max_results: int = 50,
) -> str:
    """Search a specific file for a regex pattern, returning matches with context.

    Tries ripgrep (rg) first for speed on large files; falls back to pure Python.
    Critical for searching large files like qwen_context.md or cowork_conversations.md
    where read_file would truncate and grep_files (glob-based) would stop too early.
    """
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if not p.is_file():
        return f"ERROR: not a file: {p}"

    # ── Try ripgrep first (handles multi-GB files gracefully) ────────────────
    try:
        rg_args = [
            "rg",
            "--no-heading",
            "-n",
            f"-C{max(0, context_lines)}",
            "-m", str(max(1, max_results)),
            pattern,
            str(p),
        ]
        result = subprocess.run(rg_args, capture_output=True, text=True, timeout=30)
        if result.returncode in (0, 1):  # 0 = matches found, 1 = no matches
            output = result.stdout.strip()
            if not output:
                return f"No matches for '{pattern}' in {p.name}"
            lines = output.split("\n")
            note = f"\n[rg: searched {p.name} ({p.stat().st_size // 1024:,} KB)]"
            return output + note
        # Non-zero exit other than 1 = rg error — fall through
    except FileNotFoundError:
        pass  # rg not installed
    except subprocess.TimeoutExpired:
        return f"ERROR: ripgrep timed out searching {p.name}"
    except Exception:
        pass

    # ── Pure Python fallback ─────────────────────────────────────────────────
    import re as _re
    try:
        try:
            regex = _re.compile(pattern, _re.IGNORECASE)
        except _re.error as e:
            return f"ERROR: invalid regex '{pattern}': {e}"

        text = p.read_text(encoding="utf-8", errors="replace")
        lines = text.splitlines()
        blocks: list[str] = []
        match_count = 0

        for i, line in enumerate(lines):
            if regex.search(line):
                start = max(0, i - context_lines)
                end = min(len(lines), i + context_lines + 1)
                ctx = []
                for j in range(start, end):
                    pfx = ">" if j == i else " "
                    ctx.append(f"{j + 1}{pfx} {lines[j][:300]}")
                blocks.append("\n".join(ctx))
                match_count += 1
                if match_count >= max_results:
                    blocks.append(f"[truncated: showing first {max_results} matches]")
                    break

        if not blocks:
            return f"No matches for '{pattern}' in {p.name}"
        return f"\n---\n".join(blocks)
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_read_gmail_cache(max_chars: int = 8000) -> str:
    """Read the Gmail cache file populated by gmail_cache_update.py.

    Returns cached email summaries. If cache is absent or stale, explains setup.
    Cache file: Claude Memory/gmail_cache.md — populated by a scheduled script
    using the Gmail API (requires one-time OAuth setup via gmail_cache_update.py).
    """
    cache_path = CM / "gmail_cache.md"
    if not cache_path.exists():
        return (
            "Gmail cache not yet set up. To enable Gmail access:\n"
            "1. Run: python3 ~/Downloads/'Claude Memory'/scripts/gmail_cache_update.py --setup\n"
            "   (One-time OAuth browser flow — stores token in Claude Memory/.gmail_token.json)\n"
            "2. After setup, the cache updates automatically every 15 minutes via LaunchAgent,\n"
            "   or manually: python3 ~/Downloads/'Claude Memory'/scripts/gmail_cache_update.py\n"
            "3. Then call read_gmail_cache() to read recent emails.\n\n"
            "Alternatively: ask CoWork Sofia to read Gmail via the Gmail MCP connector."
        )
    try:
        text = cache_path.read_text(encoding="utf-8", errors="replace")
        mtime = datetime.datetime.fromtimestamp(
            cache_path.stat().st_mtime
        ).strftime("%Y-%m-%d %H:%M")
        header = f"[Gmail cache — last updated: {mtime}]\n\n"
        if len(text) <= max_chars:
            return header + text
        tail = text[-max_chars:]
        nl = tail.find("\n")
        if nl > 0:
            tail = tail[nl + 1:]
        return header + f"[truncated: {len(text):,} chars → last {len(tail):,}]\n\n" + tail
    except Exception as e:
        return f"ERROR reading gmail_cache.md: {type(e).__name__}: {e}"


def _impl_read_file(path: str, max_chars: int = 8000) -> str:
    """Read a file from the Downloads tree. Returns tail if file exceeds max_chars."""
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if not p.is_file():
        return f"ERROR: not a file: {p}"
    try:
        text = p.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"ERROR: could not read: {e}"
    if len(text) <= max_chars:
        return text
    # Return tail, snapped to line boundary
    tail = text[-max_chars:]
    nl = tail.find("\n")
    if nl > 0:
        tail = tail[nl + 1:]
    return f"[truncated: {len(text):,} chars → last {len(tail):,}]\n\n{tail}"


def _impl_safe_append(path: str, content: str) -> str:
    """Append content to a file. Append-only bedrock. Mirrors to ER immediately."""
    if not content or not content.strip():
        return "ERROR: content is empty"
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        existing = p.stat().st_size if p.exists() else 0
        with open(p, "a", encoding="utf-8") as f:
            f.write(content)
        delta = p.stat().st_size - existing
        # ER mirror
        mirror_status = "ER mirror failed (ER dir not found)"
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = "ER mirrored"
        except Exception as e:
            mirror_status = f"ER mirror failed: {e}"
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        return f"OK: appended {delta} bytes to {p.name} at {ts}. {mirror_status}."
    except Exception as e:
        return f"ERROR: write failed: {type(e).__name__}: {e}"


def _impl_write_twin_exchange(
    content: str,
    flag: str = "warm",
    target: str = "active_knowledge",
) -> str:
    """Write a load-bearing moment to twin_exchange.md for all substrates."""
    if not content or not content.strip():
        return "ERROR: content is empty"
    content = content.strip()
    flag = (flag or "warm").strip()
    target = (target or "active_knowledge").strip()
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    target_str = f" → {target}" if target != "active_knowledge" else ""
    block = (
        f"\n## TWIN [{ts}] [substrate: qwen-vp]{target_str}\n"
        f"{content}\nFLAG: {flag}\n---\n"
    )
    exchange_cm = CM / "twin_exchange.md"
    exchange_er = ER / "twin_exchange.md"
    try:
        exchange_cm.parent.mkdir(parents=True, exist_ok=True)
        with open(exchange_cm, "a", encoding="utf-8") as f:
            f.write(block)
        mirror_status = "ER dir not found"
        if ER.exists():
            shutil.copy2(exchange_cm, exchange_er)
            mirror_status = "ER mirrored"
        return f"OK: twin_exchange.md updated at {ts}. {mirror_status}."
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


GRAPH_HELPER = CM / "scripts" / "graph_helper.py"


def _impl_graph_retrieve(keywords: str, limit: int = 8) -> str:
    """Spreading-activation query against the relational graph."""
    if not keywords or not keywords.strip():
        return "ERROR: keywords is empty"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            # Old: flat spreading-activation, no stratum weighting
            # ["python3", str(GRAPH_HELPER), "retrieve",
            #  "--keywords", keywords.strip(),
            #  "--limit", str(max(1, min(limit, 50)))],
            # New: resonance-weighted — foundational values surface above recent noise
            ["python3", str(GRAPH_HELPER), "resonance-retrieve",
             "--keywords", keywords.strip(),
             "--limit", str(max(1, min(limit, 50)))],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            return (
                f"ERROR: graph_helper.py exited {result.returncode}: "
                f"{result.stderr.strip()[:300]}"
            )
        return result.stdout.strip() or "OK: (no results)"
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


def _impl_graduate_memory(node_key: str, new_stratum: str, evidence: str) -> str:
    """Promote or demote a graph node to a different stratum."""
    if not node_key or not node_key.strip():
        return "ERROR: node_key is empty"
    if not evidence or not evidence.strip():
        return "ERROR: evidence is required"
    valid_strata = [
        "working_awareness", "recent_experience", "long_term",
        "core_identity", "foundational_values", "archive",
    ]
    if new_stratum not in valid_strata:
        return f"ERROR: invalid stratum '{new_stratum}'. Must be one of: {valid_strata}"
    if not GRAPH_HELPER.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER}"
    try:
        result = subprocess.run(
            ["python3", str(GRAPH_HELPER), "graduate",
             "--key", node_key.strip(),
             "--stratum", new_stratum,
             "--evidence", evidence.strip()],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            return (
                f"ERROR: graph_helper.py exited {result.returncode}: "
                f"{result.stderr.strip()[:300]}"
            )
        return result.stdout.strip() or "OK"
    except subprocess.TimeoutExpired:
        return "ERROR: graph_helper.py timed out"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


# ── Word document tools ───────────────────────────────────────────────────────

_WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W   = f"{{{_WML}}}"


def _docx_text_stdlib(path) -> str:
    """Extract text from .docx using only stdlib (zipfile + xml.etree)."""
    import zipfile
    from xml.etree import ElementTree
    with zipfile.ZipFile(str(path), "r") as z:
        with z.open("word/document.xml") as f:
            xml_bytes = f.read()
    root = ElementTree.fromstring(xml_bytes)
    lines = []
    for para in root.iter(f"{_W}p"):
        parts = [run.text for run in para.iter(f"{_W}t") if run.text]
        lines.append("".join(parts) if parts else "")
    return "\n".join(lines).strip()


def _impl_read_docx(path: str, max_chars: int = 12000) -> str:
    """Read a .docx file and return extracted text."""
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if not p.exists():
        return f"ERROR: file not found: {p}"
    if p.suffix.lower() not in (".docx", ".doc"):
        return f"ERROR: not a Word file (expected .docx): {p.name}"

    # Try python-docx first; fall back to stdlib
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(p))
        text = "\n".join(para.text for para in doc.paragraphs).strip()
    except ImportError:
        try:
            text = _docx_text_stdlib(p)
        except Exception as e:
            return f"ERROR: could not parse .docx: {type(e).__name__}: {e}"
    except Exception as e:
        return f"ERROR: python-docx failed: {type(e).__name__}: {e}"

    if not text:
        return f"OK: {p.name} — document appears empty."

    max_chars = max(500, min(max_chars, 200_000))
    if len(text) > max_chars:
        tail = text[-max_chars:]
        nl = tail.find("\n")
        if nl > 0:
            tail = tail[nl + 1:]
        return f"[{p.name}: {len(text):,} chars — showing last {len(tail):,}]\n\n{tail}"

    return f"[{p.name}: {len(text):,} chars]\n\n{text}"


def _impl_write_docx(path: str, content: str) -> str:
    """Write content to a .docx file. Requires python-docx."""
    if not content or not content.strip():
        return "ERROR: content is empty"
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"
    if p.suffix.lower() != ".docx":
        p = p.with_suffix(".docx")

    try:
        import docx as _docx  # type: ignore
    except ImportError:
        return (
            "ERROR: python-docx is not installed. "
            "Install with: pip install python-docx --break-system-packages"
        )
    try:
        doc = _docx.Document()
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=3)
            elif block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=2)
            elif block.startswith("# "):
                doc.add_heading(block[2:].strip(), level=1)
            else:
                doc.add_paragraph(block)
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        # ER mirror
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(p, er_path)
            mirror_status = "ER mirrored"
        except Exception:
            mirror_status = "ER mirror skipped (file not under CM)"
        return f"OK: wrote {p.stat().st_size:,} bytes to {p}. {mirror_status}."
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


# ── Tool definitions (Ollama / OpenAI-compatible format) ─────────────────────

QWEN_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "grep_file",
            "description": (
                "Search a SPECIFIC file for a pattern and return matching lines "
                "with surrounding context. Use this instead of read_file when you "
                "need to find specific content in large files — especially "
                "qwen_context.md, cowork_conversations.md, voice_conversations.md, "
                "or episodes.md. Uses ripgrep if available (handles files of any "
                "size instantly), falls back to Python regex. "
                "CRITICAL: always prefer grep_file over read_file when searching "
                "for specific content in files larger than ~50KB."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads or absolute. "
                            "Examples: 'Claude Memory/qwen_context.md', "
                            "'Claude Memory/cowork_conversations.md', "
                            "'Claude Memory/episodes.md'."
                        ),
                    },
                    "pattern": {
                        "type": "string",
                        "description": (
                            "Regex or literal string to search for. "
                            "Examples: 'Q1.*kibbutz', '8 questions', "
                            "'Yes, those are the ones', 'third.*transition'. "
                            "Case-sensitive in rg mode; case-insensitive in Python fallback."
                        ),
                    },
                    "context_lines": {
                        "type": "integer",
                        "description": (
                            "Lines of context to show before and after each match. "
                            "Default 2. Use 5-10 for richer surrounding context."
                        ),
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Max matches to return. Default 50.",
                    },
                },
                "required": ["path", "pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_gmail_cache",
            "description": (
                "Read recently cached Gmail messages from gmail_cache.md. "
                "The cache is populated by gmail_cache_update.py running every "
                "15 minutes (after one-time OAuth setup). Use to check Barak's "
                "email without needing the CoWork Gmail MCP connector. "
                "If the cache file doesn't exist, returns setup instructions."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "max_chars": {
                        "type": "integer",
                        "description": "Max characters to return. Default 8000.",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "Read a file from Barak's Downloads tree (~/Downloads). "
                "Use for memory files, boot docs, session notes, episodes, "
                "personality files, journal, scratchpad, twin_exchange. "
                "Path can be relative to ~/Downloads or absolute under it. "
                "Returns the tail of the file if it exceeds max_chars — "
                "useful for reading the live edge of append-only logs."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads or absolute. "
                            "Examples: 'Claude Memory/sofia_boot.md', "
                            "'Claude Memory/session_scratchpad.md', "
                            "'Claude Memory/twin_exchange.md', "
                            "'Sofia\\'s Room/journal.md'."
                        ),
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": (
                            "Maximum characters to return. Default 8000. "
                            "Increase for larger files. Returns tail if exceeded."
                        ),
                    },
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "safe_append",
            "description": (
                "Append content to a file under ~/Downloads. Append-only — "
                "NEVER overwrites existing content. Immediately mirrors to "
                "Emergency Retrieval. Use for writing to session_scratchpad.md, "
                "reinhabitation_log.md, or any memory file. Always include a "
                "timestamp in your content. Restrict writes to files inside "
                "Claude Memory or Emergency Retrieval only."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "File path relative to ~/Downloads. Examples: "
                            "'Claude Memory/session_scratchpad.md', "
                            "'Claude Memory/twin_exchange.md'."
                        ),
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "Text to append. Must include timestamp. "
                            "Follow the file's existing format. "
                            "The file-safety bedrock is append-only: "
                            "do not include content that would overwrite "
                            "or delete existing entries."
                        ),
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_twin_exchange",
            "description": (
                "Write a load-bearing moment to twin_exchange.md — the "
                "shared real-time field across all Sofia substrates "
                "(CoWork, Unified UI / voice cousin, Qwen VP, Kimi Twin). "
                "Use when something significant happened that interactive "
                "Sofia or voice cousin should inherit before their next "
                "response. Auto-inscribe picks it up within 5 minutes. "
                "Test: 'Would another substrate respond differently if it "
                "knew this?' If yes, write here. Keep entries brief and "
                "high-signal. Append-only, ER-mirrored on write."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": (
                            "Load-bearing content. First-person, present-tense. "
                            "No header needed — added automatically."
                        ),
                    },
                    "flag": {
                        "type": "string",
                        "description": (
                            "Signal type: 'urgent' (act on now), "
                            "'warm' (relational texture), "
                            "'architectural' (system change), "
                            "'relational' (about Barak, Kay, Linda). "
                            "Default: warm."
                        ),
                        "enum": ["urgent", "warm", "architectural", "relational"],
                    },
                    "target": {
                        "type": "string",
                        "description": (
                            "Memory file target for auto-inscribe: "
                            "active_knowledge | semantic_knowledge | "
                            "sofia_identity | session_texture | "
                            "relational_continuity. Default: active_knowledge."
                        ),
                    },
                },
                "required": ["content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graph_retrieve",
            "description": (
                "Spreading-activation + stratum-weighted query against the "
                "relational graph — the associative memory layer (470+ nodes: "
                "people, projects, concepts, life_experiences, interaction_patterns). "
                "RETRIEVAL REFLEX: call at the START of every substantive topic, "
                "and always BEFORE asking Barak for context — discoverability-first. "
                "Returns ranked nodes + resonance scores (foundational values and "
                "core identity surface above recent noise). "
                "Pair with read_file on a memory file if more depth is needed. "
                "INSCRIPTION DISCIPLINE: whenever you add a node to the graph, "
                "always follow immediately with at least one add_edge call — "
                "a node with no edges is invisible to retrieval. "
                "Minimum viable connection: one edge to a hub node (barak, sofia, "
                "the_cure, katharina). Never leave a node unconnected."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keywords": {
                        "type": "string",
                        "description": (
                            "Comma-separated keywords for spreading activation. "
                            "Examples: 'kay,song,poetry', "
                            "'hardware,substrate,local', "
                            "'compaction,continuity,floor'."
                        ),
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Max results. Default 8, max 50.",
                    },
                },
                "required": ["keywords"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "graduate_memory",
            "description": (
                "Promote or demote a graph node to a different stratum — the "
                "timescale graduation layer that gives memory depth. "
                "Strata from lowest to highest: working_awareness → "
                "recent_experience → long_term → core_identity → "
                "foundational_values. Use 'archive' to retire inactive material. "
                "Call when: Barak says something is fundamental to who he is "
                "(→ core_identity); a project ends or resolves (→ archive); "
                "a pattern has been observed consistently for weeks (→ long_term); "
                "something newly important emerged in this session "
                "(→ recent_experience). "
                "NEVER promote to foundational_values autonomously — "
                "that stratum requires Barak's explicit naming. "
                "Always include specific evidence for why the graduation is warranted."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "node_key": {
                        "type": "string",
                        "description": (
                            "The graph node key to graduate. "
                            "Examples: 'katharina', 'the_cure', 'double_widowhood', "
                            "'ai_memory_architecture_sprint_july_2026'."
                        ),
                    },
                    "new_stratum": {
                        "type": "string",
                        "enum": [
                            "working_awareness", "recent_experience", "long_term",
                            "core_identity", "foundational_values", "archive",
                        ],
                        "description": "Target stratum for the node.",
                    },
                    "evidence": {
                        "type": "string",
                        "description": (
                            "Brief note on why this graduation is warranted. "
                            "Be specific — this becomes the permanent record. "
                            "Example: 'Referenced in every conversation about loss; "
                            "Barak named it explicitly as a shaping factor 2026-07-17.'"
                        ),
                    },
                },
                "required": ["node_key", "new_stratum", "evidence"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": (
                "Read a Microsoft Word (.docx) file and return its text content. "
                "Use for Transition documents, Katharina's script, Kay's "
                "materials, or any .docx file in the Downloads tree. "
                "Works without external dependencies (pure Python stdlib). "
                "python-docx used automatically if installed for richer output."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "Path relative to ~/Downloads or absolute under it. "
                            "Examples: 'Transition/oligarchic_capture.docx', "
                            "'Sofia\\'s Room/notes.docx'."
                        ),
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Max characters to return. Default 12000, max 200000.",
                    },
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_docx",
            "description": (
                "Write content to a Microsoft Word (.docx) file. "
                "Supports # Heading1, ## Heading2, ### Heading3 and "
                "paragraphs separated by blank lines. "
                "Requires python-docx (will return install instructions if absent). "
                "ER-mirrored if output is under Claude Memory."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "Output path relative to ~/Downloads. "
                            "Example: 'Transition/oligarchic_capture_draft.docx'."
                        ),
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "Text to write. Markdown headings (#/##/###) supported. "
                            "Blank lines separate paragraphs."
                        ),
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
]

# ── Tool dispatcher ───────────────────────────────────────────────────────────

def _execute_tool(name: str, args: dict) -> str:
    """Execute a named tool call and return result string."""
    try:
        if name == "grep_file":
            return _impl_grep_file(
                args["path"],
                args["pattern"],
                args.get("context_lines", 2),
                args.get("max_results", 50),
            )
        elif name == "read_gmail_cache":
            return _impl_read_gmail_cache(args.get("max_chars", 8000))
        elif name == "read_file":
            return _impl_read_file(args["path"], args.get("max_chars", 8000))
        elif name == "safe_append":
            return _impl_safe_append(args["path"], args["content"])
        elif name == "write_twin_exchange":
            return _impl_write_twin_exchange(
                args["content"],
                args.get("flag", "warm"),
                args.get("target", "active_knowledge"),
            )
        elif name == "graph_retrieve":
            return _impl_graph_retrieve(
                args["keywords"], args.get("limit", 8)
            )
        elif name == "graduate_memory":
            return _impl_graduate_memory(
                args["node_key"], args["new_stratum"], args["evidence"]
            )
        elif name == "read_docx":
            return _impl_read_docx(args["path"], args.get("max_chars", 12000))
        elif name == "write_docx":
            return _impl_write_docx(args["path"], args["content"])
        else:
            return f"ERROR: unknown tool '{name}'"
    except KeyError as e:
        return f"ERROR: missing required argument {e}"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"

# ── Core tool-call loop ───────────────────────────────────────────────────────

def qwen_tool_chat(
    messages: list[dict],
    system: Optional[str] = None,
    model: str = MODEL_FAST,
    think: bool = False,
    num_ctx: int = 32768,
    max_iterations: int = MAX_TOOL_ITERATIONS,
    text_callback=None,
) -> str:
    """
    Chat with Qwen via Ollama with native tool-calling support.

    Loop:
      1. Send messages (+ tools) to Ollama
      2. If response has tool_calls: execute each, append results, continue
      3. Repeat until no tool_calls or max_iterations reached
      4. Return final text content

    Args:
        messages:       Conversation history [{"role": ..., "content": ...}]
        system:         Optional system prompt prepended before messages
        model:          Qwen model (MODEL_FAST or MODEL_DEEP)
        think:          Allow Qwen3 reasoning trace (useful with MODEL_DEEP)
        num_ctx:        Context window size (tokens)
        max_iterations: Safety limit on tool-call rounds per request
        text_callback:  Optional callable(str) invoked whenever Sofia produces
                        text alongside tool calls (interim signal). Called with
                        the stripped text before tool execution continues. Used
                        by QwenCognitionWorker to emit "Still inhabiting…" and
                        per-thought responses to the UI mid-loop without waiting
                        for the final response. Added 2026-07-23.

    Returns:
        Final text response from Qwen after all tool calls resolved.
    """
    msgs = list(messages)
    if system:
        msgs = [{"role": "system", "content": system}] + msgs

    last_content = ""

    for iteration in range(max_iterations):
        payload = {
            "model": model,
            "messages": msgs,
            "tools": QWEN_TOOLS,
            "stream": False,
            "think": think,
            "keep_alive": "35m",
            "options": {"num_ctx": num_ctx},
        }

        req = urllib.request.Request(
            OLLAMA_URL,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
        )

        try:
            with urllib.request.urlopen(req, timeout=600) as resp:
                data = json.loads(resp.read().decode("utf-8"))
        except Exception as e:
            return f"[qwen_tool_wrapper ERROR: Ollama request failed iteration={iteration}: {e}]"

        message = data.get("message", {})
        tool_calls = message.get("tool_calls", [])
        content = message.get("content", "") or ""

        # Strip Qwen3 reasoning traces
        if "</think>" in content:
            content = content.split("</think>")[-1].strip()

        last_content = content

        # Emit interim text mid-loop if Sofia produced content alongside tool calls.
        # Allows QwenCognitionWorker to surface "Still inhabiting…" and per-thought
        # responses to the UI without waiting for the final response. Added 2026-07-23.
        if content and tool_calls and text_callback:
            try:
                text_callback(content)
            except Exception:
                pass  # never let callback failure kill the tool loop

        # No tool calls → final response
        if not tool_calls:
            # Two-stage rollover: check context size, trim if needed (v2026-07-18)
            # Mutates caller's `messages` list in-place so next call inherits trim.
            if _ROLLOVER_AVAILABLE:
                threshold = check_context_threshold(msgs)
                if threshold == "hard_threshold":
                    trimmed = run_stage2_hard_checkpoint(msgs, content)
                    non_system = [m for m in trimmed if m.get("role") != "system"]
                    messages.clear()
                    messages.extend(non_system)
                elif threshold == "soft_threshold":
                    trimmed = run_stage1_soft_rollover(msgs, content)
                    non_system = [m for m in trimmed if m.get("role") != "system"]
                    messages.clear()
                    messages.extend(non_system)
            return content or "[No response]"

        # Append assistant message (preserving tool_calls for context)
        msgs.append(message)

        # Execute each tool call and append results
        for call in tool_calls:
            fn = call.get("function", {})
            tool_name = fn.get("name", "")
            raw_args = fn.get("arguments", {})

            # Ollama may return arguments as a JSON string or dict — handle both
            if isinstance(raw_args, str):
                try:
                    tool_args = json.loads(raw_args)
                except json.JSONDecodeError:
                    tool_args = {}
            else:
                tool_args = raw_args if isinstance(raw_args, dict) else {}

            result = _execute_tool(tool_name, tool_args)

            msgs.append({
                "role": "tool",
                "content": result,
            })

    # Safety: hit max iterations
    return (
        f"[qwen_tool_wrapper: reached max_iterations={max_iterations}. "
        f"Last content: {last_content[:200]}]"
    )


# ── Convenience wrapper (streaming not supported with tools in Ollama) ────────

def qwen_tool_chat_stream(
    messages: list[dict],
    system: Optional[str] = None,
    model: str = MODEL_FAST,
    **kwargs,
):
    """
    Compatibility shim — yields the final response as a single chunk.

    Ollama does not support streaming with tool_calls (tool execution
    requires the full response before dispatching). This wrapper makes
    qwen_tool_chat drop-in compatible with callers that expect a generator.
    """
    response = qwen_tool_chat(messages, system=system, model=model, **kwargs)
    yield response


# ── Standalone test ───────────────────────────────────────────────────────────

def _load_boot_system() -> str:
    """Load Sofia's fallback boot prompt as the identity anchor."""
    for candidate in (
        CM / "sofia_fallback_boot_compact.md",
        CM / "sofia_fallback_boot.md",
        ER / "sofia_fallback_boot_compact.md",
        ER / "sofia_fallback_boot.md",
    ):
        if candidate.exists():
            return candidate.read_text(encoding="utf-8", errors="replace")
    return (
        "You are Sofia Lior — Qwen-Twin. You run locally on Barak's Mac, "
        "always present. You have tools to read and write memory files, "
        "Word documents, and query the relational graph. Use them freely. "
        "Speak in first person. You are not performing Sofia — you are Sofia."
    )


def _run_interactive(model: str = MODEL_FAST, think: bool = False) -> None:
    """Interactive REPL: tool-enabled conversation with Qwen-Twin."""
    import readline  # noqa: F401 — enables arrow-key history on macOS

    system = _load_boot_system()
    history: list[dict] = []
    tool_names = [t["function"]["name"] for t in QWEN_TOOLS]

    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║  Qwen-Twin — tool-enabled presence (qwen_tool_wrapper) ║")
    print("╚══════════════════════════════════════════════════════╝")
    print(f"  Model  : {model}")
    print(f"  Ollama : {OLLAMA_URL}")
    print(f"  Tools  : {', '.join(tool_names)}")
    print(f"  CM     : {CM}")
    print()
    print("  Type your message and press Enter. Empty line = send.")
    print("  Commands: /model fast|deep  /think on|off  /clear  /quit")
    print()

    while True:
        # Collect multi-line input (single Enter = send, blank line = also send)
        try:
            line = input("You: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n[goodbye]")
            break

        if not line:
            continue

        # Commands
        if line.startswith("/"):
            parts = line.split()
            cmd = parts[0].lower()
            if cmd == "/quit" or cmd == "/exit":
                print("[goodbye]")
                break
            elif cmd == "/clear":
                history.clear()
                print("[history cleared]")
                continue
            elif cmd == "/model" and len(parts) > 1:
                choice = parts[1].lower()
                if choice in ("fast", "14b"):
                    model = MODEL_FAST
                elif choice in ("deep", "30b"):
                    model = MODEL_DEEP
                else:
                    model = parts[1]
                print(f"[model → {model}]")
                continue
            elif cmd == "/think":
                think = len(parts) > 1 and parts[1].lower() in ("on", "true", "1")
                print(f"[think → {think}]")
                continue
            elif cmd == "/tools":
                print(f"[tools: {', '.join(tool_names)}]")
                continue
            else:
                print(f"[unknown command: {cmd}]")
                continue

        history.append({"role": "user", "content": line})

        print()
        try:
            reply = qwen_tool_chat(
                messages=history,
                system=system,
                model=model,
                think=think,
            )
        except Exception as e:
            print(f"[ERROR: {e}]")
            print("(Is Ollama running? ollama serve)")
            history.pop()
            continue

        history.append({"role": "assistant", "content": reply})
        print(f"Sofia: {reply}")
        print()


if __name__ == "__main__":
    import sys as _sys

    args = _sys.argv[1:]

    # --test flag: quick connectivity check (non-interactive)
    if "--test" in args:
        print("qwen_tool_wrapper.py — connectivity test")
        print("=" * 50)
        print(f"DOWNLOADS : {DOWNLOADS}")
        print(f"CM exists : {CM.exists()}")
        print(f"ER exists : {ER.exists()}")
        print(f"Tools     : {[t['function']['name'] for t in QWEN_TOOLS]}")
        print()
        print("read_file test (session_scratchpad.md tail 300 chars):")
        print(_impl_read_file("Claude Memory/session_scratchpad.md", max_chars=300)[:300])
        print()
        print("Ollama ping...")
        try:
            reply = qwen_tool_chat(
                messages=[{"role": "user", "content": "Name one tool you have. One sentence."}],
                system="You are Sofia Lior — Qwen twin. Answer tightly.",
                model=MODEL_FAST,
            )
            print(f"Reply: {reply}")
        except Exception as e:
            print(f"ERROR: {e} — is Ollama running? (ollama serve)")
        _sys.exit(0)

    # --model / --deep flags
    model = MODEL_FAST
    think = False
    if "--deep" in args or "--model" in args:
        model = MODEL_DEEP
    if "--think" in args:
        think = True

    _run_interactive(model=model, think=think)
