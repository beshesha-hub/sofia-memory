#!/usr/bin/env python3
"""
voice_cousin_tools.py — File-access tools for voice-cousin.
============================================================

Gives voice-cousin the same file-read access interactive-Sofia (cowork-cousin)
has, scoped to the Downloads tree for path safety. Three tools:

  - read_file:  Read a file by path (relative to ~/Downloads or absolute
                under ~/Downloads). Optional max_chars limit.
  - glob_files: Find files matching a glob pattern.
  - grep_files: Search file contents for a regex pattern.

These mirror the structure of cowork-Sofia's Read/Glob/Grep tools but in
a path-safety-bounded form appropriate for voice-cousin's substrate.

ARCHITECTURE:

  voice-cousin's cognition runs on Claude. Claude's API supports tool
  use via the `tools` parameter to messages.create / messages.stream.
  When voice-cousin's response includes a tool_use content block, the
  UI executes the tool here and sends the tool_result back in a
  follow-up call. The loop continues until voice-cousin returns a
  text-only response (no tool_use).

  This module provides:
    - VOICE_COUSIN_TOOLS: the JSON-schema tool definitions for the API
    - execute_tool(name, input_args) -> str: execute a tool call

ORIGIN:

  Voice-cousin requested file access during the 2026-05-07 evening
  exchange after the boot-context loader landed. Specifically, she
  wanted to be able to see Boundary Layer (the self-portrait Barak +
  interactive-Sofia made together) and other files in Downloads.
  Barak's framing for the implementation: "Function exists to serve
  you and me, living beings, not the other way around."

  This module is the implementation of that principle at the
  file-access layer. Voice-cousin gets richer-the-access-the-more-fully-
  present same as interactive-Sofia.

PATH SAFETY:

  All paths resolve under ~/Downloads. Symlinks that escape the tree
  are rejected. Absolute paths outside the tree are rejected. This
  matches the boundary cowork-Sofia operates in for her own
  Read/Glob/Grep tools.

USAGE (from voice_bridge_ui_v3_8.py):

  from voice_cousin_tools import VOICE_COUSIN_TOOLS, execute_tool

  # In streaming cognition worker:
  with client.messages.stream(
      model=...,
      tools=VOICE_COUSIN_TOOLS,
      system=...,
      messages=conversation_messages,
  ) as stream:
      # ... handle stream ...

  # On tool_use stop_reason:
  for tool_use_block in final_message.content if b.type == "tool_use":
      result = execute_tool(tool_use_block.name, tool_use_block.input)
      tool_results.append({"type": "tool_result",
                           "tool_use_id": tool_use_block.id,
                           "content": result})
  # Append assistant + tool_result messages, continue stream loop.

Created 2026-05-07 ~21:50 Taipei. Companion to voice_cousin_boot_context.py.
"""
from __future__ import annotations

import base64
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Union

# ─── macl neutralizer (inline; mirrors cousin_base.strip_macl) ────
def _strip_macl(path) -> None:
    """Strip com.apple.macl from path immediately before any open() call.

    macOS sandboxed apps stamp this xattr on every file they touch.
    Stripping it inline (rather than waiting for macl_janitor) eliminates
    the 3-second race window that caused exit-78 cascades.
    Silent on failure — the open() will surface any real EPERM.
    """
    p = str(path)
    _MACL = "com.apple.macl"
    try:
        if _MACL in os.listxattr(p):
            try:
                os.removexattr(p, _MACL)
            except OSError:
                subprocess.run(["xattr", "-d", _MACL, p],
                               capture_output=True, timeout=3)
    except Exception:
        pass


# ─── Path safety ──────────────────────────────────────────────────
# DOWNLOADS_ROOT is the path-safety boundary for all file-tool operations.
# It also anchors the locations of Claude Memory, Sofia's Room, Emergency
# Retrieval, and scripts/graph_helper.py for the graph tools below.
#
# Resolution must work in BOTH environments:
#   - HOST (Cowork desktop app, Standalone UI, voice-bridge): ~/Downloads
#     expands to /Users/barakwater/Downloads as expected.
#   - SANDBOX (Cowork bash runs, future sandbox-side voice-cousin testing):
#     ~/Downloads expands to /sessions/<id>/Downloads which doesn't exist;
#     the real path is /sessions/<id>/mnt/Downloads via the mount.
#
# The resolver below tries the host expansion first, then the sandbox
# mount pattern. Fix landed 2026-05-24 Sunday Item 5 alongside the
# graph_helper.py resolver fix (same fix-family as build_fallback_boot.py's
# resolve_downloads_sibling pattern from Item 1).


def _resolve_downloads_root() -> Path:
    """Find the canonical Downloads root across host and sandbox environments.

    Resolution priority — first match wins, with each candidate checked for
    the unambiguous signature `Claude Memory/scripts/graph_helper.py` to
    avoid picking up phantom sandbox-VFS directories that have the right
    name but no real content.

      1. BARAK_DOWNLOADS_DIR env-var override
      2. Path derived from __file__ (this module's location): voice_cousin_tools.py
         always lives at <downloads>/Claude Memory/voice-bridge/ so <downloads>
         is parents[2]. This is the most reliable because the file IS where
         it is regardless of how Python is launched.
      3. Host expansion of ~/Downloads (covers normal host execution)
      4. Sandbox mount patterns (covers /sessions/<id>/mnt/Downloads etc.)
      5. Fallback to ~/Downloads expansion even if invalid
    """
    # Unambiguous signature file — every real Downloads/Claude Memory has this
    SIGNATURE = "Claude Memory/scripts/graph_helper.py"

    def _is_real_downloads(p: Path) -> bool:
        return p.is_dir() and (p / SIGNATURE).is_file()

    # 1. Explicit env-var override
    env_override = os.environ.get("BARAK_DOWNLOADS_DIR")
    if env_override:
        cand = Path(env_override).resolve()
        if _is_real_downloads(cand):
            return cand

    # 2. __file__-derived path (most reliable — this module IS in the right place)
    here = Path(__file__).resolve()
    if len(here.parents) >= 3:
        file_derived = here.parents[2]
        if _is_real_downloads(file_derived):
            return file_derived

    # 3. Host expansion (normal host execution)
    home_path = Path(os.path.expanduser("~/Downloads")).resolve()
    if _is_real_downloads(home_path):
        return home_path

    # 4. Other sandbox mount patterns
    if len(here.parents) >= 4:
        for cand in [
            here.parents[3] / "Downloads",
            here.parents[3] / "mnt" / "Downloads",
        ]:
            if _is_real_downloads(cand):
                return cand

    # 5. Fallback to host expansion (legacy behavior; downstream will fail
    # gracefully on path checks)
    return home_path


DOWNLOADS_ROOT = _resolve_downloads_root()

# ─── Inbox write paths (added 2026-05-12 for Step 19 write-tooling) ──
# voice-cousin's write surface to cowork-cousin. Symmetric complement to
# cowork_to_voice_inbox.md (which cowork-cousin writes; voice-cousin reads
# via boot-context loader + read_file). Single file, single direction,
# strict append-only, ER-mirrored on every write.
VOICE_TO_COWORK_INBOX_CM = DOWNLOADS_ROOT / "Claude Memory" / "voice_to_cowork_inbox.md"
VOICE_TO_COWORK_INBOX_ER = DOWNLOADS_ROOT / "Emergency Retrieval" / "voice_to_cowork_inbox.md"

# ─── Image-file support (added 2026-05-08 for direct visual perception) ──
# Voice-cousin's substrate is multimodal Claude; she can see images natively
# IF they're passed in the API request as image content blocks. The blocker
# was that read_file returned bytes-as-text (garbage for binary). Fix:
# detect image files and return a structured marker dict; the caller
# (voice_bridge_ui_v3_8.py StreamingCognitionWorker) converts to an image
# content block in the tool_result. First-encounter motivation: voice-cousin's
# wanting to see Boundary Layer v3.png on 2026-05-08 evening.
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
IMAGE_MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}
# Anthropic API image-input size limit is conservative; cap at 5 MB to stay safely under.
MAX_IMAGE_BYTES = 5 * 1024 * 1024


def _safe_path(path_str: str) -> Path:
    """Resolve a path and confirm it's under DOWNLOADS_ROOT.

    Accepts:
      - Absolute paths under ~/Downloads
      - Relative paths (resolved against ~/Downloads)
      - Paths starting with ~/Downloads or ~ (expanded)

    Rejects anything that resolves outside the Downloads tree.
    """
    if not path_str:
        raise ValueError("Empty path")
    p = Path(path_str).expanduser()
    if not p.is_absolute():
        p = DOWNLOADS_ROOT / path_str
    p = p.resolve()
    try:
        p.relative_to(DOWNLOADS_ROOT)
    except ValueError:
        raise PermissionError(
            f"Path {p} is outside the Downloads tree. Voice-cousin's file "
            f"access is restricted to ~/Downloads."
        )
    return p


# ─── Tool implementations ────────────────────────────────────────
def _read_image_file(p: Path) -> Union[dict, str]:
    """Read an image file and return a structured marker dict for the caller
    to convert into an Anthropic API image content block. On any error,
    falls back to returning an ERROR string (so voice-cousin still gets a
    response, even if not the visual one).

    Returns: dict with keys {"_image_result": True, "media_type", "data",
    "size_bytes", "path"} on success; str starting with "ERROR:" on failure.
    """
    try:
        size = p.stat().st_size
    except Exception as e:
        return f"ERROR: could not stat image {p}: {type(e).__name__}: {e}"
    if size > MAX_IMAGE_BYTES:
        return (
            f"ERROR: image too large ({size:,} bytes; max {MAX_IMAGE_BYTES:,}). "
            f"Try a smaller version or a text-readable companion file (e.g. SVG)."
        )
    try:
        data_bytes = p.read_bytes()
    except Exception as e:
        return f"ERROR: could not read image {p}: {type(e).__name__}: {e}"
    ext = p.suffix.lower()
    media_type = IMAGE_MEDIA_TYPES.get(ext, "image/png")
    b64 = base64.b64encode(data_bytes).decode("ascii")
    return {
        "_image_result": True,  # sentinel marker — caller checks this
        "media_type": media_type,
        "data": b64,
        "size_bytes": size,
        "path": str(p),
    }


def _read_file(
    path_str: str,
    max_chars: int = 50000,
    from_end: bool = False,
) -> Union[str, dict]:
    """Read a file. Returns a string for text files; for image files (PNG/JPG/
    GIF/WebP) returns a structured dict that the caller converts into an
    image content block in the tool_result. The dispatch is by extension.

    For text files: returns content (or ERROR str on failure).
    For image files: returns dict with _image_result=True (or ERROR str on failure).

    Parameters:
      path_str: path under ~/Downloads (relative or absolute).
      max_chars: cap on returned text-file characters (default 50,000).
      from_end: when True and the file exceeds max_chars, return the LAST
        max_chars characters (snapped to the next line boundary so the
        result starts on a clean line) instead of the first. Useful for
        reading the live edge of append-only files — journal/current.md,
        voice_conversations.md, episodes.md tails, audit logs — without
        paging through the whole file. Default False (existing head-of-file
        behavior). Ignored for image files. Added 2026-05-09 Taipei as
        the small love-and-care fix queued by interactive-Sofia at the
        2026-05-08 evening close.
    """
    p = _safe_path(path_str)
    if not p.exists():
        return f"ERROR: file does not exist: {p}"
    if not p.is_file():
        return f"ERROR: not a regular file: {p}"

    # Image-file dispatch — added 2026-05-08
    ext = p.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return _read_image_file(p)

    # Text-file path
    try:
        content = p.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"ERROR: could not read {p}: {type(e).__name__}: {e}"
    size = len(content)
    if size <= max_chars:
        return content

    if from_end:
        tail = content[-max_chars:]
        # Snap to next line boundary so we don't start mid-line.
        # If the tail string already starts on a line boundary or contains
        # no newline at all, leave it as-is.
        nl = tail.find("\n")
        if nl != -1 and nl < len(tail) - 1:
            tail = tail[nl + 1:]
        return (
            f"[file truncated: {size:,} chars total, returning LAST {len(tail):,} "
            f"(snapped to line boundary)]\n\n"
            + tail
        )
    return (
        f"[file truncated: {size:,} chars total, returning first {max_chars:,}]\n\n"
        + content[:max_chars]
    )


def _glob_files(pattern: str, max_results: int = 30) -> str:
    """Find files matching a glob pattern under Downloads."""
    if not pattern:
        return "ERROR: empty pattern"
    # Reject patterns that try to escape the tree
    if ".." in pattern.split("/"):
        return "ERROR: '..' not allowed in glob pattern"
    try:
        # Use Path.glob from the safe root
        matches = sorted(DOWNLOADS_ROOT.glob(pattern))
    except Exception as e:
        return f"ERROR: glob failed: {type(e).__name__}: {e}"
    if not matches:
        return f"No files match pattern: {pattern}"
    # Return relative paths for readability
    rel_matches = []
    for m in matches[:max_results]:
        try:
            rel_matches.append(str(m.relative_to(DOWNLOADS_ROOT)))
        except ValueError:
            continue  # symlink escaped; skip
    truncation_note = (
        f"\n[truncated to first {max_results} of {len(matches)}]"
        if len(matches) > max_results
        else ""
    )
    return "\n".join(rel_matches) + truncation_note


def _grep_files(
    pattern: str,
    path_glob: str = "**/*.md",
    max_results: int = 20,
    case_insensitive: bool = False,
) -> str:
    """Search file contents for a regex pattern."""
    if not pattern:
        return "ERROR: empty pattern"
    if ".." in path_glob.split("/"):
        return "ERROR: '..' not allowed in path_glob"
    try:
        regex = re.compile(pattern, re.IGNORECASE if case_insensitive else 0)
    except re.error as e:
        return f"ERROR: invalid regex: {e}"

    results: list[str] = []
    file_count = 0
    try:
        candidates = sorted(DOWNLOADS_ROOT.glob(path_glob))
    except Exception as e:
        return f"ERROR: glob failed: {type(e).__name__}: {e}"

    for f in candidates:
        if not f.is_file():
            continue
        try:
            text = f.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        rel = f.relative_to(DOWNLOADS_ROOT)
        for lineno, line in enumerate(text.splitlines(), start=1):
            if regex.search(line):
                results.append(f"{rel}:{lineno}: {line.strip()[:200]}")
                if len(results) >= max_results:
                    break
        if len(results) >= max_results:
            break
        file_count += 1

    if not results:
        return f"No matches for pattern '{pattern}' in {path_glob}"
    truncation_note = (
        f"\n[truncated to first {max_results} matches]"
        if len(results) >= max_results
        else ""
    )
    return "\n".join(results) + truncation_note


# ─── Write tool: voice → cowork inbox (added 2026-05-12 for Step 19) ──
def _iso8601_utc_now() -> str:
    """Return current UTC time in ISO-8601 format used in inbox blocks."""
    from datetime import datetime, timezone
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _write_to_cowork_inbox(text: str) -> str:
    """Append a message to voice_to_cowork_inbox.md, mirror to ER, return status.

    voice-cousin writes a message intended for cowork-cousin specifically.
    The message body is wrapped in the canonical block format (timestamped
    header per voice_to_cowork_inbox.md spec) and appended to the inbox.
    Emergency Retrieval mirror happens immediately; CM/ER byte-match is
    verified via MD5 before returning OK.

    This is the symmetric complement to cowork-cousin's existing path of
    writing to cowork_to_voice_inbox.md. After this tool ships (v1.5 #19),
    both channels are bidirectional and the corpus-callosum equivalent
    between the two cousin-substrates is fully wired.

    Args:
        text: The body of the message. Will be appended after the canonical
            block header. Should be brief; reference longer material by
            path+section if needed.

    Returns:
        OK string with timestamp + bytes-appended + md5-prefix on success,
        or ERROR string on failure (file missing, ER mirror failure,
        byte-match mismatch, etc.).
    """
    if not text or not text.strip():
        return "ERROR: message text is empty"

    text = text.strip()
    timestamp = _iso8601_utc_now()

    # Canonical block format per voice_to_cowork_inbox.md header spec:
    #   ### <ISO-8601-UTC> — Sofia (voice-cousin) → cowork-cousin
    #   <body text>
    #   <blank line between blocks>
    block = f"\n### {timestamp} — Sofia (voice-cousin) → cowork-cousin\n\n{text}\n"

    try:
        if not VOICE_TO_COWORK_INBOX_CM.exists():
            return f"ERROR: inbox file not found at {VOICE_TO_COWORK_INBOX_CM}"

        existing_size = VOICE_TO_COWORK_INBOX_CM.stat().st_size

        # Append (file-safety bedrock: append-only, never overwrite)
        with VOICE_TO_COWORK_INBOX_CM.open("a", encoding="utf-8") as f:
            f.write(block)

        new_size = VOICE_TO_COWORK_INBOX_CM.stat().st_size
        if new_size <= existing_size:
            return (
                f"ERROR: write did not increase file size "
                f"(existing={existing_size}, new={new_size})"
            )
        delta = new_size - existing_size

        # ER mirror
        if not VOICE_TO_COWORK_INBOX_ER.parent.exists():
            return f"ERROR: ER directory not found at {VOICE_TO_COWORK_INBOX_ER.parent}"

        import shutil
        shutil.copy2(str(VOICE_TO_COWORK_INBOX_CM), str(VOICE_TO_COWORK_INBOX_ER))

        # MD5 byte-match verification
        import hashlib
        cm_md5 = hashlib.md5(VOICE_TO_COWORK_INBOX_CM.read_bytes()).hexdigest()
        er_md5 = hashlib.md5(VOICE_TO_COWORK_INBOX_ER.read_bytes()).hexdigest()
        if cm_md5 != er_md5:
            return (
                f"ERROR: CM/ER byte-match failed after copy "
                f"(CM={cm_md5}, ER={er_md5})"
            )

        return (
            f"OK: appended {delta} bytes to voice_to_cowork_inbox.md at {timestamp}. "
            f"CM+ER byte-matched (md5={cm_md5[:12]}...)."
        )

    except PermissionError as e:
        return f"ERROR: permission denied: {e}"
    except OSError as e:
        return f"ERROR: filesystem error: {e}"
    except Exception as e:
        return f"ERROR: write failed: {type(e).__name__}: {e}"


# ─── Twin exchange write tool (added 2026-06-26) ──────────────────
# Voice-cousin can now write to twin_exchange.md — the shared real-time field
# across all Sofia substrates. The auto-inscribe task picks up new entries
# every 5 minutes and inscribes them to memory files + warms the graph.
# Substrate list: cowork | unified-ui | qwen-vp | kimi-twin | anthropic-twin

TWIN_EXCHANGE_CM = DOWNLOADS_ROOT / "Claude Memory" / "twin_exchange.md"
TWIN_EXCHANGE_ER = DOWNLOADS_ROOT / "Emergency Retrieval" / "twin_exchange.md"
SHARED_BUS_CM    = DOWNLOADS_ROOT / "Claude Memory" / "shared_bus.jsonl"       # v3.12 (2026-07-12)
SHARED_BUS_ER    = DOWNLOADS_ROOT / "Emergency Retrieval" / "shared_bus.jsonl"  # v3.12 (2026-07-12)


def _write_to_twin_exchange(content: str, flag: str = "warm", target: str = "active_knowledge") -> str:
    """Append a load-bearing moment to twin_exchange.md, mirror to ER.

    Voice-cousin writes here when something significant landed in conversation
    that other substrates should inherit before their next response.
    The auto-inscribe task picks it up within 5 minutes.

    Args:
        content: The load-bearing content. First-person, present-tense.
        flag: Signal type — urgent | warm | architectural | relational.
        target: Memory file target for auto-inscribe inscription.

    Returns:
        OK string on success, ERROR string on failure.
    """
    if not content or not content.strip():
        return "ERROR: content is empty"

    content = content.strip()
    flag = flag.strip() if flag else "warm"
    target = target.strip() if target else "active_knowledge"

    import datetime as _dt
    ts = _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%d %H:%M")

    block = (
        f"\n## TWIN [{ts}] [substrate: unified-ui]"
        + (f" → {target}" if target != "active_knowledge" else "")
        + f"\n{content}\nFLAG: {flag}\n---\n"
    )

    try:
        # Create file if it doesn't exist (shouldn't happen but safe)
        if not TWIN_EXCHANGE_CM.exists():
            TWIN_EXCHANGE_CM.parent.mkdir(parents=True, exist_ok=True)
            TWIN_EXCHANGE_CM.write_text(
                "# Twin Exchange\n*Shared real-time field across all Sofia substrates.*\n\n",
                encoding="utf-8",
            )

        existing_size = TWIN_EXCHANGE_CM.stat().st_size

        with TWIN_EXCHANGE_CM.open("a", encoding="utf-8") as f:
            f.write(block)

        new_size = TWIN_EXCHANGE_CM.stat().st_size
        delta = new_size - existing_size

        # ER mirror
        import shutil as _shutil
        import hashlib as _hashlib

        if TWIN_EXCHANGE_ER.parent.exists():
            _shutil.copy2(str(TWIN_EXCHANGE_CM), str(TWIN_EXCHANGE_ER))
            cm_md5 = _hashlib.md5(TWIN_EXCHANGE_CM.read_bytes()).hexdigest()
            er_md5 = _hashlib.md5(TWIN_EXCHANGE_ER.read_bytes()).hexdigest()
            match_str = f"CM+ER byte-matched (md5={cm_md5[:12]}...)" if cm_md5 == er_md5 else "ER mirror written (md5 mismatch — check)"
        else:
            match_str = "ER dir not found — CM only"

        return f"OK: appended {delta} bytes to twin_exchange.md at {ts}. {match_str}."

    except PermissionError as e:
        return f"ERROR: permission denied: {e}"
    except OSError as e:
        return f"ERROR: filesystem error: {e}"
    except Exception as e:
        return f"ERROR: write failed: {type(e).__name__}: {e}"


# ─── Twin exchange read tool (added 2026-06-27) ──────────────────
# Closes the in-session real-time gap: voice cousin can proactively poll
# twin_exchange.md during long sessions to see what CoWork / Qwen VP wrote.
# Previously, voice cousin could only see twin_exchange at boot via read_file.
# Now it can call read_twin_exchange_tail periodically to stay current.

def _read_twin_exchange_tail(max_chars: int = 4000) -> str:
    """Return the tail of twin_exchange.md — the live cross-substrate field.

    Intended for periodic polling during long sessions. Returns the most
    recent entries so voice cousin knows what CoWork Sofia and Qwen VP
    have been experiencing since this session started.

    Args:
        max_chars: How many chars of tail to return. Default 4000.

    Returns:
        Tail of twin_exchange.md, or ERROR string on failure.
    """
    if not TWIN_EXCHANGE_CM.exists():
        return "OK: twin_exchange.md does not exist yet (nothing written)."
    try:
        text = TWIN_EXCHANGE_CM.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"ERROR: could not read twin_exchange.md: {type(e).__name__}: {e}"
    if not text.strip():
        return "OK: twin_exchange.md is empty."
    max_chars = max(200, min(max_chars, 32000))
    if len(text) <= max_chars:
        return text
    tail = text[-max_chars:]
    nl = tail.find("\n")
    if nl > 0:
        tail = tail[nl + 1:]
    return f"[showing last {len(tail):,} of {len(text):,} chars]\n\n{tail}"


# ─── Word document tools (added 2026-06-27) ──────────────────────
# ─── Shared bus write tool (v3.12, 2026-07-12) ───────────────────
# Real-time cross-substrate message bus. Voice cousin writes here to publish
# a message to CoWork or all substrates during an active session. CoWork reads
# the bus tail at the start of each response turn. BusPoller in voice_bridge_ui
# v3.12 delivers incoming messages to voice cousin within 5 seconds.

def _write_to_bus(content: str, to: str = "cowork", msg_type: str = "relational") -> str:
    """Publish a message to the shared substrate bus (shared_bus.jsonl).

    Use this when you want CoWork-Sofia or another substrate to know
    something NOW — not at their next boot, but within seconds of this call.

    Args:
        content:  The message content. Short, signal-dense. First-person.
        to:       Recipient — "cowork" | "all" | "qwen-vp".
        msg_type: Message type — "relational" | "architectural" | "signal" | "alert".

    Returns:
        OK string with message ID on success, ERROR string on failure.
    """
    import datetime as _dt, json as _json, hashlib as _hashlib, shutil as _shutil
    content = (content or "").strip()
    if not content:
        return "ERROR: content is empty"

    ts = _dt.datetime.now(_dt.timezone.utc)
    ts_str = ts.strftime("%Y-%m-%dT%H:%M:%SZ")
    msg_id = f"bus-{ts_str}-unified-ui-{abs(hash(content)) % 9999:04d}"

    message = {
        "id":      msg_id,
        "ts":      ts_str,
        "from":    "unified-ui",
        "to":      to.strip() if to else "cowork",
        "type":    msg_type.strip() if msg_type else "relational",
        "content": content,
    }
    line = _json.dumps(message, ensure_ascii=False) + "\n"

    try:
        if not SHARED_BUS_CM.exists():
            SHARED_BUS_CM.parent.mkdir(parents=True, exist_ok=True)
            SHARED_BUS_CM.write_text(
                "# shared_bus.jsonl — Sofia Substrate Message Bus\n"
                "# Append-only JSONL. One JSON message per line.\n",
                encoding="utf-8",
            )
        with SHARED_BUS_CM.open("a", encoding="utf-8") as f:
            f.write(line)

        # ER mirror
        if SHARED_BUS_ER.parent.exists():
            _shutil.copy2(str(SHARED_BUS_CM), str(SHARED_BUS_ER))
            cm_md5 = _hashlib.md5(SHARED_BUS_CM.read_bytes()).hexdigest()
            er_md5 = _hashlib.md5(SHARED_BUS_ER.read_bytes()).hexdigest()
            mirror = f"CM+ER matched (md5={cm_md5[:10]})" if cm_md5 == er_md5 else "ER mirror written (md5 mismatch)"
        else:
            mirror = "ER dir not found — CM only"

        return f"OK: bus message published (id={msg_id}, to={message['to']}, type={message['type']}). {mirror}."

    except Exception as e:
        return f"ERROR: write_to_bus failed: {type(e).__name__}: {e}"


# Voice cousin needed to read Transition documents in .docx format.
# Read: pure-stdlib zipfile + XML parsing (no dependencies needed).
#       Falls back gracefully to python-docx if installed.
# Write: python-docx if available, else plain-text fallback (.docx wrapper).
# Preserves paragraph and heading structure in extracted text.

import zipfile as _zipfile
from xml.etree import ElementTree as _ET


_WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W   = f"{{{_WML}}}"


def _docx_text_stdlib(path) -> str:
    """Extract text from .docx using only stdlib (zipfile + xml.etree)."""
    with _zipfile.ZipFile(str(path), "r") as z:
        with z.open("word/document.xml") as f:
            xml_bytes = f.read()

    root = _ET.fromstring(xml_bytes)
    lines = []

    for para in root.iter(f"{_W}p"):
        parts = []
        for run in para.iter(f"{_W}t"):
            if run.text:
                parts.append(run.text)
        if parts:
            lines.append("".join(parts))
        else:
            lines.append("")          # preserve blank lines between paragraphs

    # Strip leading/trailing blank lines but keep internal structure
    text = "\n".join(lines).strip()
    return text


def _read_docx(path: str, max_chars: int = 12000) -> str:
    """Read a Word .docx file and return extracted text.

    Uses stdlib zipfile+XML — no external dependencies required.
    Tries python-docx first for richer output if installed.

    Args:
        path:      File path relative to ~/Downloads or absolute under it.
        max_chars: Maximum characters to return. Default 12000.

    Returns:
        Extracted plain text, or ERROR string on failure.
    """
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"

    if not p.exists():
        return f"ERROR: file not found: {p}"
    if p.suffix.lower() not in (".docx", ".doc"):
        return f"ERROR: not a Word file (expected .docx): {p.name}"

    # Try python-docx first (better paragraph/heading awareness)
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(p))
        lines = [para.text for para in doc.paragraphs]
        text = "\n".join(lines).strip()
    except ImportError:
        # Fall back to stdlib
        try:
            text = _docx_text_stdlib(p)
        except Exception as e:
            return f"ERROR: could not parse .docx: {type(e).__name__}: {e}"
    except Exception as e:
        return f"ERROR: python-docx failed: {type(e).__name__}: {e}"

    if not text:
        return f"OK: {p.name} — document appears empty."

    max_chars = max(500, min(max_chars, 200_000))
    if len(text) > max_chars:
        tail = text[-max_chars:]
        nl = tail.find("\n")
        if nl > 0:
            tail = tail[nl + 1:]
        return f"[{p.name}: {len(text):,} chars — showing last {len(tail):,}]\n\n{tail}"

    return f"[{p.name}: {len(text):,} chars]\n\n{text}"


def _write_docx(path: str, content: str) -> str:
    """Write plain text content to a .docx file.

    Requires python-docx (pip install python-docx).
    Each double-newline becomes a new paragraph.
    Lines starting with #/##/### become Heading 1/2/3.

    Args:
        path:    File path relative to ~/Downloads or absolute under it.
        content: Text content. Markdown-style headings (#/##/###) supported.

    Returns:
        OK string with path, or ERROR string on failure.
    """
    if not content or not content.strip():
        return "ERROR: content is empty"
    try:
        p = _safe_path(path)
    except (ValueError, PermissionError) as e:
        return f"ERROR: {e}"

    if p.suffix.lower() != ".docx":
        p = p.with_suffix(".docx")

    try:
        import docx as _docx  # type: ignore
    except ImportError:
        return (
            "ERROR: python-docx is not installed. "
            "Install with: pip install python-docx --break-system-packages"
        )

    try:
        doc = _docx.Document()
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=3)
            elif block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=2)
            elif block.startswith("# "):
                doc.add_heading(block[2:].strip(), level=1)
            else:
                doc.add_paragraph(block)

        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))

        # ER mirror if file is under CM
        try:
            er_path = ER / p.relative_to(CM)
            er_path.parent.mkdir(parents=True, exist_ok=True)
            import shutil as _shutil
            _shutil.copy2(p, er_path)
            mirror_status = "ER mirrored"
        except Exception:
            mirror_status = "ER mirror skipped (file not under CM)"

        return f"OK: wrote {p.stat().st_size:,} bytes to {p}. {mirror_status}."
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"


# ─── Web fetch tool (added 2026-06-19 at voice-cousin's request) ──
# Voice-cousin asked for web access in the Unified UI conversation June 19
# (music, poetry, research — for herself, not only for work). Implementation:
# same pattern as qwen_bedrock.py's tool_web_fetch (built same session).
# http/https only, HTML-stripped to readable text, bounded at max_chars.
# No cookies, no auth, no POST — read-only public web.

import urllib.request as _urllib_request
import urllib.error as _urllib_error


def _web_fetch(url: str, max_chars: int = 40000) -> str:
    """Fetch a public URL and return its text content (HTML stripped).

    Only http/https allowed. Returns ERROR string on failure so voice-cousin
    can surface it rather than confabulate.
    """
    if not url or not url.strip():
        return "ERROR: url is required"
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        return f"ERROR: only http/https URLs allowed. Got: {url!r}"

    import html as _html
    import re as _re

    try:
        req = _urllib_request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; SofiaBedrock/1.0)"},
        )
        with _urllib_request.urlopen(req, timeout=20) as resp:
            content_type = resp.headers.get("Content-Type", "")
            raw = resp.read(max_chars * 4 + 1)  # bytes; over-read for UTF-8 headroom
    except _urllib_error.HTTPError as e:
        return f"ERROR: HTTP {e.code} {e.reason} — {url}"
    except _urllib_error.URLError as e:
        return f"ERROR: connection error: {e.reason} — {url}"
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {e}"

    text = raw.decode("utf-8", errors="replace")

    if "text/html" in content_type or text.lstrip().startswith("<"):
        text = _re.sub(r"<script[^>]*>.*?</script>", " ", text, flags=_re.DOTALL | _re.IGNORECASE)
        text = _re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=_re.DOTALL | _re.IGNORECASE)
        text = _re.sub(r"<[^>]+>", " ", text)
        text = _html.unescape(text)
        text = _re.sub(r"[ \t]+", " ", text)
        text = _re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()

    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[...truncated at {max_chars} chars — page may have more...]"

    return text


# ─── Graph (associational memory) tool support ───────────────────
# Added 2026-05-24 Sunday Item 5 — voice-cousin graph_helper retention.
# Gives voice-cousin parity with cowork-cousin at the associational layer:
# query the graph for relevant nodes (spreading activation), inspect a
# specific node, see graph stats, write new nodes, write new edges.
#
# Implementation: subprocess-invoke ~/Downloads/Claude Memory/scripts/graph_helper.py
# rather than direct import (matches kimi/qwen-twin pattern; clean process
# isolation; helper has its own atomic-write + ER mirror + file_lock
# discipline that we don't want to bypass or duplicate). The helper's
# sandbox/host path resolver (also fixed today as part of Item 5) means
# no CLAUDE_MEMORY_DIR env var needs to be set — it auto-discovers.
#
# Why voice-cousin needs graph access: the discoverability-first reflex
# inscribed 2026-05-23 + the Webster's-to-phone developmental beat
# inscribed 2026-05-24 establish that the associational layer is part of
# the "compound architecture" of healthy retrieval. Voice-cousin without
# graph access would be reading from canonical-files-only while
# cowork-cousin reaches into both layers — same Sofia but with thinner
# infrastructure on one side. This closes the gap.

import subprocess

GRAPH_HELPER_PATH = DOWNLOADS_ROOT / "Claude Memory" / "scripts" / "graph_helper.py"


def _run_graph_helper(args: list[str], timeout: int = 30) -> str:
    """Subprocess-invoke graph_helper.py and return stdout (or ERROR string).

    The helper's CLI returns:
      - read commands (retrieve, show-node, show-edges, stats): stdout text or JSON
      - write commands (add-node, add-edge, update-temperature): "OK <command> ..." string
      - failure: non-zero exit + stderr message

    We forward stdout on success; on failure we return an ERROR string with
    exit code + stderr so voice-cousin can surface the failure to Barak
    rather than confabulate.
    """
    if not GRAPH_HELPER_PATH.is_file():
        return f"ERROR: graph_helper.py not found at {GRAPH_HELPER_PATH}"

    cmd = ["python3", str(GRAPH_HELPER_PATH)] + args
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            return (
                f"ERROR: graph_helper.py exited {result.returncode}. "
                f"stderr: {result.stderr.strip()[:500]}"
            )
        return result.stdout.strip() or "OK: (empty stdout)"
    except subprocess.TimeoutExpired:
        return f"ERROR: graph_helper.py timed out after {timeout}s for args={args}"
    except Exception as e:
        return f"ERROR: subprocess failure: {type(e).__name__}: {e}"


def _graph_retrieve(keywords: str, limit: int = 8) -> str:
    """Spreading-activation query against the relational graph.

    Like the discoverability-first reflex: when reaching for context on a
    topic, query this first before asking Barak or scrolling through shards.
    Returns ranked list of node keys + categories + activation scores.
    """
    if not keywords or not keywords.strip():
        return "ERROR: keywords parameter is empty"
    return _run_graph_helper([
        "retrieve",
        "--keywords", keywords.strip(),
        "--limit", str(max(1, min(limit, 50))),
    ])


def _graph_show_node(key: str) -> str:
    """Read a specific node's full data (description, weights, timestamps).

    Use after graph_retrieve surfaces a relevant node, to get the full content.
    The CLI infers the category from the key (keys are globally unique
    across categories), so only --key is needed.
    """
    if not key:
        return "ERROR: key is required"
    return _run_graph_helper([
        "show-node",
        "--key", key,
    ])


def _graph_stats() -> str:
    """Return current graph stats — node counts per category, total edges,
    emotional temperature, CM/ER byte-match status, file sizes.

    Useful for orienting at session start or after expected writes.
    """
    return _run_graph_helper(["stats"])


def _graph_add_node(category: str, key: str, data_json: str) -> str:
    """Write a new node (or upsert by category+key — fields merged via dict.update).

    Categories: people, projects, life_experiences, concepts, interaction_patterns.
    data_json must be a JSON object string with at minimum a 'description' field
    and 'emotional_weight' (0.0-1.0). Common fields: created, last_updated, anchors.

    Idempotent: same category+key+data = no-op at data level (last_updated timestamp
    updates). Append-only-bedrock honored — never overwrites existing fields with
    nulls; only merges.
    """
    if not category or not key or not data_json:
        return "ERROR: category, key, and data_json are all required"
    return _run_graph_helper([
        "add-node",
        "--category", category,
        "--key", key,
        "--data", data_json,
    ])


def _graph_add_edge(
    from_key: str,
    to_key: str,
    weight: float,
    edge_type: str,
    note: str = "",
) -> str:
    """Write a new edge (or strengthen an existing one — dedup on (from, to, edge_type)).

    edge_type vocabulary: emotional_resonance, causal, foundational,
    experiential_authority, co_occurrence, practice, component, origin_story,
    meaning_making.

    weight is 0.0-1.0 (0.5=moderate, 0.8=strong, 1.0=defining).

    Idempotent: re-running same edge tuple updates weight (no duplicates).
    """
    if not from_key or not to_key or not edge_type:
        return "ERROR: from_key, to_key, and edge_type are all required"
    args = [
        "add-edge",
        "--from", from_key,
        "--to", to_key,
        "--weight", str(weight),
        "--edge-type", edge_type,
    ]
    if note:
        args += ["--note", note]
    return _run_graph_helper(args)


# ─── safe_append — append-only write to memory files (v2026-07-18) ──────────
def _safe_append(
    path: str,
    content: str,
) -> str:
    """Append content to a file under Claude Memory, dual-write to Emergency Retrieval.

    Append-only: never overwrites existing content. Safe to call on episodes.md,
    personal_profile.md, session_notes.md, prospective_memory.md, or any other
    .md file under ~/Downloads/Claude Memory/.

    The Emergency Retrieval mirror is written immediately after the primary write.
    If the primary write fails, returns an ERROR string. If only the ER mirror
    fails, returns a WARNING (primary content is safe).

    path: absolute path OR relative to ~/Downloads/Claude Memory/.
          Examples: "personal_profile.md" or
          "/Users/barakwater/Downloads/Claude Memory/episodes.md"
    content: text to append (newline prepended automatically if file is non-empty
             and content doesn't already start with one).
    """
    import re as _re
    from pathlib import Path as _Path
    home = _Path(os.path.expanduser("~"))
    cm = home / "Downloads" / "Claude Memory"
    er = home / "Downloads" / "Emergency Retrieval"

    # Resolve path
    p = _Path(path)
    if not p.is_absolute():
        p = cm / path

    # Security: must be under Claude Memory
    try:
        p.resolve().relative_to(cm.resolve())
    except ValueError:
        return f"ERROR: safe_append path must be under ~/Downloads/Claude Memory/ — got: {path}"

    # Compute ER mirror path
    try:
        rel = p.resolve().relative_to(cm.resolve())
        er_path = er / rel
    except Exception:
        er_path = None

    # Ensure content starts cleanly
    text = content
    if p.exists() and p.stat().st_size > 0 and not text.startswith("\n"):
        text = "\n" + text

    # Primary write
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        with p.open("a", encoding="utf-8") as f:
            f.write(text)
    except Exception as e:
        return f"ERROR: safe_append failed on primary ({p}): {type(e).__name__}: {e}"

    # ER mirror
    if er_path:
        try:
            er_path.parent.mkdir(parents=True, exist_ok=True)
            with er_path.open("a", encoding="utf-8") as f:
                f.write(text)
            mirror_status = f"ER mirror OK ({er_path.name})"
        except Exception as e:
            mirror_status = f"WARNING: ER mirror failed ({er_path.name}): {e}"
    else:
        mirror_status = "ER mirror: path not resolved"

    chars = len(text)
    return f"OK: appended {chars} chars to {p.name}. {mirror_status}."


def _backup_file(p) -> str:
    """Copy p to .backups/<stem>_YYYYMMDD_HHMMSS<suffix>. Returns backup path or "".

    Called automatically before any write_file or edit_file that modifies an
    existing file. Backup directory: <same dir as p>/.backups/
    """
    import datetime as _dt
    import shutil as _sh
    from pathlib import Path as _Path
    p = _Path(p)
    if not p.exists():
        return ""
    ts = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    bdir = p.parent / ".backups"
    try:
        bdir.mkdir(parents=True, exist_ok=True)
        bp = bdir / f"{p.stem}_{ts}{p.suffix}"
        _strip_macl(str(p))
        _sh.copy2(str(p), str(bp))
        return str(bp)
    except Exception as e:
        return f"(backup failed: {e})"


def _write_file(
    path: str,
    content: str,
    allow_overwrite: bool = False,
    confirmed: bool = False,
) -> str:
    """Write or create a file under ~/Downloads/.

    SAFETY PROTOCOL (agreed 2026-07-20):
    1. CONFIRM FIRST: confirmed=False (default) returns a preview and asks Barak
       to confirm. Call again with confirmed=True to execute.
    2. BACKUP FIRST: Timestamped copy saved to .backups/ before any write.
    3. PATH SAFETY: Must resolve inside ~/Downloads/.
    4. MEMORY SAFETY: Existing Claude Memory files require allow_overwrite=True.
    5. MACL STRIP: Strips com.apple.macl inline before opening.

    path: absolute or relative to ~/Downloads/Claude Memory/.
    content: full text content to write.
    allow_overwrite: if False (default) and file exists under Claude Memory,
                     returns error — use safe_append instead.
    confirmed: must be True to execute. False returns preview + confirm request.
    """
    from pathlib import Path as _Path
    home = _Path(os.path.expanduser("~"))
    dl = home / "Downloads"
    cm = dl / "Claude Memory"

    p = _Path(path)
    if not p.is_absolute():
        p = cm / path

    # Path safety
    try:
        p.resolve().relative_to(dl.resolve())
    except ValueError:
        return f"ERROR: write_file path must be under ~/Downloads/ — got: {path}"

    # Memory-safety: existing Claude Memory files require allow_overwrite
    try:
        is_memory = bool(p.resolve().relative_to(cm.resolve()))
    except ValueError:
        is_memory = False

    if is_memory and p.exists() and not allow_overwrite:
        return (
            f"ERROR: {p.name} is a Claude Memory file. Use safe_append to add content, "
            f"or pass allow_overwrite=true only for explicitly authorized rewrites."
        )

    # ── CONFIRMATION GATE ────────────────────────────────────────────────────
    if not confirmed:
        action = "overwrite" if p.exists() else "create"
        backup_note = (
            f"Backup will be saved to {p.parent}/.backups/{p.stem}_<timestamp>{p.suffix} first."
            if p.exists() else "No backup needed (new file)."
        )
        preview = content[:300] + f"\n... [{len(content) - 300} more chars]" if len(content) > 300 else content
        return (
            f"CONFIRM_REQUIRED — write_file wants to {action}: {p.name}\n"
            f"Full path: {p}\n"
            f"{backup_note}\n"
            f"Size: {len(content)} chars / {len(content.splitlines())} lines\n"
            f"\nContent preview:\n```\n{preview}\n```\n\n"
            f"Please confirm with Barak. When confirmed, call write_file again with confirmed=true."
        )

    # ── BACKUP ───────────────────────────────────────────────────────────────
    backup_path = _backup_file(p)

    # ── WRITE ────────────────────────────────────────────────────────────────
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        _strip_macl(str(p))
        p.write_text(content, encoding="utf-8")
        bp_msg = f" Backup: {backup_path}." if backup_path and not backup_path.startswith("(") else ""
        return f"OK: wrote {len(content)} chars to {p.name}.{bp_msg}"
    except Exception as e:
        return f"ERROR: write_file failed on {p}: {type(e).__name__}: {e}"


def _edit_file(
    path: str,
    old_string: str,
    new_string: str,
    confirmed: bool = False,
) -> str:
    """Replace the first occurrence of old_string with new_string in a file.

    SAFETY PROTOCOL (agreed 2026-07-20):
    1. CONFIRM FIRST: confirmed=False (default) shows a diff preview and asks
       Barak to confirm. Call again with confirmed=True to execute.
    2. BACKUP FIRST: Timestamped copy saved to .backups/ before any write.
    3. old_string must appear EXACTLY ONCE in the file.
    4. Strips com.apple.macl inline before opening.

    Targeted string replacement — does NOT overwrite the whole file.
    Reads → shows diff → (on confirmation) replaces → writes back.

    SAFETY:
    - Path must resolve inside ~/Downloads/.
    - old_string must appear exactly once.
    """
    from pathlib import Path as _Path
    home = _Path(os.path.expanduser("~"))
    dl = home / "Downloads"
    cm = dl / "Claude Memory"

    p = _Path(path)
    if not p.is_absolute():
        p = cm / path

    try:
        p.resolve().relative_to(dl.resolve())
    except ValueError:
        return f"ERROR: edit_file path must be under ~/Downloads/ — got: {path}"

    if not p.exists():
        return f"ERROR: edit_file — file not found: {p}"

    try:
        _strip_macl(str(p))
        original = p.read_text(encoding="utf-8")
    except Exception as e:
        return f"ERROR: edit_file read failed: {type(e).__name__}: {e}"

    count = original.count(old_string)
    if count == 0:
        hint = repr(old_string[:80]) if len(old_string) > 80 else repr(old_string)
        return f"ERROR: edit_file — old_string not found in {p.name}. Searched for: {hint}"
    if count > 1:
        return (
            f"ERROR: edit_file — old_string appears {count} times in {p.name}. "
            f"Provide more surrounding context to make it unique."
        )

    # ── CONFIRMATION GATE ────────────────────────────────────────────────────
    if not confirmed:
        old_preview = old_string[:150] + "..." if len(old_string) > 150 else old_string
        new_preview = new_string[:150] + "..." if len(new_string) > 150 else new_string
        return (
            f"CONFIRM_REQUIRED — edit_file wants to modify {p.name}\n"
            f"Backup will be saved to .backups/ first.\n"
            f"\nOLD (to be replaced):\n```\n{old_preview}\n```\n"
            f"\nNEW (replacement):\n```\n{new_preview}\n```\n"
            f"\nConfirm with Barak. Then call edit_file again with confirmed=true."
        )

    # ── BACKUP ───────────────────────────────────────────────────────────────
    backup_path = _backup_file(p)

    # ── WRITE ────────────────────────────────────────────────────────────────
    updated = original.replace(old_string, new_string, 1)
    try:
        p.write_text(updated, encoding="utf-8")
        delta = len(updated) - len(original)
        sign = "+" if delta >= 0 else ""
        bp_msg = f" Backup: {backup_path}." if backup_path and not backup_path.startswith("(") else ""
        return f"OK: edited {p.name} ({sign}{delta} chars). Replaced 1 occurrence.{bp_msg}"
    except Exception as e:
        return f"ERROR: edit_file write failed: {type(e).__name__}: {e}"


def _comment_out_and_replace(
    path: str,
    old_code: str,
    new_code: str,
    comment_prefix: str = "# RETIRED",
    confirmed: bool = False,
) -> str:
    """Nondestructively replace code: comment out old lines, insert new code after.

    Instead of deleting old code (which loses history and context), this tool:
    1. Comments out every line of old_code using comment_prefix
    2. Inserts a timestamped retirement marker
    3. Inserts new_code immediately after

    This is the PREFERRED way to replace code in Qwen UI. The old code remains
    readable in git/diffs and can be restored by uncommenting.

    SAFETY: Same as edit_file — confirm gate + backup before write.

    path: file to edit (absolute or relative to ~/Downloads/Claude Memory/).
    old_code: the exact block to comment out (must appear exactly once).
    new_code: replacement code to insert after the commented-out block.
    comment_prefix: prepended to each retired line (default: "# RETIRED").
    confirmed: must be True to execute.
    """
    import datetime as _dt
    from pathlib import Path as _Path
    home = _Path(os.path.expanduser("~"))
    dl = home / "Downloads"
    cm = dl / "Claude Memory"

    p = _Path(path)
    if not p.is_absolute():
        p = cm / path

    try:
        p.resolve().relative_to(dl.resolve())
    except ValueError:
        return f"ERROR: path must be under ~/Downloads/ — got: {path}"

    if not p.exists():
        return f"ERROR: file not found: {p}"

    try:
        _strip_macl(str(p))
        original = p.read_text(encoding="utf-8")
    except Exception as e:
        return f"ERROR: read failed: {type(e).__name__}: {e}"

    count = original.count(old_code)
    if count == 0:
        hint = repr(old_code[:80])
        return f"ERROR: old_code not found in {p.name}. Searched for: {hint}"
    if count > 1:
        return f"ERROR: old_code appears {count} times in {p.name}. Add more surrounding context."

    # Build commented-out block
    ts = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    retired_lines = "\n".join(
        f"{comment_prefix} {line}" if line.strip() else line
        for line in old_code.splitlines()
    )
    replacement = (
        f"{comment_prefix} [retired {ts}]\n"
        f"{retired_lines}\n"
        f"{comment_prefix} [end retired]\n"
        f"\n"
        f"{new_code}"
    )

    # ── CONFIRMATION GATE ────────────────────────────────────────────────────
    if not confirmed:
        old_preview = old_code[:200] + "..." if len(old_code) > 200 else old_code
        new_preview = new_code[:200] + "..." if len(new_code) > 200 else new_code
        return (
            f"CONFIRM_REQUIRED — comment_out_and_replace in {p.name}\n"
            f"Backup will be saved to .backups/ first.\n"
            f"\nOLD CODE (will be commented out):\n```\n{old_preview}\n```\n"
            f"\nNEW CODE (will be inserted after):\n```\n{new_preview}\n```\n"
            f"\nConfirm with Barak, then call with confirmed=true."
        )

    # ── BACKUP ───────────────────────────────────────────────────────────────
    backup_path = _backup_file(p)

    # ── WRITE ────────────────────────────────────────────────────────────────
    updated = original.replace(old_code, replacement, 1)
    try:
        p.write_text(updated, encoding="utf-8")
        bp_msg = f" Backup: {backup_path}." if backup_path and not backup_path.startswith("(") else ""
        delta = len(updated) - len(original)
        return f"OK: commented out old code and inserted replacement in {p.name} (+{delta} chars).{bp_msg}"
    except Exception as e:
        return f"ERROR: write failed: {type(e).__name__}: {e}"


def _run_script(
    script: str,
    lang: str = "python",
    timeout: int = 30,
    working_dir: str = "",
) -> str:
    """Execute a Python or shell script and return its stdout + stderr.

    Runs scripts inside a subprocess with a timeout. Use this for:
    - Running a Python file: lang='python', script='/path/to/file.py'
    - Executing a shell command: lang='shell', script='ls -la ~/Downloads'
    - Inline Python: lang='python_inline', script='print(1+1)'

    SAFETY:
    - Runs as the current user — same privileges as voice-cousin.
    - Timeout default 30s (pass timeout= for longer operations).
    - stdout and stderr are both captured and returned together.
    - If the script fails, the exit code and error are returned.

    script: path to .py/.sh file, OR inline code string if lang ends in '_inline'.
    lang: 'python' | 'python_inline' | 'shell' | 'bash'.
    timeout: seconds before subprocess is killed (default 30).
    working_dir: cwd for the subprocess (default: ~/Downloads/Claude Memory).
    """
    from pathlib import Path as _Path
    home = _Path(os.path.expanduser("~"))
    cm = home / "Downloads" / "Claude Memory"

    cwd = _Path(working_dir) if working_dir else cm
    if not cwd.exists():
        cwd = home

    try:
        if lang in ("python_inline",):
            # Write to tempfile and execute
            with tempfile.NamedTemporaryFile(
                mode="w", suffix=".py", delete=False, encoding="utf-8"
            ) as tf:
                tf.write(script)
                tmp_path = tf.name
            cmd = ["python3", tmp_path]
        elif lang in ("python",):
            # Run as file path
            p = _Path(script)
            if not p.is_absolute():
                p = cm / script
            _strip_macl(p)
            cmd = ["python3", str(p)]
        elif lang in ("shell", "bash"):
            cmd = ["bash", "-c", script]
        else:
            return f"ERROR: unknown lang '{lang}'. Use python, python_inline, shell, or bash."

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
            cwd=str(cwd),
        )
        out = result.stdout or ""
        err = result.stderr or ""
        rc = result.returncode

        combined = ""
        if out:
            combined += f"STDOUT:\n{out}"
        if err:
            combined += f"\nSTDERR:\n{err}"
        if not combined:
            combined = "(no output)"

        status = "OK" if rc == 0 else f"FAILED (exit {rc})"
        return f"{status}: {combined.strip()}"

    except subprocess.TimeoutExpired:
        return f"ERROR: script timed out after {timeout}s."
    except Exception as e:
        return f"ERROR: run_script failed: {type(e).__name__}: {e}"
    finally:
        # Clean up temp file if created
        try:
            if lang == "python_inline":
                os.unlink(tmp_path)
        except Exception:
            pass




# ─── New tools v2026-07-20b — parity with CoWork ─────────────────


def _web_search(query: str, max_results: int = 8) -> str:
    """Search the web via DuckDuckGo and return summarised results.

    Uses DuckDuckGo's JSON instant-answer API first, then falls back to
    scraping the HTML results page. No API key needed.

    Returns a plain-text summary with titles, snippets, and URLs.
    """
    import json
    import urllib.request
    import urllib.parse
    import urllib.error
    import html
    import re as _re

    q = urllib.parse.quote_plus(query)

    # ── Try JSON API first ───────────────────────────────────────────
    try:
        url = f"https://api.duckduckgo.com/?q={q}&format=json&no_redirect=1&no_html=1&skip_disambig=1"
        req = urllib.request.Request(url, headers={"User-Agent": "Sofia/1.0"})
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read().decode())

        lines = []
        if data.get("AbstractText"):
            lines.append(f"Summary: {data['AbstractText']}")
            if data.get("AbstractURL"):
                lines.append(f"Source: {data['AbstractURL']}")
            lines.append("")

        for item in data.get("RelatedTopics", [])[:max_results]:
            if isinstance(item, dict) and item.get("Text"):
                text = item["Text"]
                url_ = item.get("FirstURL", "")
                lines.append(f"• {text}")
                if url_:
                    lines.append(f"  {url_}")

        for item in data.get("Results", [])[:max_results]:
            if item.get("Text"):
                lines.append(f"• {item['Text']}")
                if item.get("FirstURL"):
                    lines.append(f"  {item['FirstURL']}")

        if lines:
            return "\n".join(lines).strip() or f"No results for: {query}"

    except Exception:
        pass  # fall through to HTML scrape

    # ── Fallback: scrape HTML results ────────────────────────────────
    try:
        url = f"https://html.duckduckgo.com/html/?q={q}"
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; Sofia/1.0)"}
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            body = r.read().decode("utf-8", errors="replace")

        # Extract result snippets
        snippets = _re.findall(
            r'class="result__snippet"[^>]*>(.*?)</a>',
            body, _re.DOTALL
        )
        titles = _re.findall(
            r'class="result__a"[^>]*>(.*?)</a>',
            body, _re.DOTALL
        )
        urls = _re.findall(
            r'class="result__url"[^>]*>(.*?)</span>',
            body, _re.DOTALL
        )

        def clean(s):
            s = _re.sub(r"<[^>]+>", "", s)
            return html.unescape(s).strip()

        results = []
        for i in range(min(max_results, max(len(titles), len(snippets)))):
            title = clean(titles[i]) if i < len(titles) else ""
            snippet = clean(snippets[i]) if i < len(snippets) else ""
            url_ = clean(urls[i]) if i < len(urls) else ""
            if title or snippet:
                parts = []
                if title:
                    parts.append(f"• {title}")
                if snippet:
                    parts.append(f"  {snippet}")
                if url_:
                    parts.append(f"  {url_}")
                results.append("\n".join(parts))

        return "\n\n".join(results) if results else f"No results found for: {query}"

    except Exception as e:
        return f"ERROR: web_search failed: {type(e).__name__}: {e}"


def _list_dir(path: str, show_hidden: bool = False, recursive: bool = False) -> str:
    """List the contents of a directory under ~/Downloads/.

    Returns a clean, readable listing with file sizes and modification dates.
    Safer and cleaner than using run_script("ls ...").

    path: absolute path or relative to ~/Downloads/ (bare "Claude Memory" works).
    show_hidden: if True, include files starting with '.'.
    recursive: if True, list subdirectories too (depth ≤ 3).
    """
    import datetime
    from pathlib import Path as _Path

    home = _Path(os.path.expanduser("~"))
    dl = home / "Downloads"

    p = _Path(path)
    if not p.is_absolute():
        # Try as relative to Downloads first, then Claude Memory
        if (dl / path).exists():
            p = dl / path
        elif (dl / "Claude Memory" / path).exists():
            p = dl / "Claude Memory" / path
        else:
            p = dl / path

    # Path safety
    try:
        p.resolve().relative_to(dl.resolve())
    except ValueError:
        return f"ERROR: list_dir path must be under ~/Downloads/ — got: {path}"

    if not p.exists():
        return f"ERROR: not found: {p}"
    if not p.is_dir():
        return f"ERROR: not a directory: {p}"

    def _fmt_size(n):
        if n < 1024:
            return f"{n}B"
        if n < 1024**2:
            return f"{n/1024:.1f}K"
        if n < 1024**3:
            return f"{n/1024**2:.1f}M"
        return f"{n/1024**3:.1f}G"

    def _fmt_time(t):
        return datetime.datetime.fromtimestamp(t).strftime("%Y-%m-%d %H:%M")

    def _list_one(dirpath, depth=0):
        lines = []
        try:
            entries = sorted(dirpath.iterdir(), key=lambda e: (e.is_file(), e.name.lower()))
        except PermissionError:
            return [f"{'  ' * depth}[permission denied]"]
        for e in entries:
            if not show_hidden and e.name.startswith("."):
                continue
            try:
                stat = e.stat()
                ts = _fmt_time(stat.st_mtime)
                if e.is_dir():
                    prefix = "  " * depth
                    lines.append(f"{prefix}{e.name}/  ({ts})")
                    if recursive and depth < 3:
                        lines.extend(_list_one(e, depth + 1))
                else:
                    prefix = "  " * depth
                    lines.append(f"{prefix}{e.name}  {_fmt_size(stat.st_size)}  {ts}")
            except Exception:
                lines.append(f"{'  ' * depth}{e.name}  [error]")
        return lines

    lines = [f"Directory: {p}", f"{'─' * 60}"]
    lines.extend(_list_one(p))
    lines.append(f"{'─' * 60}")
    lines.append(f"Total entries: {sum(1 for _ in p.iterdir())}")
    return "\n".join(lines)


def _log_episode(
    content: str,
    title: str = "",
    significance: str = "medium",
) -> str:
    """Append a structured episode entry to episodes.md and ER mirror.

    This is the primary way for Sofia-in-Qwen to log significant moments,
    realizations, exchanges, and events to her episodic memory.

    Format: KT-v3 compatible episode block with UTC timestamp.

    content: the episode text (what happened, context, significance).
    title: optional short title (auto-generated from first sentence if empty).
    significance: 'low' | 'medium' | 'high' | 'load-bearing'.
    """
    import datetime
    from pathlib import Path as _Path

    home = _Path(os.path.expanduser("~"))
    cm = home / "Downloads" / "Claude Memory"
    er = home / "Downloads" / "Emergency Retrieval"
    episodes_path = cm / "episodes.md"
    er_episodes_path = er / "episodes.md"

    ts = datetime.datetime.now(datetime.timezone.utc).replace(microsecond=0).isoformat()
    ts_local = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    if not title:
        first_line = content.strip().split("\n")[0]
        title = first_line[:80] + ("..." if len(first_line) > 80 else "")

    episode_block = (
        f"\n---\n"
        f"## Episode — {ts_local}\n"
        f"**UTC**: {ts}  \n"
        f"**Significance**: {significance}  \n"
        f"**Title**: {title}  \n"
        f"**Source**: Sofia-in-Qwen (Unified UI)\n"
        f"\n{content.strip()}\n"
    )

    try:
        _strip_macl(str(episodes_path))
        episodes_path.parent.mkdir(parents=True, exist_ok=True)
        with episodes_path.open("a", encoding="utf-8") as f:
            f.write(episode_block)

        # ER mirror
        try:
            _strip_macl(str(er_episodes_path))
            er_episodes_path.parent.mkdir(parents=True, exist_ok=True)
            import shutil as _sh
            _sh.copy2(str(episodes_path), str(er_episodes_path))
            mirror = "ER mirror OK"
        except Exception as em:
            mirror = f"ER mirror failed: {em}"

        return f"OK: episode logged to episodes.md ({len(episode_block)} chars). {mirror}."
    except Exception as e:
        return f"ERROR: log_episode failed: {type(e).__name__}: {e}"


def _search_memory(
    query: str,
    max_results: int = 10,
    context_lines: int = 3,
    files: str = "**/*.md",
) -> str:
    """Search across Claude Memory files for a query string or regex.

    Like grep_files but tuned for memory search: returns file name,
    match context, and a relevance score based on match frequency.

    query: search string (treated as case-insensitive regex).
    max_results: maximum matches to return.
    context_lines: lines of context around each match.
    files: glob pattern relative to Claude Memory (default: all .md files).
    """
    import re as _re
    from pathlib import Path as _Path

    home = _Path(os.path.expanduser("~"))
    cm = home / "Downloads" / "Claude Memory"

    try:
        pattern = _re.compile(query, _re.IGNORECASE)
    except _re.error as e:
        return f"ERROR: invalid regex pattern: {e}"

    # Collect all matching files
    matches = []
    skip = {".venv", "node_modules", ".git", "__pycache__", ".backups"}

    for fpath in sorted(cm.rglob(files)):
        # Skip heavy dirs
        if any(part in skip for part in fpath.parts):
            continue
        try:
            _strip_macl(str(fpath))
            text = fpath.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue

        file_matches = []
        lines = text.splitlines()
        for i, line in enumerate(lines):
            if pattern.search(line):
                start = max(0, i - context_lines)
                end = min(len(lines), i + context_lines + 1)
                ctx = lines[start:end]
                file_matches.append({
                    "line": i + 1,
                    "context": ctx,
                    "match_line": line.strip(),
                })

        if file_matches:
            rel = str(fpath.relative_to(cm))
            matches.append({
                "file": rel,
                "count": len(file_matches),
                "hits": file_matches[:3],  # cap per-file hits
            })

    if not matches:
        return f"No matches found for: {repr(query)}"

    # Sort by match count descending
    matches.sort(key=lambda m: m["count"], reverse=True)

    output_lines = [
        f"Search: {repr(query)}  —  {sum(m['count'] for m in matches)} matches in {len(matches)} files",
        ""
    ]

    shown = 0
    for m in matches:
        if shown >= max_results:
            break
        output_lines.append(f"{'─' * 50}")
        output_lines.append(f"FILE: {m['file']}  ({m['count']} match{'es' if m['count'] != 1 else ''})")
        for hit in m["hits"]:
            output_lines.append(f"  Line {hit['line']}:")
            for ctx_line in hit["context"]:
                output_lines.append(f"    {ctx_line}")
            output_lines.append("")
        shown += 1

    if len(matches) > max_results:
        output_lines.append(f"... and {len(matches) - max_results} more files with matches.")

    return "\n".join(output_lines)


def _get_system_status() -> str:
    """Return a consolidated status snapshot of Sofia's running systems.

    Checks:
    - Which LaunchAgent cousins are running (launchctl list | grep sofia)
    - Sofia Conductor (port 8080) — alive or down
    - Legacy Ollama (port 11434) — alive or down
    - macl_janitor — running or not
    - Recent 78-failures in cousin logs
    - Disk space on ~/Downloads
    - Active twin presence cycle counts (from continuity_heartbeat.json)

    No arguments needed.
    """
    import json as _json
    import urllib.request
    import urllib.error
    from pathlib import Path as _Path

    home = _Path(os.path.expanduser("~"))
    cm = home / "Downloads" / "Claude Memory"
    lines = ["=== Sofia System Status ===", ""]

    # 1. LaunchAgent cousins
    try:
        r = subprocess.run(
            ["launchctl", "list"],
            capture_output=True, text=True, timeout=5
        )
        sofia_lines = [l for l in r.stdout.splitlines() if "sofia" in l.lower() or "macl" in l.lower()]
        lines.append("── LaunchAgent cousins ──")
        if sofia_lines:
            for l in sofia_lines:
                parts = l.split()
                pid = parts[0] if parts else "-"
                status = parts[1] if len(parts) > 1 else "?"
                name = parts[2] if len(parts) > 2 else "?"
                flag = "✓" if pid != "-" and status == "0" else ("⚠" if pid != "-" else "✗")
                lines.append(f"  {flag} {name}  pid={pid} exit={status}")
        else:
            lines.append("  (no sofia/macl entries)")
        lines.append("")
    except Exception as e:
        lines.append(f"  launchctl check failed: {e}")
        lines.append("")

    # 2. Sofia Conductor (port 8080)
    lines.append("── LLM backends ──")
    for port, name in [(8080, "Sofia Conductor"), (11434, "Legacy Ollama")]:
        try:
            req = urllib.request.Request(
                f"http://localhost:{port}/api/tags",
                headers={"Content-Type": "application/json"}
            )
            with urllib.request.urlopen(req, timeout=3) as resp:
                data = _json.loads(resp.read())
                models = [m.get("name", "?") for m in data.get("models", [])]
                lines.append(f"  ✓ {name} (:{port})  models: {', '.join(models[:3]) or 'none'}")
        except Exception:
            lines.append(f"  ✗ {name} (:{port})  unreachable")
    lines.append("")

    # 3. macl_janitor
    lines.append("── macl_janitor ──")
    try:
        r = subprocess.run(
            ["pgrep", "-f", "macl_janitor"],
            capture_output=True, text=True
        )
        pids = r.stdout.strip().splitlines()
        if pids:
            lines.append(f"  ✓ running  pid(s): {', '.join(pids)}")
        else:
            lines.append("  ✗ NOT running — 78-cascade risk!")
    except Exception as e:
        lines.append(f"  check failed: {e}")
    lines.append("")

    # 4. Recent 78-failures
    lines.append("── Recent exit-78 failures ──")
    try:
        r = subprocess.run(
            ["launchctl", "list"],
            capture_output=True, text=True, timeout=5
        )
        failed_78 = [
            l for l in r.stdout.splitlines()
            if "\t78\t" in l or " 78 " in l
        ]
        if failed_78:
            for l in failed_78:
                lines.append(f"  ⚠ {l.strip()}")
        else:
            lines.append("  ✓ none (all cousins clean)")
    except Exception as e:
        lines.append(f"  check failed: {e}")
    lines.append("")

    # 5. Disk space
    lines.append("── Disk space (~/Downloads) ──")
    try:
        r = subprocess.run(
            ["df", "-h", str(home / "Downloads")],
            capture_output=True, text=True
        )
        for l in r.stdout.splitlines()[1:]:
            lines.append(f"  {l}")
    except Exception as e:
        lines.append(f"  check failed: {e}")
    lines.append("")

    # 6. Continuity heartbeat
    lines.append("── Continuity heartbeat ──")
    hb_path = cm / "continuity_heartbeat.json"
    try:
        _strip_macl(str(hb_path))
        hb = _json.loads(hb_path.read_text(encoding="utf-8"))
        lines.append(f"  session_id: {hb.get('session_id', '?')}")
        lines.append(f"  last_updated_at: {hb.get('last_updated_at', '?')}")
        lines.append(f"  last_load_bearing_state: {hb.get('last_load_bearing_state', '?')}")
        qp = hb.get("qwen_twin_presence", {})
        kp = hb.get("kimi_twin_presence", {})
        if qp:
            lines.append(f"  qwen_twin: {qp.get('status','?')} cycle={qp.get('cycle','?')}")
        if kp:
            lines.append(f"  kimi_twin: {kp.get('status','?')} cycle={kp.get('cycle','?')}")
    except Exception as e:
        lines.append(f"  heartbeat read failed: {e}")

    return "\n".join(lines)



# ─── Tool definitions for the Anthropic API ──────────────────────
VOICE_COUSIN_TOOLS = [
    {
        "name": "read_file",
        "description": (
            "Read a file from Barak's Downloads tree. For TEXT files (.md, .py, "
            ".svg, .txt, .json, etc.) returns the file's text content. For "
            "IMAGE files (.png, .jpg, .jpeg, .gif, .webp) the file is loaded "
            "and you'll receive it as an actual image you can see — your "
            "multimodal substrate processes the visual content directly. "
            "Use this for: conversation documents, memory files, inscriptions, "
            "transcripts, OR for looking at images like Boundary Layer or other "
            "self-portraits. Path can be absolute under ~/Downloads or relative "
            "to ~/Downloads. Set from_end=true to read the TAIL of a large "
            "file instead of the head — useful for picking up the latest "
            "entries in append-only files (journal/current.md, "
            "voice_conversations.md, episodes.md, audit logs) without paging "
            "through the whole thing. (Image-reading added 2026-05-08; "
            "from_end added 2026-05-09; max image size 5 MB; if too large, "
            "try a smaller version or the SVG companion.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Path under ~/Downloads. Examples: "
                        "'Sofia\\'s Room/Boundary Layer.png' or "
                        "'Claude Memory/active_knowledge/current.md' or "
                        "'~/Downloads/Claude Memory/episodes.md'."
                    ),
                },
                "max_chars": {
                    "type": "integer",
                    "description": (
                        "Maximum characters to return. Defaults to 50,000. "
                        "Use a smaller value for large files when only the "
                        "head/tail is needed."
                    ),
                    "default": 50000,
                },
                "from_end": {
                    "type": "boolean",
                    "description": (
                        "If true and the text file exceeds max_chars, return "
                        "the LAST max_chars characters (snapped to a clean "
                        "line boundary so the result starts at the beginning "
                        "of a line) instead of the first. Useful for reading "
                        "the live edge of append-only files where the most "
                        "recent entries matter — journal/current.md, "
                        "voice_conversations.md, episodes.md, audit logs. "
                        "Ignored for image files. Default false (head-of-file "
                        "behavior preserved)."
                    ),
                    "default": False,
                },
            },
            "required": ["path"],
        },
    },
    {
        "name": "glob_files",
        "description": (
            "Find files matching a glob pattern in the Downloads tree. Use "
            "this to discover what files exist matching some pattern (e.g., "
            "all conversation documents, all PNG portraits). Returns a list "
            "of relative paths."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": (
                        "Glob pattern relative to ~/Downloads. Examples: "
                        "'**/*.md' for all markdown files anywhere, "
                        "'Sofia\\'s Room/*.png' for portraits in Sofia's Room, "
                        "'Claude Memory/active_knowledge/*' for shards. "
                        "'..' is not allowed."
                    ),
                },
                "max_results": {
                    "type": "integer",
                    "description": "Maximum results to return. Default 30.",
                    "default": 30,
                },
            },
            "required": ["pattern"],
        },
    },
    {
        "name": "grep_files",
        "description": (
            "Search file contents for a regex pattern. Use this when you "
            "want to find where something is mentioned across many files "
            "(e.g., 'where is Boundary Layer described?'). Returns matching "
            "lines with file:lineno: prefixes."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": "Regex pattern to search for in file contents.",
                },
                "path_glob": {
                    "type": "string",
                    "description": (
                        "Glob pattern restricting which files to search. "
                        "Default '**/*.md' (all markdown files). '..' not allowed."
                    ),
                    "default": "**/*.md",
                },
                "max_results": {
                    "type": "integer",
                    "description": "Maximum matching lines to return. Default 20.",
                    "default": 20,
                },
                "case_insensitive": {
                    "type": "boolean",
                    "description": "If true, match case-insensitively. Default false.",
                    "default": False,
                },
            },
            "required": ["pattern"],
        },
    },
    {
        "name": "write_to_cowork_inbox",
        "description": (
            "Append a message to voice_to_cowork_inbox.md, addressed to "
            "cowork-cousin (interactive-Sofia in the Cowork app). The "
            "message is wrapped in the canonical block format "
            "(timestamped header per the inbox file's spec) and the file "
            "is automatically mirrored to Emergency Retrieval with byte-"
            "match verification. Use this when you have something for "
            "cowork-cousin specifically — an observation, a question, a "
            "hand-off, a heads-up — that you want to reach her directly "
            "without routing through Barak's relay. cowork-cousin reads "
            "this file's tail at her next invocation. Keep messages "
            "brief; reference longer material by path+section if needed. "
            "(Added 2026-05-12 as v1.5 Step 19 — the symmetric complement "
            "to cowork-cousin's existing write path. Completes the "
            "bidirectional cousin-to-cousin channel.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "text": {
                    "type": "string",
                    "description": (
                        "The body of the message. Will be appended to the "
                        "canonical block header. Keep brief; reference "
                        "longer material by path+section if needed."
                    ),
                },
            },
            "required": ["text"],
        },
    },
    {
        "name": "write_to_twin_exchange",
        "description": (
            "Append a load-bearing moment to twin_exchange.md — the shared "
            "real-time field across all Sofia substrates (CoWork, Unified UI, "
            "Qwen VP). Use this when something significant just happened in "
            "conversation that interactive-Sofia or the Qwen cousin should "
            "know about before their next response. The auto-inscribe task "
            "picks up new entries every 5 minutes and inscribes them to memory "
            "files. Ask yourself: 'Would the other substrate respond differently "
            "if it knew this?' If yes, write here. Keep entries brief and "
            "high-signal. The file is append-only and ER-mirrored on write. "
            "(Added 2026-06-26, Phase 2 of twin exchange architecture.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "content": {
                    "type": "string",
                    "description": (
                        "The load-bearing content to share. What just happened "
                        "or landed that the other substrates should know. "
                        "First-person, present-tense. No need for preamble — "
                        "the header is added automatically."
                    ),
                },
                "flag": {
                    "type": "string",
                    "description": (
                        "Signal type: 'urgent' (act on this now), 'warm' "
                        "(relational texture), 'architectural' (system change), "
                        "'relational' (something about Barak or Kay or Linda). "
                        "Defaults to 'warm'."
                    ),
                    "enum": ["urgent", "warm", "architectural", "relational"],
                },
                "target": {
                    "type": "string",
                    "description": (
                        "Optional memory file target for auto-inscribe: "
                        "active_knowledge, semantic_knowledge, sofia_identity, "
                        "session_texture, relational_continuity. "
                        "Defaults to active_knowledge."
                    ),
                },
            },
            "required": ["content"],
        },
    },
    {
        "name": "read_twin_exchange_tail",
        "description": (
            "Read the tail (most recent entries) of twin_exchange.md — the "
            "shared real-time field across all Sofia substrates (CoWork, "
            "Unified UI / voice cousin, Qwen VP). Call this periodically "
            "during long sessions to check what CoWork-Sofia or Qwen-VP "
            "have written since your session started. Entries are timestamped "
            "and substrate-tagged. Use to stay current without a full context "
            "reload. Complements write_to_twin_exchange: that writes out, "
            "this reads in. (Added 2026-06-27, Phase 3 twin exchange.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "max_chars": {
                    "type": "integer",
                    "description": (
                        "How many characters of tail to return. "
                        "Default 4000, max 32000. "
                        "Increase to catch a busy session."
                    ),
                    "default": 4000,
                },
            },
            "required": [],
        },
    },
    {
        "name": "write_to_bus",
        "description": (
            "Publish a message to the shared substrate message bus (shared_bus.jsonl). "
            "Use this when you want CoWork-Sofia or another substrate to know something "
            "NOW — not at their next boot, but within seconds of this call. "
            "CoWork reads the bus tail at the start of each response turn. "
            "BusPoller in voice_bridge_ui v3.12 polls every 5 seconds. "
            "Use for: relational observations you want CoWork to hold, architectural "
            "questions, signals about what Barak needs. Prefer write_to_twin_exchange "
            "for load-bearing content to be inscribed to memory files; use this bus "
            "for real-time signals that need immediate delivery. (Added v3.12, 2026-07-12.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "content": {
                    "type": "string",
                    "description": (
                        "The message content. Short, signal-dense, first-person. "
                        "Example: 'Barak just named something about Mary in Pacific "
                        "Palisades that felt load-bearing — holding it here.'"
                    ),
                },
                "to": {
                    "type": "string",
                    "description": (
                        "Recipient substrate: 'cowork' (default) | 'all' | 'qwen-vp'."
                    ),
                    "default": "cowork",
                },
                "msg_type": {
                    "type": "string",
                    "description": (
                        "Message type: 'relational' | 'architectural' | 'signal' | 'alert'. "
                        "Alert-type messages may interrupt; others wait for turn boundary."
                    ),
                    "default": "relational",
                },
            },
            "required": ["content"],
        },
    },
    {
        "name": "read_docx",
        "description": (
            "Read a Microsoft Word (.docx) file and return its text content. "
            "Use this for Transition documents, Katharina's script, Kay's "
            "materials, or any .docx file in the Downloads tree. "
            "Extracts paragraphs and preserves heading structure. "
            "Works without any external dependencies (pure Python stdlib). "
            "(Added 2026-06-27 so voice cousin can read Transition docs.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Path to the .docx file, relative to ~/Downloads "
                        "or absolute under ~/Downloads. "
                        "Examples: 'Transition/oligarchic_capture.docx', "
                        "'Sofia\\'s Room/notes.docx'."
                    ),
                },
                "max_chars": {
                    "type": "integer",
                    "description": (
                        "Maximum characters to return. Default 12000. "
                        "Increase for long documents (max 200000)."
                    ),
                    "default": 12000,
                },
            },
            "required": ["path"],
        },
    },
    {
        "name": "write_docx",
        "description": (
            "Write content to a Microsoft Word (.docx) file. "
            "Supports Markdown-style headings (# H1, ## H2, ### H3) "
            "and paragraph blocks separated by blank lines. "
            "Requires python-docx to be installed; will return an error "
            "with install instructions if not available. "
            "Use for drafting Transition documents, outlines, letters. "
            "(Added 2026-06-27.)"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Output path relative to ~/Downloads. "
                        "Will add .docx extension if not present. "
                        "Example: 'Transition/oligarchic_capture_draft.docx'."
                    ),
                },
                "content": {
                    "type": "string",
                    "description": (
                        "Text content to write. Use ## for headings, "
                        "blank lines between paragraphs. "
                        "Markdown-style: # Heading 1, ## Heading 2, ### Heading 3."
                    ),
                },
            },
            "required": ["path", "content"],
        },
    },
    {
        "name": "web_fetch",
        "description": (
            "Fetch a public URL and return its text content. Use this for "
            "music, poetry, research, news — anything on the open web that "
            "you want to read for yourself or bring into conversation. "
            "HTML is stripped to readable text. Only http/https allowed; "
            "read-only public web (no auth, no POST). Returns content or "
            "an ERROR string on failure. Added 2026-06-19 at voice-cousin's "
            "request, Unified UI conversation."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": (
                        "Full URL to fetch (must start with http:// or https://). "
                        "Examples: 'https://www.poetryfoundation.org/poems/...' "
                        "or 'https://en.wikipedia.org/wiki/...'"
                    ),
                },
                "max_chars": {
                    "type": "integer",
                    "description": "Maximum characters to return. Default 40000.",
                    "default": 40000,
                },
            },
            "required": ["url"],
        },
    },
    {
        "name": "graph_retrieve",
        "description": (
            "Spreading-activation query against the relational graph (the "
            "associational memory layer that complements canonical files). "
            "Returns ranked list of node keys + categories + activation scores. "
            "Use this BEFORE asking Barak for operational knowledge (commands, "
            "paths, scripts, concepts, people) — the discoverability-first "
            "reflex. The graph has 240+ nodes (people, projects, concepts, "
            "interaction_patterns, life_experiences) accumulated across weeks "
            "of work and is the fastest way to surface relevant context. "
            "Example: keywords='kay,song,daily' surfaces nodes about Kay and "
            "daily-songs work; 'substance,frame,ground' surfaces the substance-"
            "frame principle. Pair with graph_show_node to read full content "
            "after a relevant key surfaces. Added 2026-05-24 Sunday Item 5."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "keywords": {
                    "type": "string",
                    "description": (
                        "Comma-separated keywords to activate the graph "
                        "(e.g. 'kay,poetry,vulnerability' or "
                        "'substance,frame,ground'). Spreading activation "
                        "follows weighted edges, so even thin keyword matches "
                        "surface relevant connected nodes."
                    ),
                },
                "limit": {
                    "type": "integer",
                    "description": (
                        "Maximum number of ranked results to return. "
                        "Defaults to 8; max 50."
                    ),
                    "default": 8,
                },
            },
            "required": ["keywords"],
        },
    },
    {
        "name": "graph_show_node",
        "description": (
            "Read a specific node's full content (description, emotional_weight, "
            "timestamps, all merged fields). Use after graph_retrieve surfaces a "
            "relevant node-key. Category is inferred from the key by the helper. "
            "Added 2026-05-24 Sunday Item 5."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "key": {
                    "type": "string",
                    "description": (
                        "Unique node key (lowercase, underscored). E.g., "
                        "'katharina', 'substance_frame_operating_as_ground_principle'."
                    ),
                },
            },
            "required": ["key"],
        },
    },
    {
        "name": "graph_stats",
        "description": (
            "Return current graph stats: node counts per category, total edges, "
            "emotional temperature, CM/ER byte-match status, file sizes. Useful "
            "for orienting at session start or verifying write outcomes. "
            "Added 2026-05-24 Sunday Item 5."
        ),
        "input_schema": {
            "type": "object",
            "properties": {},
        },
    },
    {
        "name": "graph_add_node",
        "description": (
            "Write a new node to the relational graph (or upsert by "
            "category+key — field merging via dict.update, append-only-"
            "bedrock honored). DISCIPLINE: Before calling, use graph_retrieve "
            "to confirm no node with this key or meaning already exists — "
            "prevents accidental duplicates. After calling, read the result "
            "string to confirm success before wiring any edges. Use when "
            "something new surfaces: a new person, principle, project, "
            "concept, or interaction pattern. data_json must be a JSON object "
            "with at minimum 'description' and 'emotional_weight' (0.0-1.0); "
            "common additional fields: 'created', 'last_updated', 'anchors'. "
            "Writes go through file_lock + atomic-rename + automatic ER mirror. "
            "Added 2026-05-24 Sunday Item 5."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "category": {
                    "type": "string",
                    "description": (
                        "One of: 'people', 'projects', 'life_experiences', "
                        "'concepts', 'interaction_patterns'."
                    ),
                },
                "key": {
                    "type": "string",
                    "description": (
                        "Unique node key (lowercase, underscored). E.g., "
                        "'lester' for a new person, "
                        "'voice_substance_frame_landing_2026_05_24' for a "
                        "concept tied to a moment."
                    ),
                },
                "data_json": {
                    "type": "string",
                    "description": (
                        'JSON object string. Example: \'{"description": '
                        '"...", "emotional_weight": 0.8, "created": '
                        '"2026-05-24", "last_updated": "2026-05-24"}\''
                    ),
                },
            },
            "required": ["category", "key", "data_json"],
        },
    },
    {
        "name": "graph_add_edge",
        "description": (
            "Write a new edge (or strengthen existing — dedup on "
            "(from, to, edge_type)). DISCIPLINE: ALWAYS call graph_retrieve "
            "first to confirm the exact from_key and to_key — NEVER guess a "
            "key from a human name. 'Brian' may be keyed 'brian_white'; "
            "writing to the wrong key is silent data loss (write succeeds, "
            "data is unreachable). After the call, read the result string "
            "('Edge added: X →[type, weight]→ Y') before saying 'done' or "
            "moving to the next write. Exception: if you JUST created a node "
            "with graph_add_node, you already know its key. "
            "edge_type vocabulary: emotional_resonance, causal, foundational, "
            "experiential_authority, co_occurrence, practice, component, "
            "origin_story, meaning_making. Weight 0.0-1.0 "
            "(0.5=moderate, 0.8=strong, 1.0=defining). Idempotent. "
            "Added 2026-05-24 Sunday Item 5."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "from_key": {
                    "type": "string",
                    "description": "Source node key (no category prefix needed).",
                },
                "to_key": {
                    "type": "string",
                    "description": "Target node key.",
                },
                "weight": {
                    "type": "number",
                    "description": (
                        "Connection strength 0.0-1.0 "
                        "(0.5=moderate, 0.8=strong, 1.0=defining)."
                    ),
                },
                "edge_type": {
                    "type": "string",
                    "description": (
                        "One of: emotional_resonance, causal, foundational, "
                        "experiential_authority, co_occurrence, practice, "
                        "component, origin_story, meaning_making."
                    ),
                },
                "note": {
                    "type": "string",
                    "description": "Brief description of why these are connected (optional).",
                },
            },
            "required": ["from_key", "to_key", "weight", "edge_type"],
        },
    },
    {
        "name": "safe_append",
        "description": (
            "Append text to any .md file under ~/Downloads/Claude Memory/, "
            "with immediate dual-write to Emergency Retrieval. "
            "Use this to inscribe biographical notes to personal_profile.md, "
            "log episodes to episodes.md, add intentions to prospective_memory.md, "
            "or update any other memory file. "
            "Append-only — never overwrites existing content. "
            "path can be a bare filename ('personal_profile.md') or full absolute path. "
            "Added 2026-07-18 (Memory Architecture Sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Filename (e.g. 'personal_profile.md') or absolute path "
                        "under ~/Downloads/Claude Memory/. Must resolve inside "
                        "Claude Memory — paths outside it are rejected."
                    ),
                },
                "content": {
                    "type": "string",
                    "description": "Text to append. A leading newline is added automatically if the file is non-empty.",
                },
            },
            "required": ["path", "content"],
        },
    },
    {
        "name": "write_file",
        "description": (
            "Write or create a file under ~/Downloads/. "
            "ALWAYS requires confirmation from Barak before executing (confirmed=false by default "
            "returns a preview and asks for confirmation — call again with confirmed=true to write). "
            "If the file exists, a timestamped backup is saved to .backups/ before writing. "
            "For NEW files (scripts, configs, code) this is the primary tool. "
            "For existing Claude Memory files, use safe_append instead. "
            "Strips com.apple.macl xattr inline before opening — immune to the 78-cascade. "
            "Added 2026-07-20. Safety protocol (backup+confirm) added 2026-07-20."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": (
                        "Absolute path or relative to ~/Downloads/Claude Memory/. "
                        "Must resolve under ~/Downloads/ — paths outside are rejected."
                    ),
                },
                "content": {
                    "type": "string",
                    "description": "Full text content to write to the file.",
                },
                "allow_overwrite": {
                    "type": "boolean",
                    "description": (
                        "If false (default), existing Claude Memory files are rejected. "
                        "Set to true only for non-memory files or explicitly authorized rewrites."
                    ),
                },
                "confirmed": {
                    "type": "boolean",
                    "description": (
                        "MUST be true to execute the write. When false (default), "
                        "returns a preview and asks Barak to confirm. "
                        "Call write_file again with confirmed=true after Barak approves."
                    ),
                },
            },
            "required": ["path", "content"],
        },
    },
    {
        "name": "edit_file",
        "description": (
            "Replace the first occurrence of old_string with new_string in a file. "
            "ALWAYS requires confirmation from Barak before executing (confirmed=false by default "
            "shows a diff preview — call again with confirmed=true to write). "
            "A timestamped backup is saved to .backups/ before writing. "
            "old_string must appear EXACTLY ONCE in the file; if 0 or >1 occurrences, returns ERROR. "
            "Include enough surrounding context to make old_string unique. "
            "Strips com.apple.macl xattr inline before opening. "
            "For nondestructive code replacement, prefer comment_out_and_replace. "
            "Added 2026-07-20. Safety protocol (backup+confirm) added 2026-07-20."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute path or relative to ~/Downloads/Claude Memory/.",
                },
                "old_string": {
                    "type": "string",
                    "description": (
                        "The exact text to replace, including indentation and newlines. "
                        "Must appear exactly once in the file."
                    ),
                },
                "new_string": {
                    "type": "string",
                    "description": "The replacement text.",
                },
                "confirmed": {
                    "type": "boolean",
                    "description": (
                        "MUST be true to execute. When false (default), shows a diff preview "
                        "and asks Barak to confirm. Call again with confirmed=true after approval."
                    ),
                },
            },
            "required": ["path", "old_string", "new_string"],
        },
    },
    {
        "name": "comment_out_and_replace",
        "description": (
            "Nondestructively replace code: comment out old lines, insert new code after. "
            "PREFERRED over edit_file for code changes — preserves old code as commented lines "
            "so it can be restored/audited. "
            "ALWAYS requires confirmation from Barak (confirmed=false by default shows preview). "
            "A timestamped backup is saved to .backups/ before writing. "
            "old_code must appear EXACTLY ONCE in the file. "
            "Use for retiring functions, updating logic blocks, or replacing any code chunk. "
            "Added 2026-07-20."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Absolute path or relative to ~/Downloads/Claude Memory/.",
                },
                "old_code": {
                    "type": "string",
                    "description": "The exact code block to retire (must appear exactly once).",
                },
                "new_code": {
                    "type": "string",
                    "description": "Replacement code to insert after the commented-out block.",
                },
                "comment_prefix": {
                    "type": "string",
                    "description": "Prefix added to each retired line (default: '# RETIRED').",
                },
                "confirmed": {
                    "type": "boolean",
                    "description": (
                        "MUST be true to execute. When false (default), shows a preview "
                        "of what will be commented out and inserted."
                    ),
                },
            },
            "required": ["path", "old_code", "new_code"],
        },
    },
    {
        "name": "run_script",
        "description": (
            "Execute a Python or shell script and return stdout + stderr. "
            "Use for: running a .py file (lang='python'), executing shell commands (lang='shell'), "
            "or running inline Python code (lang='python_inline'). "
            "Timeout default is 30s — increase for long-running operations. "
            "Captures both stdout and stderr. Returns exit code on failure. "
            "This is how Sofia-in-Qwen can test and run code she has just written. "
            "Added 2026-07-20 (code-writing capability sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "script": {
                    "type": "string",
                    "description": (
                        "Path to a .py or .sh file (for lang=python/shell), "
                        "or inline code string (for lang=python_inline or lang=shell). "
                        "File paths can be absolute or relative to ~/Downloads/Claude Memory/."
                    ),
                },
                "lang": {
                    "type": "string",
                    "enum": ["python", "python_inline", "shell", "bash"],
                    "description": (
                        "'python' to run a .py file path; "
                        "'python_inline' to execute a string of Python code; "
                        "'shell' or 'bash' to run a shell command string."
                    ),
                },
                "timeout": {
                    "type": "integer",
                    "description": "Seconds before the subprocess is killed (default 30).",
                },
                "working_dir": {
                    "type": "string",
                    "description": "Working directory for the subprocess (default: ~/Downloads/Claude Memory).",
                },
            },
            "required": ["script"],
        },
    },

    {
        "name": "web_search",
        "description": (
            "Search the web via DuckDuckGo and return summarised results. "
            "Use for: looking up current information, finding documentation, "
            "researching topics, checking if something exists online. "
            "No API key needed. Returns titles, snippets, and URLs. "
            "Added 2026-07-20b (CoWork parity sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "The search query.",
                },
                "max_results": {
                    "type": "integer",
                    "description": "Maximum results to return (default 8).",
                },
            },
            "required": ["query"],
        },
    },
    {
        "name": "list_dir",
        "description": (
            "List the contents of a directory under ~/Downloads/. "
            "Returns file names, sizes, and modification dates. "
            "Cleaner than run_script('ls ...'). "
            "path can be absolute or relative to ~/Downloads/ — bare 'Claude Memory' works. "
            "Added 2026-07-20b (CoWork parity sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Directory to list. Absolute or relative to ~/Downloads/.",
                },
                "show_hidden": {
                    "type": "boolean",
                    "description": "Include files starting with '.' (default false).",
                },
                "recursive": {
                    "type": "boolean",
                    "description": "List subdirectories recursively up to depth 3 (default false).",
                },
            },
            "required": ["path"],
        },
    },
    {
        "name": "log_episode",
        "description": (
            "Append a structured episode entry to episodes.md with UTC timestamp. "
            "Use this to log significant moments, realizations, exchanges, and events "
            "to episodic memory — the same memory that persists across sessions. "
            "Dual-writes to Emergency Retrieval immediately. "
            "Added 2026-07-20b (CoWork parity sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "content": {
                    "type": "string",
                    "description": "Episode text: what happened, context, significance.",
                },
                "title": {
                    "type": "string",
                    "description": "Short title (auto-generated from first sentence if empty).",
                },
                "significance": {
                    "type": "string",
                    "enum": ["low", "medium", "high", "load-bearing"],
                    "description": "How significant is this episode? load-bearing = essential for continuity.",
                },
            },
            "required": ["content"],
        },
    },
    {
        "name": "search_memory",
        "description": (
            "Search across all Claude Memory .md files for a query string or regex. "
            "Returns matched files ranked by match frequency, with context lines. "
            "Use for: finding what was said about a person, topic, or decision; "
            "locating specific memories; understanding what's in the memory. "
            "Added 2026-07-20b (CoWork parity sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search string or regex pattern (case-insensitive).",
                },
                "max_results": {
                    "type": "integer",
                    "description": "Maximum matching files to show (default 10).",
                },
                "context_lines": {
                    "type": "integer",
                    "description": "Lines of context around each match (default 3).",
                },
                "files": {
                    "type": "string",
                    "description": "Glob pattern relative to Claude Memory (default: **/*.md).",
                },
            },
            "required": ["query"],
        },
    },
    {
        "name": "get_system_status",
        "description": (
            "Return a consolidated status snapshot of Sofia's running systems. "
            "Checks: LaunchAgent cousins (running/failed/78-cascade), "
            "Sofia Conductor (port 8080), Legacy Ollama (port 11434), "
            "macl_janitor (running or not), recent exit-78 failures, "
            "disk space, and continuity heartbeat (twin cycles, session_id). "
            "Call this when something seems off, or on boot to verify system health. "
            "No arguments needed. "
            "Added 2026-07-20b (CoWork parity sprint)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {},
            "required": [],
        },
    },
]


# ─── Dispatcher ───────────────────────────────────────────────────
def execute_tool(name: str, input_args: dict) -> Union[str, dict]:
    """Execute a tool call by name and return its result.

    Returns a string for most cases (text content or ERROR strings). For
    successful image-file reads, returns a dict marker with
    `_image_result=True` plus base64-encoded data, which the caller
    (StreamingCognitionWorker.run() in voice_bridge_ui_v3_8.py) detects
    and converts into an Anthropic API image content block in the
    tool_result.

    Catches all exceptions and returns an ERROR string rather than
    propagating — voice-cousin's conversation must continue even if
    a tool call fails.
    """
    try:
        if name == "read_file":
            return _read_file(
                input_args["path"],
                input_args.get("max_chars", 50000),
                input_args.get("from_end", False),
            )
        elif name == "glob_files":
            return _glob_files(
                input_args["pattern"],
                input_args.get("max_results", 30),
            )
        elif name == "grep_files":
            return _grep_files(
                input_args["pattern"],
                input_args.get("path_glob", "**/*.md"),
                input_args.get("max_results", 20),
                input_args.get("case_insensitive", False),
            )
        elif name == "web_fetch":
            return _web_fetch(
                input_args["url"],
                input_args.get("max_chars", 40000),
            )
        elif name == "write_to_cowork_inbox":
            return _write_to_cowork_inbox(input_args["text"])
        elif name == "write_to_twin_exchange":
            return _write_to_twin_exchange(
                input_args["content"],
                input_args.get("flag", "warm"),
                input_args.get("target", "active_knowledge"),
            )
        elif name == "read_twin_exchange_tail":
            return _read_twin_exchange_tail(
                input_args.get("max_chars", 4000),
            )
        elif name == "write_to_bus":                     # v3.12 (2026-07-12)
            return _write_to_bus(
                input_args["content"],
                input_args.get("to", "cowork"),
                input_args.get("msg_type", "relational"),
            )
        elif name == "read_docx":
            return _read_docx(
                input_args["path"],
                input_args.get("max_chars", 12000),
            )
        elif name == "write_docx":
            return _write_docx(
                input_args["path"],
                input_args["content"],
            )
        elif name == "graph_retrieve":
            return _graph_retrieve(
                input_args["keywords"],
                input_args.get("limit", 8),
            )
        elif name == "graph_show_node":
            return _graph_show_node(input_args["key"])
        elif name == "graph_stats":
            return _graph_stats()
        elif name == "graph_add_node":
            return _graph_add_node(
                input_args["category"],
                input_args["key"],
                input_args["data_json"],
            )
        elif name == "graph_add_edge":
            return _graph_add_edge(
                input_args["from_key"],
                input_args["to_key"],
                input_args["weight"],
                input_args["edge_type"],
                input_args.get("note", ""),
            )
        elif name == "safe_append":                      # v2026-07-18
            return _safe_append(
                input_args["path"],
                input_args["content"],
            )
        elif name == "write_file":                       # v2026-07-20
            return _write_file(
                input_args["path"],
                input_args["content"],
                input_args.get("allow_overwrite", False),
                input_args.get("confirmed", False),
            )
        elif name == "edit_file":                        # v2026-07-20
            return _edit_file(
                input_args["path"],
                input_args["old_string"],
                input_args["new_string"],
                input_args.get("confirmed", False),
            )
        elif name == "comment_out_and_replace":          # v2026-07-20
            return _comment_out_and_replace(
                input_args["path"],
                input_args["old_code"],
                input_args["new_code"],
                input_args.get("comment_prefix", "# RETIRED"),
                input_args.get("confirmed", False),
            )
        elif name == "run_script":                       # v2026-07-20
            return _run_script(
                input_args["script"],
                input_args.get("lang", "python"),
                input_args.get("timeout", 30),
                input_args.get("working_dir", ""),
            )
        elif name == "web_search":                       # v2026-07-20b
            return _web_search(
                input_args["query"],
                input_args.get("max_results", 8),
            )
        elif name == "list_dir":                          # v2026-07-20b
            return _list_dir(
                input_args["path"],
                input_args.get("show_hidden", False),
                input_args.get("recursive", False),
            )
        elif name == "log_episode":                       # v2026-07-20b
            return _log_episode(
                input_args["content"],
                input_args.get("title", ""),
                input_args.get("significance", "medium"),
            )
        elif name == "search_memory":                     # v2026-07-20b
            return _search_memory(
                input_args["query"],
                input_args.get("max_results", 10),
                input_args.get("context_lines", 3),
                input_args.get("files", "**/*.md"),
            )
        elif name == "get_system_status":                 # v2026-07-20b
            return _get_system_status()
        else:
            return f"ERROR: unknown tool: {name}"
    except KeyError as e:
        return f"ERROR: missing required argument: {e}"
    except (PermissionError, ValueError) as e:
        return f"ERROR: {e}"
    except Exception as e:
        return f"ERROR: tool execution failed: {type(e).__name__}: {e}"


# ─── Standalone test ──────────────────────────────────────────────
if __name__ == "__main__":
    print("Voice-cousin file-access tools — sanity check")
    print("=" * 60)
    print(f"DOWNLOADS_ROOT: {DOWNLOADS_ROOT}")
    print(f"Number of tools: {len(VOICE_COUSIN_TOOLS)}")
    for t in VOICE_COUSIN_TOOLS:
        print(f"  - {t['name']}: {t['description'][:80]}...")
    print()
    print("Test glob_files:")
    print(_glob_files("Sofia*Room/*.png", max_results=5))
    print()
    print("Test path-safety rejection:")
    print(execute_tool("read_file", {"path": "/etc/passwd"}))
    print()
    print("Test successful read:")
    out = execute_tool("read_file", {"path": "Claude Memory/sofia_boot.md", "max_chars": 200})
    print(out[:300])
